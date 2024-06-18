import datetime

from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtGui import *
from PyQt5.QtCore import *
from PyQt5.QtWidgets import *
from window_portal import *
import sys
import os
from sqlite3 import connect
import time
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, Mm, Length, RGBColor, Cm
from subprocess import call
from selenium import webdriver
from selenium.webdriver.common.by import By
import shutil
from selenium.webdriver.common.keys import Keys
from selenium.webdriver import FirefoxOptions



class Ui_MainWindow(object):
    number=0
    def setupUi(self, MainWindow):
        if Ui_MainWindow.number==0:
                Ui_MainWindow.number=1
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(int(960*suma3/Ui_MainWindow.number), int(618*suma3/Ui_MainWindow.number))
        MainWindow.setMinimumSize(QtCore.QSize(int(400*suma3/Ui_MainWindow.number), int(0*suma3/Ui_MainWindow.number)))
        MainWindow.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.gridLayout_2 = QtWidgets.QGridLayout(self.centralwidget)
        self.gridLayout_2.setObjectName("gridLayout_2")
        self.tableWidget = QtWidgets.QTableWidget(self.centralwidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Expanding)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.tableWidget.sizePolicy().hasHeightForWidth())
        self.tableWidget.setSizePolicy(sizePolicy)
        self.tableWidget.setMinimumSize(QtCore.QSize(int(567*suma3/Ui_MainWindow.number), int(0*suma3/Ui_MainWindow.number)))
        self.tableWidget.setMaximumSize(QtCore.QSize(int(2000*suma3/Ui_MainWindow.number), int(26777215*suma3/Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.tableWidget.setFont(font)
        self.tableWidget.setStyleSheet("QTableWidget{\n"
f"    border: {2*suma3}px solid;\n"
"    border-color: rgb(41, 213, 202);\n"
f"    border-radius: {2*suma3}%;\n"
"}\n"
"\n"
"QTableWidget::item:hover{\n"
"    background-color: rgb(41, 213, 202);\n"
"}")
        self.tableWidget.setLineWidth(1)
        self.tableWidget.setColumnCount(2)
        self.tableWidget.setObjectName("tableWidget")
        self.tableWidget.verticalHeader().setStyleSheet("QHeaderView { font-size: 10pt; }")

        item = QtWidgets.QTableWidgetItem()
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        font.setBold(True)
        font.setWeight(75)
        item.setFont(font)
        self.tableWidget.setHorizontalHeaderItem(0, item)
        item = QtWidgets.QTableWidgetItem()
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        font.setBold(True)
        font.setWeight(75)
        item.setFont(font)

        self.tableWidget.setHorizontalHeaderItem(1, item)
        # self.tableWidget.horizontalHeader().resizeSection(1, int(450))


        table = self.tableWidget
        header = table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.Stretch)
        header.setSectionResizeMode(1,  QtWidgets.QHeaderView.ResizeToContents)

        self.gridLayout_2.addWidget(self.tableWidget, int(0), int(0), int(1), int(1))
        self.widget = QtWidgets.QWidget(self.centralwidget)
        self.widget.setMinimumSize(QtCore.QSize(int(69*suma3/Ui_MainWindow.number), int(600*suma3/Ui_MainWindow.number)))
        self.widget.setMaximumWidth(int(69*suma3/Ui_MainWindow.number))


        self.widget.setStyleSheet(f"border-radius: {2*suma3/Ui_MainWindow.number}%;\n"
"background-color: rgb(41, 213, 202);")
        self.widget.setObjectName("widget")
        self.verticalLayout_4 = QtWidgets.QVBoxLayout(self.widget)
        self.verticalLayout_4.setObjectName("verticalLayout_4")
        self.add_menu = QtWidgets.QPushButton(self.widget)
        self.add_menu.setIcon(QIcon('1.png'))
        self.add_menu.setIconSize(QSize(int(45 * suma3/Ui_MainWindow.number), int(45 * suma3/Ui_MainWindow.number)))
        self.add_menu.setMinimumSize(QtCore.QSize(int(55*suma3/Ui_MainWindow.number), int(55*suma3/Ui_MainWindow.number)))
        self.add_menu.setStyleSheet("QPushButton{\n"
f"    border: 1px solid;\n"
"    background-color: rgb(255, 255, 255);\n"
f"    border-radius: {27*suma3/Ui_MainWindow.number}%;\n"
"    \n"
"}\n"
"QPushButton:hover{\n"
"    \n"
"    background-color: rgb(223, 223, 223);\n"
"}")
        self.add_menu.setText("")
        self.add_menu.setObjectName("add_menu")
        self.verticalLayout_4.addWidget(self.add_menu)
        self.add_metoduk = QtWidgets.QPushButton(self.widget)
        self.add_metoduk.setMinimumSize(QtCore.QSize(int(55*suma3/Ui_MainWindow.number), int(55*suma3/Ui_MainWindow.number)))
        self.add_metoduk.setIcon(QIcon('2.png'))
        self.add_metoduk.setIconSize(QSize(int(50 * suma3/Ui_MainWindow.number), int(50 * suma3/Ui_MainWindow.number)))
        self.add_metoduk.setStyleSheet("QPushButton{\n"
"    border: 1px solid;\n"
"    background-color: rgb(255, 255, 255);\n"
f"    border-radius: {27*suma3/Ui_MainWindow.number}%;\n"
"    \n"
"}\n"
"QPushButton:hover{\n"
"    background-color: rgb(223, 223, 223);\n"
"}")
        self.add_metoduk.setText("")
        self.add_metoduk.setObjectName("add_metoduk")
        self.verticalLayout_4.addWidget(self.add_metoduk)
        self.message = QtWidgets.QPushButton(self.widget)
        self.message.setMinimumSize(QtCore.QSize(int(55*suma3/Ui_MainWindow.number), int(55*suma3/Ui_MainWindow.number)))
        self.message.setIcon(QIcon('chenge.png'))
        self.message.setIconSize(QSize(int(40 * suma3/Ui_MainWindow.number), int(40 * suma3/Ui_MainWindow.number)))
        self.message.setStyleSheet("QPushButton{\n"
"    border: 1px solid;\n"
"    background-color: rgb(255, 255, 255);\n"
f"    border-radius: {27*suma3/Ui_MainWindow.number}%;\n"
"    \n"
"}\n"
"QPushButton:hover{\n"
"    background-color: rgb(223, 223, 223);\n"
"}")
        self.message.setText("")
        self.message.setObjectName("message")
        self.verticalLayout_4.addWidget(self.message)

        self.add_post = QtWidgets.QPushButton(self.widget)
        self.add_post.setMinimumSize(
            QtCore.QSize(int(55 * suma3 / Ui_MainWindow.number), int(55 * suma3 / Ui_MainWindow.number)))
        self.add_post.setIcon(QIcon('ж.png'))
        self.add_post.setIconSize(QSize(int(40 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        self.add_post.setStyleSheet("QPushButton{\n"
                                "    border: 1px solid;\n"
                                "    background-color: rgb(255, 255, 255);\n"
                                f"    border-radius: {27 * suma3 / Ui_MainWindow.number}%;\n"
                                "    \n"
                                "}\n"
                                "QPushButton:hover{\n"
                                "    background-color: rgb(223, 223, 223);\n"
                                "}")
        self.add_post.setText("")
        self.add_post.setObjectName("save")
        self.verticalLayout_4.addWidget(self.add_post)


        self.save = QtWidgets.QPushButton(self.widget)
        self.save.setMinimumSize(QtCore.QSize(int(55*suma3/Ui_MainWindow.number), int(55*suma3/Ui_MainWindow.number)))
        self.save.setIcon(QIcon('save.png'))
        self.save.setIconSize(QSize(int(40 * suma3/Ui_MainWindow.number), int(40 * suma3/Ui_MainWindow.number)))
        self.save.setStyleSheet("QPushButton{\n"
"    border: 1px solid;\n"
"    background-color: rgb(255, 255, 255);\n"
f"    border-radius: {27*suma3/Ui_MainWindow.number}%;\n"
"    \n"
"}\n"
"QPushButton:hover{\n"
"    background-color: rgb(223, 223, 223);\n"
"}")
        self.save.setText("")
        self.save.setObjectName("save")
        self.verticalLayout_4.addWidget(self.save)

        self.rezervne = QtWidgets.QPushButton(self.widget)
        self.rezervne.setMinimumSize(
            QtCore.QSize(int(55 * suma3 / Ui_MainWindow.number), int(55 * suma3 / Ui_MainWindow.number)))
        self.rezervne.setIcon(QIcon('rezerv.png'))
        self.rezervne.setIconSize(QSize(int(40 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        self.rezervne.setStyleSheet("QPushButton{\n"
                                "    border: 1px solid;\n"
                                "    background-color: rgb(255, 255, 255);\n"
                                f"    border-radius: {27 * suma3 / Ui_MainWindow.number}%;\n"
                                "    \n"
                                "}\n"
                                "QPushButton:hover{\n"
                                "    background-color: rgb(223, 223, 223);\n"
                                "}")
        self.rezervne.setText("")
        self.rezervne.setObjectName("rezervne")
        self.verticalLayout_4.addWidget(self.rezervne)

        self.error = QtWidgets.QPushButton(self.widget)
        self.error.setMinimumSize(
            QtCore.QSize(int(55 * suma3 / Ui_MainWindow.number), int(55 * suma3 / Ui_MainWindow.number)))
        self.error.setIcon(QIcon('error.png'))
        self.error.setIconSize(QSize(int(40 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        self.error.setStyleSheet("QPushButton{\n"
                                "    border: 1px solid;\n"
                                "    background-color: rgb(255, 255, 255);\n"
                                f"    border-radius: {27 * suma3 / Ui_MainWindow.number}%;\n"
                                "    \n"
                                "}\n"
                                "QPushButton:hover{\n"
                                "    background-color: rgb(223, 223, 223);\n"
                                "}")
        self.error.setText("")
        self.error.setObjectName("error")
        self.verticalLayout_4.addWidget(self.error)
        self.error.hide()


        spacerItem = QtWidgets.QSpacerItem(int(20*suma3/Ui_MainWindow.number), int(300*suma3/Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Fixed)
        self.verticalLayout_4.addItem(spacerItem)
        self.plus = QtWidgets.QPushButton(self.widget)
        self.plus.setMinimumSize(QtCore.QSize(int(55 * suma3/Ui_MainWindow.number), int(55 * suma3/Ui_MainWindow.number)))
        self.plus.setIcon(QIcon('+.png'))
        self.plus.setIconSize(QSize(int(45 * suma3/Ui_MainWindow.number), int(45 * suma3/Ui_MainWindow.number)))
        self.plus.setStyleSheet("QPushButton{\n"
                                         "    border: 1px solid;\n"
                                         "    background-color: rgb(255, 255, 255);\n"
                                         f"    border-radius: {27 * suma3/Ui_MainWindow.number}%;\n"
                                         "    \n"
                                         "}\n"
                                         "QPushButton:hover{\n"
                                         "    background-color: rgb(223, 223, 223);\n"
                                         "}")
        self.plus.setText("")
        self.plus.setObjectName("plus")
        self.verticalLayout_4.addWidget(self.plus)

        self.next = QtWidgets.QPushButton(self.widget)
        self.next.setMinimumSize(QtCore.QSize(int(55 * suma3/Ui_MainWindow.number), int(55 * suma3/Ui_MainWindow.number)))
        self.next.setIcon(QIcon('1argb.png'))
        self.next.setIconSize(QSize(int(45 * suma3/Ui_MainWindow.number), int(45 * suma3/Ui_MainWindow.number)))
        self.next.setStyleSheet("QPushButton{\n"
                                "    border: 1px solid;\n"
                                "    background-color: rgb(255, 255, 255);\n"
                                f"    border-radius: {27 * suma3/Ui_MainWindow.number}%;\n"
                                "    \n"
                                "}\n"
                                "QPushButton:hover{\n"
                                "    background-color: rgb(223, 223, 223);\n"
                                "}")
        self.next.setText("")
        self.next.setObjectName("next")
        self.next.hide()
        self.verticalLayout_4.addWidget(self.next)

        self.next_2 = QtWidgets.QPushButton(self.widget)
        self.next_2.setMinimumSize(
            QtCore.QSize(int(55 * suma3 / Ui_MainWindow.number), int(55 * suma3 / Ui_MainWindow.number)))
        self.next_2.setIcon(QIcon('1argb.png'))
        self.next_2.setIconSize(QSize(int(45 * suma3 / Ui_MainWindow.number), int(45 * suma3 / Ui_MainWindow.number)))
        self.next_2.setStyleSheet("QPushButton{\n"
                                "    border: 1px solid;\n"
                                "    background-color: rgb(255, 255, 255);\n"
                                f"    border-radius: {27 * suma3 / Ui_MainWindow.number}%;\n"
                                "    \n"
                                "}\n"
                                "QPushButton:hover{\n"
                                "    background-color: rgb(223, 223, 223);\n"
                                "}")
        self.next_2.setText("")
        self.next_2.setObjectName("next_2")
        self.next_2.hide()
        self.verticalLayout_4.addWidget(self.next_2)


        self.minus = QtWidgets.QPushButton(self.widget)
        self.minus.setMinimumSize(QtCore.QSize(int(55 * suma3/Ui_MainWindow.number), int(55 * suma3/Ui_MainWindow.number)))
        self.minus.setIcon(QIcon('-.png'))
        self.minus.setIconSize(QSize(int(45 * suma3/Ui_MainWindow.number), int(45 * suma3/Ui_MainWindow.number)))
        self.minus.setStyleSheet("QPushButton{\n"
                                "    border: 1px solid;\n"
                                "    background-color: rgb(255, 255, 255);\n"
                                f"    border-radius: {27 * suma3/Ui_MainWindow.number}%;\n"
                                "    \n"
                                "}\n"
                                "QPushButton:hover{\n"
                                "    background-color: rgb(223, 223, 223);\n"
                                "}")
        self.minus.setText("")
        self.minus.setObjectName("minus")
        self.verticalLayout_4.addWidget(self.minus)


        self.previous = QtWidgets.QPushButton(self.widget)
        self.previous.setMinimumSize(QtCore.QSize(int(55 * suma3/Ui_MainWindow.number), int(55 * suma3/Ui_MainWindow.number)))
        self.previous.setIcon(QIcon('1rgb.png'))
        self.previous.setIconSize(QSize(int(45 * suma3/Ui_MainWindow.number), int(45 * suma3/Ui_MainWindow.number)))
        self.previous.setStyleSheet("QPushButton{\n"
                                 "    border: 1px solid;\n"
                                 "    background-color: rgb(255, 255, 255);\n"
                                 f"    border-radius: {27 * suma3/Ui_MainWindow.number}%;\n"
                                 "    \n"
                                 "}\n"
                                 "QPushButton:hover{\n"
                                 "    background-color: rgb(223, 223, 223);\n"
                                 "}")
        self.previous.setText("")
        self.previous.setObjectName("previous")
        self.previous.hide()
        self.verticalLayout_4.addWidget(self.previous)

        self.previous_2 = QtWidgets.QPushButton(self.widget)
        self.previous_2.setMinimumSize(
            QtCore.QSize(int(55 * suma3 / Ui_MainWindow.number), int(55 * suma3 / Ui_MainWindow.number)))
        self.previous_2.setIcon(QIcon('1rgb.png'))
        self.previous_2.setIconSize(QSize(int(45 * suma3 / Ui_MainWindow.number), int(45 * suma3 / Ui_MainWindow.number)))
        self.previous_2.setStyleSheet("QPushButton{\n"
                                    "    border: 1px solid;\n"
                                    "    background-color: rgb(255, 255, 255);\n"
                                    f"    border-radius: {27 * suma3 / Ui_MainWindow.number}%;\n"
                                    "    \n"
                                    "}\n"
                                    "QPushButton:hover{\n"
                                    "    background-color: rgb(223, 223, 223);\n"
                                    "}")
        self.previous_2.setText("")
        self.previous_2.setObjectName("previous_2")
        self.previous_2.hide()
        self.verticalLayout_4.addWidget(self.previous_2)

        self.gridLayout_2.addWidget(self.widget, int(0), int(2), int(1), int(1))

        self.frame = QtWidgets.QFrame(self.centralwidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.frame.sizePolicy().hasHeightForWidth())
        self.frame.setSizePolicy(sizePolicy)
        self.frame.setMinimumSize(QtCore.QSize(int(300 * suma3 / Ui_MainWindow.number), int(0 * suma3 / Ui_MainWindow.number)))
        self.frame.setMaximumSize(QtCore.QSize(int(0 * suma3 / Ui_MainWindow.number), int(16777215 * suma3 / Ui_MainWindow.number)))
        # self.frame.setMinimumSize(QtCore.QSize(int(300*suma3/Ui_MainWindow.number), int(0*suma3/Ui_MainWindow.number)))
        # self.frame.setMaximumSize(QtCore.QSize(int(800*suma3/Ui_MainWindow.number), int(16777215*suma3/Ui_MainWindow.number)))
        self.frame.setStyleSheet("QFrame{\n"
f"    border-radius: {2*suma3/Ui_MainWindow.number}%;\n"
f"    border: {2*suma3/Ui_MainWindow.number}px solid;\n"
"    border-color: rgb(41, 213, 202);\n"
"}")
        self.frame.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.frame.setFrameShadow(QtWidgets.QFrame.Raised)
        self.frame.setObjectName("frame")
        self.frame.hide()
        self.gridLayout = QtWidgets.QGridLayout(self.frame)
        self.gridLayout.setObjectName("gridLayout")

        spacerItem1 = QtWidgets.QSpacerItem(int(20*suma3/Ui_MainWindow.number), int(20*suma3/Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Maximum)
        self.gridLayout.addItem(spacerItem1, int(2), int(0), int(1), int(1))
        spacerItem2 = QtWidgets.QSpacerItem(int(20*suma3/Ui_MainWindow.number), int(100*suma3/Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Maximum)
        self.gridLayout.addItem(spacerItem2, int(4), int(0), int(1), int(1))
        self.verticalLayout = QtWidgets.QVBoxLayout()
        self.verticalLayout.setObjectName("verticalLayout")
        spacerItem3 = QtWidgets.QSpacerItem(int(20*suma3/Ui_MainWindow.number), int(5*suma3/Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Maximum)
        self.verticalLayout.addItem(spacerItem3)
        self.PIP = QtWidgets.QLineEdit(self.frame)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.PIP.sizePolicy().hasHeightForWidth())
        self.PIP.setSizePolicy(sizePolicy)
        self.PIP.setMinimumSize(QtCore.QSize(int(260*suma3/Ui_MainWindow.number), int(40*suma3/Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        self.PIP.setFont(font)
        self.PIP.setStyleSheet("QLineEdit{\n"
"    \n"
"    background-color: rgb(255, 255, 255);\n"
f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
f"    border: 1px solid;\n"
"}\n"
"QLineEdit:hover {\n"
f"    border: 1px solid;\n"
"    border-color: rgb(41, 213, 202);\n"
f"    border-radius: {2*suma3/Ui_MainWindow.number}%;\n"
"\n"
"    \n"
"}")
        self.PIP.setObjectName("PIP")
        self.verticalLayout.addWidget(self.PIP)



        spacerItem4 = QtWidgets.QSpacerItem(int(20*suma3/Ui_MainWindow.number), int(10*suma3/Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Maximum)
        self.verticalLayout.addItem(spacerItem4)
        self.label = QtWidgets.QLabel(self.frame)
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        self.label.setFont(font)
        self.label.setStyleSheet("border: 0px;")
        self.label.setObjectName("label")
        self.verticalLayout.addWidget(self.label)
        self.dateEdit = QtWidgets.QDateEdit(self.frame)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.dateEdit.sizePolicy().hasHeightForWidth())
        self.dateEdit.setSizePolicy(sizePolicy)
        self.dateEdit.setMinimumSize(QtCore.QSize(int(260*suma3/Ui_MainWindow.number), int(40*suma3/Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        self.dateEdit.setFont(font)
        self.dateEdit.setStyleSheet("QDateEdit{\n"
"background-color: rgb(255, 255, 255);\n"
f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
"    border: 1px solid;\n"
"}\n"
"QDateEdit:hover {\n"
f"    border: 1px solid;\n"
"    border-color: rgb(41, 213, 202);\n"
f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
"}")
        self.dateEdit.setObjectName("dateEdit")
        self.verticalLayout.addWidget(self.dateEdit)
        spacerItem5 = QtWidgets.QSpacerItem(int(20*suma3/Ui_MainWindow.number), int(20*suma3/Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Maximum)
        self.verticalLayout.addItem(spacerItem5)
        self.Povtorna_KO = QtWidgets.QCheckBox(self.frame)
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        font.setBold(False)
        font.setItalic(False)
        font.setWeight(50)
        self.Povtorna_KO.setFont(font)
        self.Povtorna_KO.setObjectName("Povtorna_KO")
        self.verticalLayout.addWidget(self.Povtorna_KO)
        spacerItem7 = QtWidgets.QSpacerItem(int(20*suma3/Ui_MainWindow.number), int(30*suma3/Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Maximum)
        self.verticalLayout.addItem(spacerItem7)
        self.label_2 = QtWidgets.QLabel(self.frame)
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        self.label_2.setFont(font)
        self.label_2.setStyleSheet("border:0px;")
        self.label_2.setObjectName("label_2")
        self.verticalLayout.addWidget(self.label_2)
        self.gridLayout.addLayout(self.verticalLayout, int(0), int(0), int(1), int(1))
        self.verticalLayout_2 = QtWidgets.QVBoxLayout()
        self.verticalLayout_2.setObjectName("verticalLayout_2")

        self.School = QtWidgets.QRadioButton(self.frame)
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        font.setBold(False)
        font.setItalic(False)
        font.setWeight(50)
        self.School.setFont(font)
        self.School.setObjectName("School")
        self.verticalLayout_2.addWidget(self.School)
        self.Doschilnul = QtWidgets.QRadioButton(self.frame)
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        font.setBold(False)
        font.setItalic(False)
        font.setWeight(50)
        self.Doschilnul.setFont(font)
        self.Doschilnul.setObjectName("Doschilnul")
        self.verticalLayout_2.addWidget(self.Doschilnul)
        spacerItem6 = QtWidgets.QSpacerItem(int(20*suma3/Ui_MainWindow.number), int(20*suma3/Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Maximum)
        self.verticalLayout_2.addItem(spacerItem6)
        self.label_3 = QtWidgets.QLabel(self.frame)
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        self.label_3.setFont(font)
        self.label_3.setStyleSheet("border:0px")
        self.label_3.setObjectName("label_3")
        self.verticalLayout_2.addWidget(self.label_3)
        self.Perent_Kontakt = QtWidgets.QLineEdit(self.frame)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.Perent_Kontakt.sizePolicy().hasHeightForWidth())
        self.Perent_Kontakt.setSizePolicy(sizePolicy)
        self.Perent_Kontakt.setMinimumSize(QtCore.QSize(int(260*suma3/Ui_MainWindow.number), int(40*suma3/Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        self.Perent_Kontakt.setFont(font)
        self.Perent_Kontakt.setStyleSheet("QLineEdit{\n"
"    \n"
"    background-color: rgb(255, 255, 255);\n"
f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
"    border: 1px solid;\n"
"}\n"
"QLineEdit:hover {\n"
"    border: 1px solid;\n"
"    border-color: rgb(41, 213, 202);\n"
"    border-radius: 4%;\n"
"\n"
"    \n"
"}")
        self.Perent_Kontakt.setObjectName("Perent_Kontakt")
        self.verticalLayout_2.addWidget(self.Perent_Kontakt)

        spacerItem60 = QtWidgets.QSpacerItem(int(20 * suma3 / Ui_MainWindow.number),
                                            int(50 * suma3 / Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum,
                                            QtWidgets.QSizePolicy.Maximum)
        self.verticalLayout_2.addItem(spacerItem60)

        self.comboBox_qestions = QtWidgets.QComboBox(MainWindow)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.comboBox_qestions.sizePolicy().hasHeightForWidth())
        self.comboBox_qestions.setSizePolicy(sizePolicy)
        self.comboBox_qestions.setMinimumSize(QtCore.QSize(int(260 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.comboBox_qestions.setFont(font)
        self.comboBox_qestions.setStyleSheet(
                "QComboBox{\n"
                "    background-color: rgb(255, 255, 255);\n"
                f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                f"    border: 1px solid;\n"
                "}\n"
                "QComboBox:hover{"
                "border-color: rgb(41, 213, 202)"

                "}"
                "\n"
                "QAbstractItemView{\n"
                "  selection-color: black;\n"
                "  selection-background-color: #05F3CC ;\n"
                "  border: 0px ;\n"

                "}\n"
                "QComboBox::drop-down{\n"
                "  border: 0 ;\n"
                "}\n"
                "")
        self.comboBox_qestions.setObjectName("comboBox_qestions")
        self.verticalLayout_2.addWidget(self.comboBox_qestions)


        self.Go_qestions = QtWidgets.QPushButton(self.frame)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.Go_qestions.sizePolicy().hasHeightForWidth())
        self.Go_qestions.setSizePolicy(sizePolicy)
        self.Go_qestions.setMinimumSize(QtCore.QSize(int(260 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.Go_qestions.setFont(font)
        self.Go_qestions.setStyleSheet("QPushButton{\n"
                                "    \n"
                                "    background-color: rgb(255, 255, 255);\n"
                                f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                f"    border: 1px solid;\n"
                                "}\n"
                                "QPushButton:hover {\n"
                                f"    border: 1px solid;\n"
                                "    border-color: rgb(41, 213, 202);\n"
                                f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                "\n"
                                "    \n"
                                "}")
        self.Go_qestions.setObjectName("Go_qestions")
        self.Go_qestions.setText("Опитування")
        self.verticalLayout_2.addWidget(self.Go_qestions)

        self.update_qestions = QtWidgets.QPushButton(self.frame)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.update_qestions.sizePolicy().hasHeightForWidth())
        self.update_qestions.setSizePolicy(sizePolicy)
        self.update_qestions.setMinimumSize(QtCore.QSize(int(260 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.update_qestions.setFont(font)
        self.update_qestions.setStyleSheet("QPushButton{\n"
                                       "    \n"
                                       "    background-color: rgb(255, 255, 255);\n"
                                       f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                       f"    border: 1px solid;\n"
                                       "}\n"
                                       "QPushButton:hover {\n"
                                       f"    border: 1px solid;\n"
                                       "    border-color: rgb(41, 213, 202);\n"
                                       f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                       "\n"
                                       "    \n"
                                       "}")
        self.update_qestions.setObjectName("Go_qestions")
        self.update_qestions.setText("Зберегти Опитування")
        self.verticalLayout_2.addWidget(self.update_qestions)

        self.show_anleta_ko = QtWidgets.QPushButton(self.frame)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.show_anleta_ko.sizePolicy().hasHeightForWidth())
        self.show_anleta_ko.setSizePolicy(sizePolicy)
        self.show_anleta_ko.setMinimumSize(
            QtCore.QSize(int(260 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.show_anleta_ko.setFont(font)
        self.show_anleta_ko.setStyleSheet("QPushButton{\n"
                                           "    \n"
                                           "    background-color: rgb(255, 255, 255);\n"
                                           f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                           f"    border: 1px solid;\n"
                                           "}\n"
                                           "QPushButton:hover {\n"
                                           f"    border: 1px solid;\n"
                                           "    border-color: rgb(41, 213, 202);\n"
                                           f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                           "\n"
                                           "    \n"
                                           "}")
        self.show_anleta_ko.setObjectName("show_anleta_ko")
        self.show_anleta_ko.setText("Результати")
        self.verticalLayout_2.addWidget(self.show_anleta_ko)


        self.gridLayout.addLayout(self.verticalLayout_2, int(3), int(0), int(1), int(1))
        self.horizontalLayout_2 = QtWidgets.QHBoxLayout()
        self.horizontalLayout_2.setObjectName("horizontalLayout_2")
        self.GoKO = QtWidgets.QPushButton(self.frame)
        self.GoKO.setMinimumSize(QtCore.QSize(int(0*suma3/Ui_MainWindow.number), int(40*suma3/Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        self.GoKO.setFont(font)
        self.GoKO.setStyleSheet("QPushButton{\n"
"    \n"
"    background-color: rgb(255, 255, 255);\n"
f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
f"    border: 1px solid;\n"
"}\n"
"QPushButton:hover {\n"
f"    border: 1px solid;\n"
"    border-color: rgb(41, 213, 202);\n"
f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
"\n"
"    \n"
"}")
        self.GoKO.setObjectName("GoKO")
        self.horizontalLayout_2.addWidget(self.GoKO)
        self.add_row = QtWidgets.QPushButton(self.frame)
        self.add_row.setMinimumSize(QtCore.QSize(int(0*suma3/Ui_MainWindow.number), int(40*suma3/Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        self.add_row.setFont(font)
        self.add_row.setStyleSheet("QPushButton{\n"
"    \n"
"    background-color: rgb(255, 255, 255);\n"
f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
f"    border: 1px solid;\n"
"}\n"
"QPushButton:hover {\n"
f"    border: 1px solid;\n"
"    border-color: rgb(41, 213, 202);\n"
f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
"\n"
"    \n"
"}")
        self.add_row.setObjectName("GoKO_2")
        self.horizontalLayout_2.addWidget(self.add_row)
        self.gridLayout.addLayout(self.horizontalLayout_2, int(5), int(0), int(1), int(1))
        self.horizontalLayout = QtWidgets.QHBoxLayout()
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.comboBox = QtWidgets.QComboBox(MainWindow)
        self.comboBox.setMinimumSize(QtCore.QSize(int(20*suma3/Ui_MainWindow.number), int(40*suma3/Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        self.comboBox.setFont(font)
        self.comboBox.setStyleSheet(
                "QComboBox{\n"
                "    background-color: rgb(255, 255, 255);\n"
                f"    border-radius: {4 * suma3/Ui_MainWindow.number}%;\n"
                f"    border: 1px solid;\n"
                "}\n"
                "QComboBox:hover{"
                "border-color: rgb(41, 213, 202)"

                "}"
                "\n"
                "QAbstractItemView{\n"
                "  selection-color: black;\n"
                "  selection-background-color: #05F3CC ;\n"
                "  border: 0px ;\n"

                "}\n"
                "QComboBox::drop-down{\n"
                "  border: 0 ;\n"
                "}\n"
"")
        self.comboBox.setObjectName("comboBox")
        self.horizontalLayout.addWidget(self.comboBox)
        self.Show_Metoduk = QtWidgets.QPushButton(self.frame)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.Show_Metoduk.sizePolicy().hasHeightForWidth())
        self.Show_Metoduk.setSizePolicy(sizePolicy)
        self.Show_Metoduk.setMinimumSize(QtCore.QSize(int(40*suma3/Ui_MainWindow.number), int(40*suma3/Ui_MainWindow.number)))
        self.Show_Metoduk.setStyleSheet("QPushButton{\n"
"    \n"
"    background-color: rgb(255, 255, 255);\n"
f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
f"    border: 1px solid;\n"
"}\n"
"QPushButton:hover {\n"
f"    border: 1px solid;\n"
"    border-color: rgb(41, 213, 202);\n"
f"    border-radius: {2*suma3/Ui_MainWindow.number}%;\n"
"\n"
"    \n"
"}")
        self.Show_Metoduk.setObjectName("Show_Metoduk")
        self.Show_Metoduk.hide()
        self.horizontalLayout.addWidget(self.Show_Metoduk)
        self.gridLayout.addLayout(self.horizontalLayout, int(1), int(0), int(1), int(1))
        self.gridLayout_2.addWidget(self.frame, int(0), int(1), int(1), int(1))

        # фільтр бар
        # self.gridLayout_2.addWidget(self.widget, 0, 1, 1, 1)
        self.frame_two = QtWidgets.QFrame(self.centralwidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.frame_two.sizePolicy().hasHeightForWidth())
        self.frame_two.setSizePolicy(sizePolicy)
        self.frame_two.setMinimumSize(QtCore.QSize(int(300*suma3/Ui_MainWindow.number), int(0*suma3/Ui_MainWindow.number)))
        self.frame_two.setMaximumSize(QtCore.QSize(int(0*suma3/Ui_MainWindow.number), int(16777215*suma3/Ui_MainWindow.number)))
        self.frame_two.setStyleSheet("QFrame{\n"
                                     f"    border-radius: {2*suma3/Ui_MainWindow.number}%;\n"
                                     f"    border: {2*suma3/Ui_MainWindow.number}px solid;\n"
                                     "    border-color: rgb(41, 213, 202);\n"
                                     "}")
        self.frame_two.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.frame_two.setFrameShadow(QtWidgets.QFrame.Raised)
        self.frame_two.setObjectName("frame_two")
        self.gridLayout = QtWidgets.QGridLayout(self.frame_two)
        self.gridLayout.setObjectName("gridLayout")
        self.formLayout = QtWidgets.QFormLayout()
        self.formLayout.setObjectName("formLayout")
        self.poshyk_ko = QtWidgets.QPushButton(self.frame_two)
        self.poshyk_ko.setMinimumSize(QtCore.QSize(int(130*suma3/Ui_MainWindow.number), int(40*suma3/Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        self.poshyk_ko.setFont(font)
        self.poshyk_ko.setStyleSheet("QPushButton{\n"
                                       "    \n"
                                       "    background-color: rgb(255, 255, 255);\n"
                                       f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                       f"    border: 1px solid;\n"
                                       "}\n"
                                       "QPushButton:hover {\n"
                                       f"    border: 1px solid;\n"
                                       "    border-color: rgb(41, 213, 202);\n"
                                       f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                       "\n"
                                       "    \n"
                                       "}")
        self.poshyk_ko.setObjectName("poshyk_ko")
        self.formLayout.setWidget(0, QtWidgets.QFormLayout.LabelRole, self.poshyk_ko)
        self.all_ko = QtWidgets.QPushButton(self.frame_two)
        self.all_ko.setMinimumSize(QtCore.QSize(int(130*suma3/Ui_MainWindow.number), int(40*suma3/Ui_MainWindow.number)))
        self.all_ko.setMaximumSize(QtCore.QSize(int(130*suma3/Ui_MainWindow.number), int(16777215*suma3/Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        self.all_ko.setFont(font)
        self.all_ko.setStyleSheet("QPushButton{\n"
                                  "    \n"
                                  "    background-color: rgb(255, 255, 255);\n"
                                  f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                  f"    border: 1px solid;\n"
                                  "}\n"
                                  "QPushButton:hover {\n"
                                  f"    border: 1px solid;\n"
                                  "    border-color: rgb(41, 213, 202);\n"
                                  f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                  "\n"
                                  "    \n"
                                  "}")
        self.all_ko.setObjectName("all_ko")
        self.formLayout.setWidget(0, QtWidgets.QFormLayout.FieldRole, self.all_ko)
        self.show_ko = QtWidgets.QPushButton(self.frame_two)
        self.show_ko.setMinimumSize(QtCore.QSize(int(130*suma3/Ui_MainWindow.number), int(40*suma3/Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        self.show_ko.setFont(font)
        self.show_ko.setStyleSheet("QPushButton{\n"
                                   "    \n"
                                   "    background-color: rgb(255, 255, 255);\n"
                                   f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                   f"    border: 1px solid;\n"
                                   "}\n"
                                   "QPushButton:hover {\n"
                                   f"    border: 1px solid;\n"
                                   "    border-color: rgb(41, 213, 202);\n"
                                   f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                   "\n"
                                   "    \n"
                                   "}")
        self.show_ko.setObjectName("show_ko")
        self.formLayout.setWidget(1, QtWidgets.QFormLayout.LabelRole, self.show_ko)
        self.delete_ko = QtWidgets.QPushButton(self.frame_two)
        self.delete_ko.setMinimumSize(QtCore.QSize(int(130*suma3/Ui_MainWindow.number), int(40*suma3/Ui_MainWindow.number)))
        self.delete_ko.setMaximumSize(QtCore.QSize(int(130*suma3/Ui_MainWindow.number), int(16777215*suma3/Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        self.delete_ko.setFont(font)
        self.delete_ko.setStyleSheet("QPushButton{\n"
                                     "    \n"
                                     "    background-color: rgb(255, 255, 255);\n"
                                     f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                     f"    border: 1px solid;\n"
                                     "}\n"
                                     "QPushButton:hover {\n"
                                     f"    border: 1px solid;\n"
                                     "    border-color: rgb(41, 213, 202);\n"
                                     f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                     "\n"
                                     "    \n"
                                     "}")
        self.delete_ko.setObjectName("delete_ko")
        self.formLayout.setWidget(1, QtWidgets.QFormLayout.FieldRole, self.delete_ko)
        self.add_row_table = QtWidgets.QPushButton(self.frame_two)
        self.add_row_table.setMinimumSize(QtCore.QSize(int(130*suma3/Ui_MainWindow.number), int(40*suma3/Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        self.add_row_table.setFont(font)
        self.add_row_table.setStyleSheet("QPushButton{\n"
                                         "    \n"
                                         "    background-color: rgb(255, 255, 255);\n"
                                         f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                         f"    border: 1px solid;\n"
                                         "}\n"
                                         "QPushButton:hover {\n"
                                         f"    border: 1px solid;\n"
                                         "    border-color: rgb(41, 213, 202);\n"
                                         f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                         "\n"
                                         "    \n"
                                         "}")
        self.add_row_table.setObjectName("add_row_table")
        self.formLayout.setWidget(2, QtWidgets.QFormLayout.LabelRole, self.add_row_table)
        self.update_dani = QtWidgets.QPushButton(self.frame_two)
        self.update_dani.setMinimumSize(QtCore.QSize(int(130*suma3/Ui_MainWindow.number), int(40*suma3/Ui_MainWindow.number)))
        self.update_dani.setMaximumSize(QtCore.QSize(int(130*suma3/Ui_MainWindow.number), int(16777215*suma3/Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        self.update_dani.setFont(font)
        self.update_dani.setStyleSheet("QPushButton{\n"
                                       "    \n"
                                       "    background-color: rgb(255, 255, 255);\n"
                                       f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                       f"    border: 1px solid;\n"
                                       "}\n"
                                       "QPushButton:hover {\n"
                                       f"    border: 1px solid;\n"
                                       "    border-color: rgb(41, 213, 202);\n"
                                       f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                       "\n"
                                       "    \n"
                                       "}")
        self.update_dani.setObjectName("update_dani")
        self.formLayout.setWidget(2, QtWidgets.QFormLayout.FieldRole, self.update_dani)

        self.avto_row = QtWidgets.QPushButton(self.frame_two)
        self.avto_row.setMinimumSize(QtCore.QSize(int(130*suma3/Ui_MainWindow.number), int(40*suma3/Ui_MainWindow.number)))
        self.avto_row.setMaximumSize(QtCore.QSize(int(130*suma3/Ui_MainWindow.number), int(16777215*suma3/Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        self.avto_row.setFont(font)
        self.avto_row.setStyleSheet("QPushButton{\n"
                                    "    \n"
                                    "    background-color: rgb(255, 255, 255);\n"
                                    f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                    "    border: 1px solid;\n"
                                    "}\n"
                                    "QPushButton:hover {\n"
                                    "    border: 1px solid;\n"
                                    "    border-color: rgb(41, 213, 202);\n"
                                    f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                    "\n"
                                    "    \n"
                                    "}")
        self.avto_row.setObjectName("avto_row")
        self.avto_row.setText("Автозаповнення")
        self.formLayout.setWidget(3, QtWidgets.QFormLayout.LabelRole, self.avto_row)

        self.all_dani = QtWidgets.QPushButton(self.frame_two)
        self.all_dani.setMinimumSize(QtCore.QSize(int(130*suma3/Ui_MainWindow.number), int(40*suma3/Ui_MainWindow.number)))
        self.all_dani.setMaximumSize(QtCore.QSize(int(130*suma3/Ui_MainWindow.number), int(16777215*suma3/Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        self.all_dani.setFont(font)
        self.all_dani.setStyleSheet("QPushButton{\n"
                                    "    \n"
                                    "    background-color: rgb(255, 255, 255);\n"
                                    f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                    f"    border: 1px solid;\n"
                                    "}\n"
                                    "QPushButton:hover {\n"
                                    f"    border: 1px solid;\n"
                                    "    border-color: rgb(41, 213, 202);\n"
                                    f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                    "\n"
                                    "    \n"
                                    "}")
        self.all_dani.setObjectName("all_dani")
        self.formLayout.setWidget(3, QtWidgets.QFormLayout.FieldRole, self.all_dani)

        self.anketa_show = QtWidgets.QPushButton(self.frame_two)
        self.anketa_show.setMinimumSize(
            QtCore.QSize(int(130 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        self.anketa_show.setMaximumSize(
            QtCore.QSize(int(130 * suma3 / Ui_MainWindow.number), int(16777215 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.anketa_show.setFont(font)
        self.anketa_show.setStyleSheet("QPushButton{\n"
                                    "    \n"
                                    "    background-color: rgb(255, 255, 255);\n"
                                    f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                    "    border: 1px solid;\n"
                                    "}\n"
                                    "QPushButton:hover {\n"
                                    "    border: 1px solid;\n"
                                    "    border-color: rgb(41, 213, 202);\n"
                                    f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                    "\n"
                                    "    \n"
                                    "}")
        self.anketa_show.setObjectName("anketa_show")
        self.anketa_show.setText("Опитування")
        self.formLayout.setWidget(4, QtWidgets.QFormLayout.LabelRole, self.anketa_show)

        self.delete_anketa = QtWidgets.QPushButton(self.frame_two)
        self.delete_anketa.setMinimumSize(
            QtCore.QSize(int(130 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        self.delete_anketa.setMaximumSize(
            QtCore.QSize(int(130 * suma3 / Ui_MainWindow.number), int(16777215 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.delete_anketa.setFont(font)
        self.delete_anketa.setStyleSheet("QPushButton{\n"
                                    "    \n"
                                    "    background-color: rgb(255, 255, 255);\n"
                                    f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                    f"    border: 1px solid;\n"
                                    "}\n"
                                    "QPushButton:hover {\n"
                                    f"    border: 1px solid;\n"
                                    "    border-color: rgb(41, 213, 202);\n"
                                    f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                    "\n"
                                    "    \n"
                                    "}")
        self.delete_anketa.setObjectName("delete_anketa")
        self.delete_anketa.setText('Видалити Оп.')
        self.formLayout.setWidget(4, QtWidgets.QFormLayout.FieldRole, self.delete_anketa)


        self.gridLayout.addLayout(self.formLayout, 3, 0, 1, 1)
        self.verticalLayout = QtWidgets.QVBoxLayout()
        self.verticalLayout.setObjectName("verticalLayout")
        self.first_name = QtWidgets.QLineEdit(self.frame_two)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.first_name.sizePolicy().hasHeightForWidth())
        self.first_name.setSizePolicy(sizePolicy)
        self.first_name.setMinimumSize(QtCore.QSize(int(260*suma3/Ui_MainWindow.number), int(40*suma3/Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        self.first_name.setFont(font)
        self.first_name.setStyleSheet("QLineEdit{\n"
                                      "    \n"
                                      "    background-color: rgb(255, 255, 255);\n"
                                      f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                      f"    border: 1px solid;\n"
                                      "}\n"
                                      "QLineEdit:hover {\n"
                                      f"    border: 1px solid;\n"
                                      "    border-color: rgb(41, 213, 202);\n"
                                      f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                      "\n"
                                      "    \n"
                                      "}")
        self.first_name.setObjectName("first_name")
        self.verticalLayout.addWidget(self.first_name)
        self.last_name = QtWidgets.QLineEdit(self.frame_two)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.last_name.sizePolicy().hasHeightForWidth())
        self.last_name.setSizePolicy(sizePolicy)
        self.last_name.setMinimumSize(QtCore.QSize(int(260*suma3/Ui_MainWindow.number), int(40*suma3/Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        self.last_name.setFont(font)
        self.last_name.setStyleSheet("QLineEdit{\n"
                                     "    \n"
                                     "    background-color: rgb(255, 255, 255);\n"
                                     f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                     f"    border: 1px solid;\n"
                                     "}\n"
                                     "QLineEdit:hover {\n"
                                     f"    border: 1px solid;\n"
                                     "    border-color: rgb(41, 213, 202);\n"
                                     f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                     "\n"
                                     "    \n"
                                     "}")
        self.last_name.setObjectName("last_name")
        self.verticalLayout.addWidget(self.last_name)
        self.pobatkovi = QtWidgets.QLineEdit(self.frame_two)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.pobatkovi.sizePolicy().hasHeightForWidth())
        self.pobatkovi.setSizePolicy(sizePolicy)
        self.pobatkovi.setMinimumSize(QtCore.QSize(int(260*suma3/Ui_MainWindow.number), int(40*suma3/Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        self.pobatkovi.setFont(font)
        self.pobatkovi.setStyleSheet("QLineEdit{\n"
                                     "    \n"
                                     "    background-color: rgb(255, 255, 255);\n"
                                     f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                     f"    border: 1px solid;\n"
                                     "}\n"
                                     "QLineEdit:hover {\n"
                                     f"    border: 1px solid;\n"
                                     "    border-color: rgb(41, 213, 202);\n"
                                     f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                     "\n"
                                     "    \n"
                                     "}")
        self.pobatkovi.setObjectName("pobatkovi")
        self.verticalLayout.addWidget(self.pobatkovi)
        self.Perent_Kontakt_2 = QtWidgets.QLineEdit(self.frame_two)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.Perent_Kontakt_2.sizePolicy().hasHeightForWidth())
        self.Perent_Kontakt_2.setSizePolicy(sizePolicy)
        self.Perent_Kontakt_2.setMinimumSize(QtCore.QSize(int(260*suma3/Ui_MainWindow.number), int(40*suma3/Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        self.Perent_Kontakt_2.setFont(font)
        self.Perent_Kontakt_2.setStyleSheet("QLineEdit{\n"
                                            "    \n"
                                            "    background-color: rgb(255, 255, 255);\n"
                                            f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                            f"    border: 1px solid;\n"
                                            "}\n"
                                            "QLineEdit:hover {\n"
                                            f"    border: 1px solid;\n"
                                            "    border-color: rgb(41, 213, 202);\n"
                                            f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                            "\n"
                                            "    \n"
                                            "}")
        self.Perent_Kontakt_2.setObjectName("Perent_Kontakt_2")
        self.verticalLayout.addWidget(self.Perent_Kontakt_2)
        spacerItem1 = QtWidgets.QSpacerItem(int(20*suma3/Ui_MainWindow.number), int(10*suma3/Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Maximum)
        self.verticalLayout.addItem(spacerItem1)
        self.Povtorna_KO_2 = QtWidgets.QCheckBox(self.frame_two)
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        font.setBold(False)
        font.setItalic(False)
        font.setWeight(50)
        self.Povtorna_KO_2.setFont(font)
        self.Povtorna_KO_2.setObjectName("Povtorna_KO_2")
        self.verticalLayout.addWidget(self.Povtorna_KO_2)
        self.Pervuna_KO = QtWidgets.QCheckBox(self.frame_two)
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        font.setBold(False)
        font.setItalic(False)
        font.setWeight(50)
        self.Pervuna_KO.setFont(font)
        self.Pervuna_KO.setObjectName("Pervuna_KO")
        self.verticalLayout.addWidget(self.Pervuna_KO)
        spacerItem2 = QtWidgets.QSpacerItem(int(20*suma3/Ui_MainWindow.number), int(10*suma3/Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Maximum)
        self.verticalLayout.addItem(spacerItem2)
        self.gridLayout.addLayout(self.verticalLayout, 0, 0, 1, 1)
        self.verticalLayout_3 = QtWidgets.QVBoxLayout()
        self.verticalLayout_3.setObjectName("verticalLayout_3")
        self.label_filter_4 = QtWidgets.QCheckBox(self.frame_two)
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        font.setBold(False)
        font.setItalic(False)
        font.setWeight(50)
        self.label_filter_4.setFont(font)
        self.label_filter_4.setObjectName("label_filter_4")
        self.verticalLayout_3.addWidget(self.label_filter_4)
        self.label_filter_2 = QtWidgets.QCheckBox(self.frame_two)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.label_filter_2.sizePolicy().hasHeightForWidth())
        self.label_filter_2.setSizePolicy(sizePolicy)
        self.label_filter_2.setMinimumSize(QtCore.QSize(int(100*suma3/Ui_MainWindow.number), int(0*suma3/Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        font.setBold(False)
        font.setItalic(False)
        font.setWeight(50)
        self.label_filter_2.setFont(font)
        self.label_filter_2.setObjectName("label_filter_2")
        self.verticalLayout_3.addWidget(self.label_filter_2)
        self.dateEdit_2 = QtWidgets.QDateEdit(self.frame_two)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.dateEdit_2.sizePolicy().hasHeightForWidth())
        self.dateEdit_2.setSizePolicy(sizePolicy)
        self.dateEdit_2.setMinimumSize(QtCore.QSize(int(260*suma3/Ui_MainWindow.number), int(40*suma3/Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        self.dateEdit_2.setFont(font)
        self.dateEdit_2.setStyleSheet("QDateEdit{\n"
                                      "background-color: rgb(255, 255, 255);\n"
                                      f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                      f"    border: 1px solid;\n"
                                      "}\n"
                                      "QDateEdit:hover {\n"
                                      f"    border: 1px solid;\n"
                                      "    border-color: rgb(41, 213, 202);\n"
                                      f"    border-radius: {4*suma3/Ui_MainWindow.number}%;\n"
                                      "}")
        self.dateEdit_2.setObjectName("dateEdit_2")
        self.verticalLayout_3.addWidget(self.dateEdit_2)
        spacerItem3 = QtWidgets.QSpacerItem(int(20*suma3/Ui_MainWindow.number), int(10*suma3/Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Maximum)
        self.verticalLayout_3.addItem(spacerItem3)
        self.comboBox_KO = QtWidgets.QComboBox(MainWindow)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.comboBox_KO.sizePolicy().hasHeightForWidth())
        self.comboBox_KO.setSizePolicy(sizePolicy)
        self.comboBox_KO.setMinimumSize(QtCore.QSize(int(260*suma3/Ui_MainWindow.number), int(60*suma3/Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14/Ui_MainWindow.number))
        self.comboBox_KO.setFont(font)
        self.comboBox_KO.setStyleSheet(
            "QComboBox{\n"
                "    background-color: rgb(255, 255, 255);\n"
                f"    border-radius: {4 * suma3/Ui_MainWindow.number}%;\n"
                f"    border: 1px solid;\n"
                "}\n"
                "QComboBox:hover{"
                "border-color: rgb(41, 213, 202)"

                "}"
                "\n"
                "QAbstractItemView{\n"
                "  selection-color: black;\n"
                "  selection-background-color: #05F3CC ;\n"
                "  border: 0px ;\n"

                "}\n"
                "QComboBox::drop-down{\n"
                "  border: 0 ;\n"
                "}\n"
            
                "")
        self.comboBox_KO.setObjectName("comboBox_KO")
        self.verticalLayout_3.addWidget(self.comboBox_KO)
        spacerItem4 = QtWidgets.QSpacerItem(int(20*suma3/Ui_MainWindow.number), int(40*suma3/Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Maximum)
        self.verticalLayout_3.addItem(spacerItem4)
        self.gridLayout.addLayout(self.verticalLayout_3, 1, 0, 1, 1)
        self.gridLayout_2.addWidget(self.frame_two, 0, 1, 1, 1)
        self.frame_two.hide()

        # фільтр бар

        #автозаповнення
        # self.gridLayout_2.addWidget(self.frame_two, 0, 2, 1, 1)


        #автозаповнення
        self.frame_avtozapovnena = QtWidgets.QFrame(self.centralwidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.frame.sizePolicy().hasHeightForWidth())
        self.frame_avtozapovnena.setSizePolicy(sizePolicy)
        self.frame_avtozapovnena.setMinimumSize(
                QtCore.QSize(int(280 * suma3 / Ui_MainWindow.number), int(0 * suma3 / Ui_MainWindow.number)))
        self.frame_avtozapovnena.setMaximumSize(
                QtCore.QSize(int(300 * suma3 / Ui_MainWindow.number), int(16777215 * suma3 / Ui_MainWindow.number)))
        self.frame_avtozapovnena.setStyleSheet("QFrame{\n"
                                 f"    border-radius: {2 * suma3 / Ui_MainWindow.number}%;\n"
                                 f"    border: {2 * suma3 / Ui_MainWindow.number}px solid;\n"
                                 "    border-color: rgb(41, 213, 202);\n"
                                 "}")
        self.frame_avtozapovnena.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.frame_avtozapovnena.setFrameShadow(QtWidgets.QFrame.Raised)
        self.frame_avtozapovnena.setObjectName("frame")
        self.frame_avtozapovnena.hide()
        self.gridLayout_avto = QtWidgets.QGridLayout(self.frame_avtozapovnena)
        self.gridLayout_avto.setObjectName("gridLayout_avto")
        self.label_irc = QtWidgets.QCheckBox(self.frame_avtozapovnena)
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        font.setBold(False)
        font.setWeight(50)
        self.label_irc.setFont(font)
        self.label_irc.setObjectName("label_irc")
        self.gridLayout_avto.addWidget(self.label_irc)

        spacerItem_avto_2 = QtWidgets.QSpacerItem(int(20 * suma3 / Ui_MainWindow.number),
                                                  int(20 * suma3 / Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum,
                                                  QtWidgets.QSizePolicy.Fixed)
        self.gridLayout_avto.addItem(spacerItem_avto_2)

        self.edit_URL = QtWidgets.QLineEdit(self.frame_avtozapovnena)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.edit_URL.sizePolicy().hasHeightForWidth())
        self.edit_URL.setSizePolicy(sizePolicy)
        self.edit_URL.setMinimumSize(
                QtCore.QSize(int(250 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        self.edit_URL.setMaximumSize(
                QtCore.QSize(int(300 * suma3 / Ui_MainWindow.number), int(16777215 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.edit_URL.setFont(font)
        self.edit_URL.setStyleSheet("QLineEdit{\n"
                                    "    \n"
                                    "    background-color: rgb(255, 255, 255);\n"
                                    f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                    "    border: 1px solid;\n"
                                    "}\n"
                                    "QLineEdit:hover {\n"
                                    "    border: 1px solid;\n"
                                    "    border-color: rgb(41, 213, 202);\n"
                                    f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                    "\n"
                                    "    \n"
                                    "}")
        self.edit_URL.setObjectName("edit_URL")
        self.gridLayout_avto.addWidget(self.edit_URL)

        self.login = QtWidgets.QLineEdit(self.frame_avtozapovnena)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.login.sizePolicy().hasHeightForWidth())
        self.login.setSizePolicy(sizePolicy)
        self.login.setMinimumSize(
                QtCore.QSize(int(250 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        self.login.setMaximumSize(
                QtCore.QSize(int(300 * suma3 / Ui_MainWindow.number), int(16777215 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.login.setFont(font)
        self.login.setLayoutDirection(QtCore.Qt.LeftToRight)
        self.login.setStyleSheet("QLineEdit{\n"
                                 "    \n"
                                 "    background-color: rgb(255, 255, 255);\n"
                                 f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                 "    border: 1px solid;\n"
                                 "}\n"
                                 "QLineEdit:hover {\n"
                                 "    border: 1px solid;\n"
                                 "    border-color: rgb(41, 213, 202);\n"
                                 f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                 "\n"
                                 "    \n"
                                 "}")
        self.login.setObjectName("login")
        self.gridLayout_avto.addWidget(self.login)

        self.password = QtWidgets.QLineEdit(self.frame_avtozapovnena)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.password.sizePolicy().hasHeightForWidth())
        self.password.setSizePolicy(sizePolicy)
        self.password.setMinimumSize(
                QtCore.QSize(int(250 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        self.password.setMaximumSize(
                QtCore.QSize(int(300 * suma3 / Ui_MainWindow.number), int(16777215 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.password.setFont(font)
        self.password.setStyleSheet("QLineEdit{\n"
                                    "    \n"
                                    "    background-color: rgb(255, 255, 255);\n"
                                    f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                    "    border: 1px solid;\n"
                                    "}\n"
                                    "QLineEdit:hover {\n"
                                    "    border: 1px solid;\n"
                                    "    border-color: rgb(41, 213, 202);\n"
                                    f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                    "\n"
                                    "    \n"
                                    "}")
        self.password.setObjectName("password")
        self.gridLayout_avto.addWidget(self.password)
        spacerItem_avto_1 = QtWidgets.QSpacerItem(int(20 * suma3 / Ui_MainWindow.number),
                                                int(20 * suma3 / Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum,
                                                QtWidgets.QSizePolicy.Fixed)
        self.gridLayout_avto.addItem(spacerItem_avto_1)


        spacerItem_avto = QtWidgets.QSpacerItem(int(20 * suma3 / Ui_MainWindow.number),
                                           int(200 * suma3 / Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum,
                                           QtWidgets.QSizePolicy.Fixed)
        self.gridLayout_avto.addItem(spacerItem_avto)

        self.btn_kompetencia = QtWidgets.QPushButton(self.frame_avtozapovnena)
        self.btn_kompetencia.setMinimumSize(
                QtCore.QSize(int(130 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.btn_kompetencia.setFont(font)
        self.btn_kompetencia.setStyleSheet("QPushButton{\n"
                                           "    \n"
                                           "    background-color: rgb(255, 255, 255);\n"
                                           f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                           "    border: 1px solid;\n"
                                           "}\n"
                                           "QPushButton:hover {\n"
                                           "    border: 1px solid;\n"
                                           "    border-color: rgb(41, 213, 202);\n"
                                           f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                           "\n"
                                           "    \n"
                                           "}")
        self.btn_kompetencia.setObjectName("btn_kompetencia")
        self.gridLayout_avto.addWidget(self.btn_kompetencia)

        self.btn_potrebu = QtWidgets.QPushButton(self.frame_avtozapovnena)
        self.btn_potrebu.setMinimumSize(
                QtCore.QSize(int(130 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.btn_potrebu.setFont(font)
        self.btn_potrebu.setStyleSheet("QPushButton{\n"
                                       "    \n"
                                       "    background-color: rgb(255, 255, 255);\n"
                                       f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                       "    border: 1px solid;\n"
                                       "}\n"
                                       "QPushButton:hover {\n"
                                       "    border: 1px solid;\n"
                                       "    border-color: rgb(41, 213, 202);\n"
                                       f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                       "\n"
                                       "    \n"
                                       "}")
        self.btn_potrebu.setObjectName("btn_potrebu")
        self.gridLayout_avto.addWidget(self.btn_potrebu)

        self.btn_recomendacii = QtWidgets.QPushButton(self.frame_avtozapovnena)
        self.btn_recomendacii.setMinimumSize(
                QtCore.QSize(int(130 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.btn_recomendacii.setFont(font)
        self.btn_recomendacii.setStyleSheet("QPushButton{\n"
                                            "    \n"
                                            "    background-color: rgb(255, 255, 255);\n"
                                            f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                            "    border: 1px solid;\n"
                                            "}\n"
                                            "QPushButton:hover {\n"
                                            "    border: 1px solid;\n"
                                            "    border-color: rgb(41, 213, 202);\n"
                                            f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                            "\n"
                                            "    \n"
                                            "}")
        self.btn_recomendacii.setObjectName("btn_recomendacii")
        self.gridLayout_avto.addWidget(self.btn_recomendacii)
        self.gridLayout_2.addWidget(self.frame_avtozapovnena, int(0), int(2), int(1), int(1))

        # Опитуваня
        self.frame_qestions = QtWidgets.QFrame(self.centralwidget)
        self.frame_qestions.setStyleSheet("QFrame{\n"
                                          f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                          "    border: 2px solid;\n"
                                          "    border-color: rgb(41, 213, 202);\n"
                                          "}")
        self.frame_qestions.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.frame_qestions.setFrameShadow(QtWidgets.QFrame.Raised)
        self.frame_qestions.setObjectName("frame_qestions")
        self.verticalLayout_2 = QtWidgets.QVBoxLayout(self.frame_qestions)
        self.verticalLayout_2.setContentsMargins(int(20 * suma3 / Ui_MainWindow.number), int(-1 * suma3 / Ui_MainWindow.number), int(20 * suma3 / Ui_MainWindow.number), int(-1 * suma3 / Ui_MainWindow.number))
        self.verticalLayout_2.setObjectName("verticalLayout_2")
        self.textBrowser = QtWidgets.QTextBrowser(self.frame_qestions)
        font = QtGui.QFont()
        font.setPointSize(int(24/ Ui_MainWindow.number))
        self.textBrowser.setFont(font)
        self.textBrowser.setStyleSheet("border-color: rgb(85, 170, 255);")
        self.textBrowser.setObjectName("textBrowser")
        self.verticalLayout_2.addWidget(self.textBrowser)
        self.horizontalLayout = QtWidgets.QHBoxLayout()
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.pushButton_yes = QtWidgets.QPushButton(self.frame_qestions)
        self.pushButton_yes.setMinimumSize(QtCore.QSize(int(0 * suma3 / Ui_MainWindow.number), int(100 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(20/ Ui_MainWindow.number))
        self.pushButton_yes.setFont(font)
        self.pushButton_yes.setStyleSheet("QPushButton{\n"
                                          "    \n"
                                          "    background-color: rgb(255, 255, 255);\n"
                                          f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                          "    border: 2px solid;\n"
                                          "}\n"
                                          "QPushButton:hover {\n"
                                          "    border: 2px solid;\n"
                                          "    border-color: rgb(85, 170, 255);\n"
                                          f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                          "    background-color: rgb(41, 213, 202);\n"
                                          "}")
        self.pushButton_yes.setObjectName("pushButton_yes")
        self.pushButton_yes.setText('Так')
        self.horizontalLayout.addWidget(self.pushButton_yes)
        self.pushButton_no = QtWidgets.QPushButton(self.frame_qestions)
        self.pushButton_no.setMinimumSize(QtCore.QSize(int(0 * suma3 / Ui_MainWindow.number), int(100 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(20/ Ui_MainWindow.number))
        font.setBold(False)
        font.setWeight(50)
        self.pushButton_no.setFont(font)
        self.pushButton_no.setStyleSheet("QPushButton{\n"
                                         "    \n"
                                         "    background-color: rgb(255, 255, 255);\n"
                                         f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                         "    border: 2px solid;\n"
                                         "}\n"
                                         "QPushButton:hover {\n"
                                         "    border: 2px solid;\n"
                                         "    border-color: rgb(85, 170, 255);\n"
                                         f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                         "    background-color: rgb(41, 213, 202);\n"
                                         "}")
        self.pushButton_no.setObjectName("pushButton_no")
        self.pushButton_no.setText('Ні')
        self.horizontalLayout.addWidget(self.pushButton_no)
        self.pushButton_mayby = QtWidgets.QPushButton(self.frame_qestions)
        self.pushButton_mayby.setMinimumSize(QtCore.QSize(int(0 * suma3 / Ui_MainWindow.number), int(100 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(20/ Ui_MainWindow.number))
        self.pushButton_mayby.setFont(font)
        self.pushButton_mayby.setStyleSheet("QPushButton{\n"
                                            "    \n"
                                            "    background-color: rgb(255, 255, 255);\n"
                                            f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                            "    border: 2px solid;\n"
                                            "}\n"
                                            "QPushButton:hover {\n"
                                            "    border: 2px solid;\n"
                                            "    border-color: rgb(85, 170, 255);\n"
                                            f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                            "    background-color: rgb(41, 213, 202);\n"
                                            "    \n"
                                            "}")
        self.pushButton_mayby.setObjectName("pushButton_mayby")
        self.pushButton_mayby.setText('Частково')
        self.horizontalLayout.addWidget(self.pushButton_mayby)
        self.pushButton_probel = QtWidgets.QPushButton(self.frame_qestions)
        self.pushButton_probel.setMinimumSize(QtCore.QSize(int(0 * suma3 / Ui_MainWindow.number), int(100 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(20/ Ui_MainWindow.number))
        self.pushButton_probel.setFont(font)
        self.pushButton_probel.setStyleSheet("QPushButton{\n"
                                             "    \n"
                                             "    background-color: rgb(255, 255, 255);\n"
                                             f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                             "    border: 2px solid;\n"
                                             "}\n"
                                             "QPushButton:hover {\n"
                                             "    border: 2px solid;\n"
                                             "    border-color: rgb(85, 170, 255);\n"
                                             f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                             "    background-color: rgb(41, 213, 202);\n"
                                             "    \n"
                                             "}")
        self.pushButton_probel.setObjectName("pushButton_probel")
        self.pushButton_probel.setText('Пропустити')
        self.horizontalLayout.addWidget(self.pushButton_probel)
        self.verticalLayout_2.addLayout(self.horizontalLayout)
        spacerItem = QtWidgets.QSpacerItem(int(20 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Expanding)
        self.verticalLayout_2.addItem(spacerItem)
        self.gridLayout_2.addWidget(self.frame_qestions, 0, 0, 1, 1)
        self.frame_qestions.hide()

        # форма журнал
        self.frame_megassine = QtWidgets.QFrame(self.centralwidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.frame_megassine.sizePolicy().hasHeightForWidth())
        self.frame_megassine.setSizePolicy(sizePolicy)
        self.frame_megassine.setMinimumSize(
            QtCore.QSize(int(300 * suma3 / Ui_MainWindow.number), int(540 * suma3 / Ui_MainWindow.number)))
        self.frame_megassine.setMaximumSize(
            QtCore.QSize(int(300 * suma3 / Ui_MainWindow.number), int(16777215 * suma3 / Ui_MainWindow.number)))
        self.frame_megassine.setStyleSheet("QFrame{\n"
                                           f"    border-radius: { 2* suma3 / Ui_MainWindow.number}%;\n"
                                           f"    border: { 2* suma3 / Ui_MainWindow.number}px solid;\n"
                                           "    border-color: rgb(41, 213, 202);\n"
                                           "}")
        self.frame_megassine.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.frame_megassine.setFrameShadow(QtWidgets.QFrame.Raised)
        self.frame_megassine.setObjectName("frame_megassine")
        self.gridLayout_post = QtWidgets.QGridLayout(self.frame_megassine)
        self.gridLayout_post.setObjectName("gridLayout")
        self.verticalLayout_post = QtWidgets.QVBoxLayout()
        self.verticalLayout_post.setObjectName("verticalLayout_post")
        self.edit_child_name = QtWidgets.QLineEdit(self.frame_megassine)
        self.edit_child_name.setMinimumSize(
            QtCore.QSize(int(0 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.edit_child_name.setFont(font)
        self.edit_child_name.setStyleSheet("QLineEdit{\n"
                                           "    background-color: rgb(255, 255, 255);\n"
                                           f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                           "    border: 1px solid;\n"
                                           "}\n"
                                           "QLineEdit:hover {\n"
                                           "    border: 1px solid;\n"
                                           "    border-color: rgb(41, 213, 202);\n"
                                           f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                           "\n"
                                           "    \n"
                                           "}")
        self.edit_child_name.setObjectName("edit_child_name")
        self.edit_child_name.setPlaceholderText('ПІП')
        self.verticalLayout_post.addWidget(self.edit_child_name)

        self.edit_unical = QtWidgets.QLineEdit(self.frame_megassine)
        self.edit_unical.setMinimumSize(
            QtCore.QSize(int(0 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.edit_unical.setFont(font)
        self.edit_unical.setStyleSheet("QLineEdit{\n"
                                           "    background-color: rgb(255, 255, 255);\n"
                                           f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                           "    border: 1px solid;\n"
                                           "}\n"
                                           "QLineEdit:hover {\n"
                                           "    border: 1px solid;\n"
                                           "    border-color: rgb(41, 213, 202);\n"
                                           f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                           "\n"
                                           "    \n"
                                           "}")
        self.edit_unical.setObjectName("edit_child_name")
        self.edit_unical.setPlaceholderText('Унікальність')
        self.verticalLayout_post.addWidget(self.edit_unical)

        self.add_child_name = QtWidgets.QPushButton(self.frame_megassine)
        self.add_child_name.setMinimumSize(
            QtCore.QSize(int(0 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.add_child_name.setFont(font)
        self.add_child_name.setCursor(QtGui.QCursor(QtCore.Qt.PointingHandCursor))
        self.add_child_name.setStyleSheet("QPushButton{\n"
                                          "    \n"
                                          "    background-color: rgb(255, 255, 255);\n"
                                          f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                          "    border: 1px solid;\n"
                                          "}\n"
                                          "QPushButton:hover {\n"
                                          "    border: 1px solid;\n"
                                          "    border-color: rgb(41, 213, 202);\n"
                                          f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                          "\n"
                                          "    \n"
                                          "}")
        self.add_child_name.setObjectName("add_child_name")
        self.verticalLayout_post.addWidget(self.add_child_name)
        spacerItem = QtWidgets.QSpacerItem(int(20 * suma3 / Ui_MainWindow.number),
                                           int(10 * suma3 / Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum,
                                           QtWidgets.QSizePolicy.Maximum)
        self.verticalLayout_post.addItem(spacerItem)
        self.comboBox_childrean_all = QtWidgets.QComboBox(MainWindow)
        self.comboBox_childrean_all.setMinimumSize(
            QtCore.QSize(int(20 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.comboBox_childrean_all.setFont(font)
        self.comboBox_childrean_all.setStyleSheet(
            "QComboBox{\n"
            "    background-color: rgb(255, 255, 255);\n"
            f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
            f"    border: 1px solid;\n"
            "}\n"
            "QComboBox:hover{"
            "border-color: rgb(41, 213, 202)"

            "}"
            "\n"
            "QAbstractItemView{\n"
            "  selection-color: black;\n"
            "  selection-background-color: #05F3CC ;\n"
            "  border: 0px ;\n"

            "}\n"
            "QComboBox::drop-down{\n"
            "  border: 0 ;\n"
            "}\n"
            "")
        self.comboBox_childrean_all.setObjectName("comboBox_childrean_all")
        self.verticalLayout_post.addWidget(self.comboBox_childrean_all)
        self.btn_delete_child = QtWidgets.QPushButton(self.frame_megassine)
        self.btn_delete_child.setMinimumSize(
            QtCore.QSize(int(0 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.btn_delete_child.setFont(font)
        self.btn_delete_child.setCursor(QtGui.QCursor(QtCore.Qt.PointingHandCursor))
        self.btn_delete_child.setStyleSheet("QPushButton{\n"
                                            "    \n"
                                            "    background-color: rgb(255, 255, 255);\n"
                                            f"    border-radius: { 4* suma3 / Ui_MainWindow.number}%;\n"
                                            "    border: 1px solid;\n"
                                            "}\n"
                                            "QPushButton:hover {\n"
                                            "    border: 1px solid;\n"
                                            "    border-color: rgb(41, 213, 202);\n"
                                            f"    border-radius: { 4* suma3 / Ui_MainWindow.number}%;\n"
                                            "\n"
                                            "    \n"
                                            "}")
        self.btn_delete_child.setObjectName("btn_delete_child")
        self.verticalLayout_post.addWidget(self.btn_delete_child)
        self.btn_update_child = QtWidgets.QPushButton(self.frame_megassine)
        self.btn_update_child.setMinimumSize(
            QtCore.QSize(int(0 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.btn_update_child.setFont(font)
        self.btn_update_child.setCursor(QtGui.QCursor(QtCore.Qt.PointingHandCursor))
        self.btn_update_child.setStyleSheet("QPushButton{\n"
                                            "    \n"
                                            "    background-color: rgb(255, 255, 255);\n"
                                            f"    border-radius: { 4* suma3 / Ui_MainWindow.number}%;\n"
                                            "    border: 1px solid;\n"
                                            "}\n"
                                            "QPushButton:hover {\n"
                                            "    border: 1px solid;\n"
                                            "    border-color: rgb(41, 213, 202);\n"
                                            f"    border-radius: { 4* suma3 / Ui_MainWindow.number}%;\n"
                                            "\n"
                                            "    \n"
                                            "}")
        self.btn_update_child.setObjectName("btn_update_child")
        self.verticalLayout_post.addWidget(self.btn_update_child)



        spacerItem = QtWidgets.QSpacerItem(int(20 * suma3 / Ui_MainWindow.number),
                                           int(20 * suma3 / Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum,
                                           QtWidgets.QSizePolicy.Maximum)
        self.verticalLayout_post.addItem(spacerItem)

        self.dateEdit_krz = QtWidgets.QComboBox(MainWindow)
        self.dateEdit_krz.setMinimumSize(
            QtCore.QSize(int(20 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.dateEdit_krz.setFont(font)
        self.dateEdit_krz.setStyleSheet(
            "QComboBox{\n"
            "    background-color: rgb(255, 255, 255);\n"
            f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
            f"    border: 1px solid;\n"
            "}\n"
            "QComboBox:hover{"
            "border-color: rgb(41, 213, 202)"

            "}"
            "\n"
            "QAbstractItemView{\n"
            "  selection-color: black;\n"
            "  selection-background-color: #05F3CC ;\n"
            "  border: 0px ;\n"

            "}\n"
            "QComboBox::drop-down{\n"
            "  border: 0 ;\n"
            "}\n"
            "")
        self.dateEdit_krz.setObjectName("comboBox_childrean_all")
        self.verticalLayout_post.addWidget(self.dateEdit_krz)


        # self.dateEdit_krz = QtWidgets.QDateEdit(self.frame_megassine)
        # self.dateEdit_krz.setDisplayFormat("yyyy")
        # self.dateEdit_krz.setDate(QtCore.QDate.currentDate())
        # sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Fixed)
        # sizePolicy.setHorizontalStretch(0)
        # sizePolicy.setVerticalStretch(0)
        # sizePolicy.setHeightForWidth(self.dateEdit_krz.sizePolicy().hasHeightForWidth())
        # self.dateEdit_krz.setSizePolicy(sizePolicy)
        # self.dateEdit_krz.setMinimumSize(
        #     QtCore.QSize(int(280 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        # font = QtGui.QFont()
        # font.setPointSize(int(14 / Ui_MainWindow.number))
        # self.dateEdit_krz.setFont(font)
        # self.dateEdit_krz.setStyleSheet("QDateEdit{\n"
        #                             "background-color: rgb(255, 255, 255);\n"
        #                             f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
        #                             "    border: 1px solid;\n"
        #                             "}\n"
        #                             "QDateEdit:hover {\n"
        #                             f"    border: 1px solid;\n"
        #                             "    border-color: rgb(41, 213, 202);\n"
        #                             f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
        #                             "}")
        # self.dateEdit_krz.setObjectName("dateEdit")
        # self.verticalLayout_post.addWidget(self.dateEdit_krz)



        self.btn_vibir_year_krz = QtWidgets.QPushButton(self.frame_megassine)
        self.btn_vibir_year_krz.setMinimumSize(
            QtCore.QSize(int(0 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.btn_vibir_year_krz.setFont(font)
        self.btn_vibir_year_krz.setCursor(QtGui.QCursor(QtCore.Qt.PointingHandCursor))
        self.btn_vibir_year_krz.setStyleSheet("QPushButton{\n"
                                            "    \n"
                                            "    background-color: rgb(255, 255, 255);\n"
                                            f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                            "    border: 1px solid;\n"
                                            "}\n"
                                            "QPushButton:hover {\n"
                                            "    border: 1px solid;\n"
                                            "    border-color: rgb(41, 213, 202);\n"
                                            f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                            "\n"
                                            "    \n"
                                            "}")
        self.btn_vibir_year_krz.setObjectName("btn_vibir_year_krz")
        self.verticalLayout_post.addWidget(self.btn_vibir_year_krz)

        spacerItem1 = QtWidgets.QSpacerItem(int(20 * suma3 / Ui_MainWindow.number),
                                            int(40 * suma3 / Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum,
                                            QtWidgets.QSizePolicy.Expanding)
        self.verticalLayout_post.addItem(spacerItem1)
        self.gridLayout_post.addLayout(self.verticalLayout_post, 0, 0, 1, 1)
        self.gridLayout_2.addWidget(self.frame_megassine, 0, 1, 1, 1)
        self.frame_megassine.hide()

        # форма для створення дописів в журнал
        self.frame_create_gurnal = QtWidgets.QFrame(self.centralwidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.frame_create_gurnal.sizePolicy().hasHeightForWidth())
        self.frame_create_gurnal.setSizePolicy(sizePolicy)
        self.frame_create_gurnal.setMinimumSize(
            QtCore.QSize(int(300 * suma3 / Ui_MainWindow.number), int(540 * suma3 / Ui_MainWindow.number)))
        self.frame_create_gurnal.setMaximumSize(
            QtCore.QSize(int(2000 * suma3 / Ui_MainWindow.number), int(16777215 * suma3 / Ui_MainWindow.number)))
        self.frame_create_gurnal.setStyleSheet("QFrame{\n"
                                               f"    border-radius: { 2* suma3 / Ui_MainWindow.number}%;\n"
                                               f"    border: { 2* suma3 / Ui_MainWindow.number}px solid;\n"
                                               "    border-color: rgb(41, 213, 202);\n"
                                               "}")
        self.frame_create_gurnal.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.frame_create_gurnal.setFrameShadow(QtWidgets.QFrame.Raised)
        self.frame_create_gurnal.setObjectName("frame_create_gurnal")
        self.gridLayout_2_post = QtWidgets.QGridLayout(self.frame_create_gurnal)
        self.gridLayout_2_post.setObjectName("gridLayout_2")
        self.verticalLayout_post_form = QtWidgets.QVBoxLayout()
        self.verticalLayout_post_form.setObjectName("verticalLayout_post_form")
        self.comboBox_form_child_all = QtWidgets.QComboBox(MainWindow)
        self.comboBox_form_child_all.setMinimumSize(
            QtCore.QSize(int(20 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.comboBox_form_child_all.setFont(font)
        self.comboBox_form_child_all.setStyleSheet(
            "QComboBox{\n"
            "    background-color: rgb(255, 255, 255);\n"
            f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
            f"    border: 1px solid;\n"
            "}\n"
            "QComboBox:hover{"
            "border-color: rgb(41, 213, 202)"

            "}"
            "\n"
            "QAbstractItemView{\n"
            "  selection-color: black;\n"
            "  selection-background-color: #05F3CC ;\n"
            "  border: 0px ;\n"

            "}\n"
            "QComboBox::drop-down{\n"
            "  border: 0 ;\n"
            "}\n"
            "")
        self.comboBox_form_child_all.setObjectName("comboBox_form_child_all")
        self.verticalLayout_post_form.addWidget(self.comboBox_form_child_all)
        self.comboBox_form_child_all.hide()
        self.horizontalLayout_2_post = QtWidgets.QHBoxLayout()
        self.horizontalLayout_2_post.setObjectName("horizontalLayout_2_post")
        self.label_id = QtWidgets.QLabel(self.frame_create_gurnal)
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.label_id.setFont(font)
        self.label_id.setStyleSheet("QLabel{\n"
                                    "    border: 2px solid;\n"
                                    "    background-color: rgb(255, 255, 255);\n"
                                    f"    border-radius: { 4* suma3 / Ui_MainWindow.number}%;\n"
                                    "    border-color: rgb(85, 170, 255);\n"
                                    "\n"
                                    "}\n"
                                    "")
        self.label_id.setObjectName("label_id")
        self.horizontalLayout_2_post.addWidget(self.label_id)
        self.lcdNumber_1 = QtWidgets.QLCDNumber(self.frame_create_gurnal)
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.lcdNumber_1.setFont(font)
        self.lcdNumber_1.setStyleSheet("QLCDNumber{\n"
                                       "    border: 2px solid;\n"
                                       "    background-color: rgb(255, 255, 255);\n"
                                       f"    border-radius: { 4* suma3 / Ui_MainWindow.number}%;\n"
                                       "    border-color: rgb(85, 170, 255);\n"
                                       "\n"
                                       "}\n"
                                       "")
        self.lcdNumber_1.setObjectName("lcdNumber_1")
        self.horizontalLayout_2_post.addWidget(self.lcdNumber_1)
        self.label_esistement = QtWidgets.QLabel(self.frame_create_gurnal)
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.label_esistement.setFont(font)
        self.label_esistement.setStyleSheet("QLabel{\n"
                                            "    border: 2px solid;\n"
                                            "    background-color: rgb(255, 255, 255);\n"
                                            f"    border-radius: { 4* suma3 / Ui_MainWindow.number}%;\n"
                                            "    border-color: rgb(85, 170, 255);\n"
                                            "\n"
                                            "}\n"
                                            "")
        self.label_esistement.setObjectName("label_esistement")
        self.horizontalLayout_2_post.addWidget(self.label_esistement)
        self.lcdNumber_2 = QtWidgets.QLCDNumber(self.frame_create_gurnal)
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.lcdNumber_2.setFont(font)
        self.lcdNumber_2.setStyleSheet("QLCDNumber{\n"
                                       "    border: 2px solid;\n"
                                       "    background-color: rgb(255, 255, 255);\n"
                                       f"    border-radius: { 4* suma3 / Ui_MainWindow.number}%;\n"
                                       "    border-color: rgb(85, 170, 255);\n"
                                       "\n"
                                       "}\n"
                                       "")
        self.lcdNumber_2.setObjectName("lcdNumber_2")
        self.horizontalLayout_2_post.addWidget(self.lcdNumber_2)

        self.dateEdit_form = QtWidgets.QDateEdit(self.frame_create_gurnal)
        self.dateEdit_form.setMinimumSize(
            QtCore.QSize(int(0 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.dateEdit_form.setFont(font)
        self.dateEdit_form.setStyleSheet("QDateEdit{\n"
                                         "    background-color: rgb(255, 255, 255);\n"
                                         f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                         "    border: 1px solid;\n"
                                         "}\n"
                                         "QDateEdit:hover {\n"
                                         "    border: 1px solid;\n"
                                         "    border-color: rgb(41, 213, 202);\n"
                                         f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                         "\n"
                                         "    \n"
                                         "}")
        self.dateEdit_form.setObjectName("dateEdit_form")
        self.dateEdit_form.setMinimumSize(QtCore.QSize(int(150 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        self.horizontalLayout_2_post.addWidget(self.dateEdit_form)
        # spacerItem = QtWidgets.QSpacerItem(int(40 * suma3 / Ui_MainWindow.number),
        #                                    int(20 * suma3 / Ui_MainWindow.number), QtWidgets.QSizePolicy.Expanding,
        #                                    QtWidgets.QSizePolicy.Minimum)
        # self.horizontalLayout_2_post.addItem(spacerItem)
        #
        self.verticalLayout_post_form.addLayout(self.horizontalLayout_2_post)
        spacerItem1 = QtWidgets.QSpacerItem(int(20 * suma3 / Ui_MainWindow.number),
                                            int(10 * suma3 / Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum,
                                            QtWidgets.QSizePolicy.Maximum)
        self.verticalLayout_post_form.addItem(spacerItem1)
        self.lineEdit_tema = QtWidgets.QLineEdit(self.frame_create_gurnal)
        self.lineEdit_tema.setMinimumSize(
            QtCore.QSize(int(0 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.lineEdit_tema.setFont(font)
        self.lineEdit_tema.setStyleSheet("QLineEdit{\n"
                                         "    background-color: rgb(255, 255, 255);\n"
                                         f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                         "    border: 1px solid;\n"
                                         "}\n"
                                         "QLineEdit:hover {\n"
                                         "    border: 1px solid;\n"
                                         "    border-color: rgb(41, 213, 202);\n"
                                         f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                         "\n"
                                         "    \n"
                                         "}")
        self.lineEdit_tema.setObjectName("lineEdit_tema")
        self.verticalLayout_post_form.addWidget(self.lineEdit_tema)
        spacerItem2 = QtWidgets.QSpacerItem(int(20 * suma3 / Ui_MainWindow.number),
                                            int(10 * suma3 / Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum,
                                            QtWidgets.QSizePolicy.Maximum)
        self.verticalLayout_post_form.addItem(spacerItem2)

        spacerItem3 = QtWidgets.QSpacerItem(int(20 * suma3 / Ui_MainWindow.number),
                                            int(10 * suma3 / Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum,
                                            QtWidgets.QSizePolicy.Maximum)
        self.verticalLayout_post_form.addItem(spacerItem3)
        self.horizontalLayout_2_post_2 = QtWidgets.QHBoxLayout()
        self.horizontalLayout_2_post_2.setObjectName("horizontalLayout_2_post_2")
        self.textEdit_form_1 = QtWidgets.QTextEdit(self.frame_create_gurnal)
        self.textEdit_form_1.setMinimumSize(
            QtCore.QSize(int(330 * suma3 / Ui_MainWindow.number), int(300 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.textEdit_form_1.setFont(font)
        self.textEdit_form_1.viewport().setProperty("cursor", QtGui.QCursor(QtCore.Qt.IBeamCursor))
        self.textEdit_form_1.setStyleSheet("QTextEdit { \n"
                                           "    border: 2px solid;\n"
                                           "    padding-left:10; \n"
                                           "    padding-top:10; \n"
                                           "    padding-bottom:10; \n"
                                           "    padding-right:10;\n"
                                           "    border-color: rgb(85, 170, 255);\n"
                                           f"    border-radius: { 2* suma3 / Ui_MainWindow.number}%;\n"
                                           "}\n"
                                           "\n"
                                           "")
        self.textEdit_form_1.setObjectName("textEdit_form_1")
        self.horizontalLayout_2_post_2.addWidget(self.textEdit_form_1)

        self.verticalLayout_show_cili = QtWidgets.QVBoxLayout()
        self.verticalLayout_show_cili.setObjectName("verticalLayout_show_cili")
        self.textEdit_form_2 = QtWidgets.QTextEdit(self.frame_create_gurnal)
        self.textEdit_form_2.setMinimumSize(QtCore.QSize(0, 0))
        self.textEdit_form_2.setMaximumSize(QtCore.QSize(int(300* suma3 / Ui_MainWindow.number), int(167772152* suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14/ Ui_MainWindow.number))
        self.textEdit_form_2.setFont(font)
        self.textEdit_form_2.viewport().setProperty("cursor", QtGui.QCursor(QtCore.Qt.IBeamCursor))
        self.textEdit_form_2.setStyleSheet("QTextEdit { \n"
                                           "    border: 2px solid;\n"
                                           "    padding-left:10; \n"
                                           "    padding-top:10; \n"
                                           "    padding-bottom:10; \n"
                                           "    padding-right:10;\n"
                                           "    border-color: rgb(85, 170, 255);\n"
                                           f"    border-radius: {2* suma3 / Ui_MainWindow.number}%;\n"
                                           "}\n"
                                           "\n"
                                           "")
        self.textEdit_form_2.setObjectName("textEdit_form_2")
        self.verticalLayout_show_cili.addWidget(self.textEdit_form_2)
        self.textEdit_form_show_cili = QtWidgets.QTextEdit(self.frame_create_gurnal)
        self.textEdit_form_show_cili.setMinimumSize(QtCore.QSize(0, 0))
        self.textEdit_form_show_cili.setMaximumSize(QtCore.QSize(int(300* suma3 / Ui_MainWindow.number), int(167772152* suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14/ Ui_MainWindow.number))
        self.textEdit_form_show_cili.setFont(font)
        self.textEdit_form_show_cili.viewport().setProperty("cursor", QtGui.QCursor(QtCore.Qt.IBeamCursor))
        self.textEdit_form_show_cili.setStyleSheet("QTextEdit { \n"
                                                   "    border: 2px solid;\n"
                                                   "    padding-left:10; \n"
                                                   "    padding-top:10; \n"
                                                   "    padding-bottom:10; \n"
                                                   "    padding-right:10;\n"
                                                   "    border-color: rgb(85, 170, 255);\n"
                                                   f"    border-radius: {2* suma3 / Ui_MainWindow.number}%;\n"
                                                   "}\n"
                                                   "\n"
                                                   "")
        self.textEdit_form_show_cili.setObjectName("textEdit_form_show_cili")
        self.verticalLayout_show_cili.addWidget(self.textEdit_form_show_cili)
        self.horizontalLayout_2_post_2.addLayout(self.verticalLayout_show_cili)

        self.verticalLayout_post_form.addLayout(self.horizontalLayout_2_post_2)
        self.gridLayout_2_post.addLayout(self.verticalLayout_post_form, 0, 0, 1, 2)
        self.gridLayout_2.addWidget(self.frame_create_gurnal, 0, 0, 1, 1)
        self.frame_create_gurnal.hide()

        # форма редагування журнала
        self.frame_gurnal_update = QtWidgets.QFrame(self.centralwidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.frame_gurnal_update.sizePolicy().hasHeightForWidth())
        self.frame_gurnal_update.setSizePolicy(sizePolicy)
        self.frame_gurnal_update.setMinimumSize(
            QtCore.QSize(int(300 * suma3 / Ui_MainWindow.number), int(540 * suma3 / Ui_MainWindow.number)))
        self.frame_gurnal_update.setMaximumSize(
            QtCore.QSize(int(2000 * suma3 / Ui_MainWindow.number), int(16777215 * suma3 / Ui_MainWindow.number)))
        self.frame_gurnal_update.setStyleSheet("QFrame{\n"
                                               f"    border-radius: {2 * suma3 / Ui_MainWindow.number}%;\n"
                                               f"    border: {2 * suma3 / Ui_MainWindow.number}px solid;\n"
                                               "    border-color: rgb(41, 213, 202);\n"
                                               "}")
        self.frame_gurnal_update.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.frame_gurnal_update.setFrameShadow(QtWidgets.QFrame.Raised)
        self.frame_gurnal_update.setObjectName("frame_gurnal_update")
        self.gridLayout_22_post = QtWidgets.QGridLayout(self.frame_gurnal_update)
        self.gridLayout_22_post.setObjectName("gridLayout_2")
        self.verticalLayout_3_post = QtWidgets.QVBoxLayout()
        self.verticalLayout_3_post.setObjectName("verticalLayout_3_post")
        self.comboBox_updade_all = QtWidgets.QComboBox(MainWindow)
        self.comboBox_updade_all.setMinimumSize(
            QtCore.QSize(int(20 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.comboBox_updade_all.setFont(font)
        self.comboBox_updade_all.setStyleSheet(
            "QComboBox{\n"
            "    background-color: rgb(255, 255, 255);\n"
            f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
            f"    border: 1px solid;\n"
            "}\n"
            "QComboBox:hover{"
            "border-color: rgb(41, 213, 202)"

            "}"
            "\n"
            "QAbstractItemView{\n"
            "  selection-color: black;\n"
            "  selection-background-color: #05F3CC ;\n"
            "  border: 0px ;\n"

            "}\n"
            "QComboBox::drop-down{\n"
            "  border: 0 ;\n"
            "}\n"
            "")
        self.comboBox_updade_all.setObjectName("comboBox_updade_all")
        self.comboBox_updade_all.hide()
        self.verticalLayout_3_post.addWidget(self.comboBox_updade_all)
        self.lineEdit_name_chenge = QtWidgets.QLineEdit(self.frame_gurnal_update)
        self.lineEdit_name_chenge.setMinimumSize(
            QtCore.QSize(int(0 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.lineEdit_name_chenge.setFont(font)
        self.lineEdit_name_chenge.setStyleSheet("QLineEdit{\n"
                                                "    background-color: rgb(255, 255, 255);\n"
                                                f"    border-radius: {3 * suma3 / Ui_MainWindow.number}%;\n"
                                                "    border: 1px solid;\n"
                                                "}\n"
                                                "QLineEdit:hover {\n"
                                                "    border: 1px solid;\n"
                                                "    border-color: rgb(41, 213, 202);\n"
                                                f"    border-radius: {2 * suma3 / Ui_MainWindow.number}%;\n"
                                                "\n"
                                                "    \n"
                                                "}")
        self.lineEdit_name_chenge.setPlaceholderText('Редагувати ПІП:')
        self.lineEdit_name_chenge.setObjectName("lineEdit_name_chenge")
        self.verticalLayout_3_post.addWidget(self.lineEdit_name_chenge)
        self.lineEdit_tema_2 = QtWidgets.QLineEdit(self.frame_gurnal_update)
        self.lineEdit_tema_2.setMinimumSize(
            QtCore.QSize(int(0 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.lineEdit_tema_2.setFont(font)
        self.lineEdit_tema_2.setStyleSheet("QLineEdit{\n"
                                           "    background-color: rgb(255, 255, 255);\n"
                                           f"    border-radius: {2 * suma3 / Ui_MainWindow.number}%;\n"
                                           "    border: 1px solid;\n"
                                           "}\n"
                                           "QLineEdit:hover {\n"
                                           "    border: 1px solid;\n"
                                           "    border-color: rgb(41, 213, 202);\n"
                                           f"    border-radius: {2 * suma3 / Ui_MainWindow.number}%;\n"
                                           "\n"
                                           "    \n"
                                           "}")
        self.lineEdit_tema_2.setObjectName("lineEdit_tema_2")
        self.lineEdit_tema_2.setPlaceholderText('Тема:')
        self.verticalLayout_3_post.addWidget(self.lineEdit_tema_2)
        self.horizontalLayout_2post = QtWidgets.QHBoxLayout()
        self.horizontalLayout_2post.setObjectName("horizontalLayout_2post")
        self.label_id_2 = QtWidgets.QLabel(self.frame_gurnal_update)
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.label_id_2.setFont(font)
        self.label_id_2.setStyleSheet("QLabel{\n"
                                      "    border: 2px solid;\n"
                                      "    background-color: rgb(255, 255, 255);\n"
                                      f"    border-radius: {2 * suma3 / Ui_MainWindow.number}%;\n"
                                      "    border-color: rgb(85, 170, 255);\n"
                                      "\n"
                                      "}\n"
                                      "")
        self.label_id_2.setObjectName("label_id_2")
        self.horizontalLayout_2post.addWidget(self.label_id_2)
        self.lcdNumber_11 = QtWidgets.QLCDNumber(self.frame_gurnal_update)
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.lcdNumber_11.setFont(font)
        self.lcdNumber_11.setStyleSheet("QLCDNumber{\n"
                                       "    border: 2px solid;\n"
                                       "    background-color: rgb(255, 255, 255);\n"
                                       f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                       "    border-color: rgb(85, 170, 255);\n"
                                       "\n"
                                       "}\n"
                                       "")
        self.lcdNumber_11.setObjectName("lcdNumber_1")
        self.horizontalLayout_2post.addWidget(self.lcdNumber_11)
        self.label_esistamete_2 = QtWidgets.QLabel(self.frame_gurnal_update)
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.label_esistamete_2.setFont(font)
        self.label_esistamete_2.setStyleSheet("QLabel{\n"
                                              "    border: 2px solid;\n"
                                              "    background-color: rgb(255, 255, 255);\n"
                                              f"    border-radius: {2 * suma3 / Ui_MainWindow.number}%;\n"
                                              "    border-color: rgb(85, 170, 255);\n"
                                              "\n"
                                              "}\n"
                                              "")
        self.label_esistamete_2.setObjectName("label_esistamete_2")
        self.horizontalLayout_2post.addWidget(self.label_esistamete_2)
        self.lcdNumber_3 = QtWidgets.QLCDNumber(self.frame_gurnal_update)
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.lcdNumber_3.setFont(font)
        self.lcdNumber_3.setStyleSheet("QLCDNumber{\n"
                                       "    border: 2px solid;\n"
                                       "    background-color: rgb(255, 255, 255);\n"
                                       f"    border-radius: {2 * suma3 / Ui_MainWindow.number}%;\n"
                                       "    border-color: rgb(85, 170, 255);\n"
                                       "\n"
                                       "}\n"
                                       "")
        self.lcdNumber_3.setObjectName("lcdNumber_3")
        self.horizontalLayout_2post.addWidget(self.lcdNumber_3)
        # spacerItem = QtWidgets.QSpacerItem(int(40 * suma3 / Ui_MainWindow.number),
        #                                    int(20 * suma3 / Ui_MainWindow.number), QtWidgets.QSizePolicy.Expanding,
        #                                    QtWidgets.QSizePolicy.Minimum)
        # self.horizontalLayout_2post.addItem(spacerItem)
        self.verticalLayout_3_post.addLayout(self.horizontalLayout_2post)

        spacerItem1 = QtWidgets.QSpacerItem(int(20 * suma3 / Ui_MainWindow.number),
                                            int(10 * suma3 / Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum,
                                            QtWidgets.QSizePolicy.Maximum)
        self.verticalLayout_3_post.addItem(spacerItem1)
        spacerItem2 = QtWidgets.QSpacerItem(int(20 * suma3 / Ui_MainWindow.number),
                                            int(10 * suma3 / Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum,
                                            QtWidgets.QSizePolicy.Maximum)
        self.verticalLayout_3_post.addItem(spacerItem2)
        self.dateEdit_form_2 = QtWidgets.QDateEdit(self.frame_gurnal_update)
        self.dateEdit_form_2.setMinimumSize(
            QtCore.QSize(int(0 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.dateEdit_form_2.setFont(font)
        self.dateEdit_form_2.setStyleSheet("QDateEdit{\n"
                                           
                                           "    background-color: rgb(255, 255, 255);\n"
                                           f"    border-radius: {2 * suma3 / Ui_MainWindow.number}%;\n"
                                           "    border: 1px solid;\n"
                                           "}\n"
                                           "QDateEdit:hover {\n"
                                           "    border: 1px solid;\n"
                                           "    border-color: rgb(41, 213, 202);\n"
                                           f"    border-radius: {2 * suma3 / Ui_MainWindow.number}%;\n"
                                           "\n"
                                           "    \n"
                                           "}")
        self.dateEdit_form_2.setObjectName("dateEdit_form_2")
        self.dateEdit_form_2.setMinimumSize(QtCore.QSize(int(150 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        self.horizontalLayout_2post.addWidget(self.dateEdit_form_2)
        spacerItem3 = QtWidgets.QSpacerItem(int(20 * suma3 / Ui_MainWindow.number),
                                            int(10 * suma3 / Ui_MainWindow.number), QtWidgets.QSizePolicy.Minimum,
                                            QtWidgets.QSizePolicy.Maximum)
        self.verticalLayout_3_post.addItem(spacerItem3)
        self.horizontalLayout_3 = QtWidgets.QHBoxLayout()
        self.horizontalLayout_3.setObjectName("horizontalLayout_3")
        self.textEdit_form2 = QtWidgets.QTextEdit(self.frame_gurnal_update)
        self.textEdit_form2.setMinimumSize(
            QtCore.QSize(int(350 * suma3 / Ui_MainWindow.number), int(230 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.textEdit_form2.setFont(font)
        self.textEdit_form2.viewport().setProperty("cursor", QtGui.QCursor(QtCore.Qt.IBeamCursor))
        self.textEdit_form2.setStyleSheet("QTextEdit { \n"
                                          "    border: 1px solid;\n"
                                          "    padding-left:10; \n"
                                          "    padding-top:10; \n"
                                          "    padding-bottom:10; \n"
                                          "    padding-right:10;\n"
                                          "    border-color: rgb(85, 170, 255);\n"
                                          f"    border-radius: {2 * suma3 / Ui_MainWindow.number}%;\n"
                                          "}\n"
                                          "\n"
                                          "")
        self.textEdit_form2.setPlaceholderText('Текст заняття:')
        self.textEdit_form2.setObjectName("textEdit_form2")
        self.horizontalLayout_3.addWidget(self.textEdit_form2)
        self.textEdit_form22 = QtWidgets.QTextEdit(self.frame_gurnal_update)
        self.textEdit_form22.setMinimumSize(QtCore.QSize(int(0 * suma3 / Ui_MainWindow.number), int(230 * suma3 / Ui_MainWindow.number)))
        self.textEdit_form22.setMaximumSize(
            QtCore.QSize(int(250 * suma3 / Ui_MainWindow.number), int(16777215 * suma3 / Ui_MainWindow.number)))
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.textEdit_form22.setFont(font)
        self.textEdit_form22.viewport().setProperty("cursor", QtGui.QCursor(QtCore.Qt.IBeamCursor))
        self.textEdit_form22.setStyleSheet("QTextEdit { \n"
                                           "    border: 1px solid;\n"                                 
                                           "    padding-left:10; \n"
                                           "    padding-top:10; \n"
                                           "    padding-bottom:10; \n"
                                           "    padding-right:10;\n"
                                           "    border-color: rgb(85, 170, 255);\n"
                                           f"    border-radius: {2 * suma3 / Ui_MainWindow.number}%;\n"
                                           "}\n"
                                           "\n"
                                           "")
        self.textEdit_form22.setPlaceholderText('Цілі:')
        self.textEdit_form22.setObjectName("textEdit_form22")
        self.horizontalLayout_3.addWidget(self.textEdit_form22)
        self.verticalLayout_3_post.addLayout(self.horizontalLayout_3)
        self.horizontalLayout_post = QtWidgets.QHBoxLayout()
        self.horizontalLayout_post.setObjectName("horizontalLayout_post")
        self.btn_show = QtWidgets.QPushButton(self.frame_gurnal_update)
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.btn_show.setFont(font)
        self.btn_show.setStyleSheet("QPushButton{\n"
                                    "    \n"
                                    "    background-color: rgb(255, 255, 255);\n"
                                    f"    border-radius: {2 * suma3 / Ui_MainWindow.number}%;\n"
                                    "    border: 1px solid;\n"
                                    "}\n"
                                    "QPushButton:hover {\n"
                                    "    border: 1px solid;\n"
                                    "    border-color: rgb(41, 213, 202);\n"
                                    f"    border-radius: {2 * suma3 / Ui_MainWindow.number}%;\n"
                                    "\n"
                                    "    \n"
                                    "}")
        self.btn_show.setObjectName("btn_show")
        self.btn_show.setMinimumSize(QtCore.QSize(int(0 * suma3 / Ui_MainWindow.number), int(30 * suma3 / Ui_MainWindow.number)))
        self.horizontalLayout_post.addWidget(self.btn_show)
        self.btn_show.hide()
        self.to_word = QtWidgets.QPushButton(self.frame_gurnal_update)
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.to_word.setFont(font)
        self.to_word.setStyleSheet("QPushButton{\n"
                                      "    \n"
                                      "    background-color: rgb(255, 255, 255);\n"
                                      f"    border-radius: {2 * suma3 / Ui_MainWindow.number}%;\n"
                                      "    border: 1px solid;\n"
                                      "}\n"
                                      "QPushButton:hover {\n"
                                      "    border: 1px solid;\n"
                                      "    border-color: rgb(41, 213, 202);\n"
                                      f"    border-radius: {2 * suma3 / Ui_MainWindow.number}%;\n"
                                      "\n"
                                      "    \n"
                                      "}")
        self.to_word.setObjectName("to_word")
        self.to_word.setMinimumSize(QtCore.QSize(int(0 * suma3 / Ui_MainWindow.number), int(30 * suma3 / Ui_MainWindow.number)))
        self.horizontalLayout_post.addWidget(self.to_word)
        # self.btn_previos_2 = QtWidgets.QPushButton(self.frame_gurnal_update)
        # font = QtGui.QFont()
        # font.setPointSize(int(14 / Ui_MainWindow.number))
        # self.btn_previos_2.setFont(font)
        # self.btn_previos_2.setStyleSheet("QPushButton{\n"
        #                                  "    \n"
        #                                  "    background-color: rgb(255, 255, 255);\n"
        #                                  f"    border-radius: {2 * suma3 / Ui_MainWindow.number}%;\n"
        #                                  "    border: 1px solid;\n"
        #                                  "}\n"
        #                                  "QPushButton:hover {\n"
        #                                  "    border: 1px solid;\n"
        #                                  "    border-color: rgb(41, 213, 202);\n"
        #                                  f"    border-radius: {2 * suma3 / Ui_MainWindow.number}%;\n"
        #                                  "\n"
        #                                  "    \n"
        #                                  "}")
        # self.btn_previos_2.setObjectName("btn_previos_2")
        # self.btn_previos_2.setMinimumSize(QtCore.QSize(int(0 * suma3 / Ui_MainWindow.number), int(30 * suma3 / Ui_MainWindow.number)))
        # self.horizontalLayout_post.addWidget(self.btn_previos_2)

        self.btn_delete_post_one = QtWidgets.QPushButton(self.frame_gurnal_update)
        font = QtGui.QFont()
        font.setPointSize(int(14 / Ui_MainWindow.number))
        self.btn_delete_post_one.setFont(font)
        self.btn_delete_post_one.setStyleSheet("QPushButton{\n"
                                         "    \n"
                                         "    background-color: rgb(255, 255, 255);\n"
                                         f"    border-radius: {2 * suma3 / Ui_MainWindow.number}%;\n"
                                         "    border: 1px solid;\n"
                                         "}\n"
                                         "QPushButton:hover {\n"
                                         "    border: 1px solid;\n"
                                         "    border-color: rgb(41, 213, 202);\n"
                                         f"    border-radius: {2 * suma3 / Ui_MainWindow.number}%;\n"
                                         "\n"
                                         "    \n"
                                         "}")
        self.btn_delete_post_one.setObjectName("btn_delete_post_one")
        self.btn_delete_post_one.setText('Видалити')
        self.btn_delete_post_one.setMinimumSize(
            QtCore.QSize(int(0 * suma3 / Ui_MainWindow.number), int(30 * suma3 / Ui_MainWindow.number)))
        self.horizontalLayout_post.addWidget(self.btn_delete_post_one)

        self.verticalLayout_3_post.addLayout(self.horizontalLayout_post)
        self.gridLayout_22_post.addLayout(self.verticalLayout_3_post, 0, 0, 1, 2)
        self.gridLayout_2.addWidget(self.frame_gurnal_update, 0, 0, 1, 1)
        self.frame_gurnal_update.hide()
        MainWindow.setCentralWidget(self.centralwidget)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateUi(self, MainWindow):
        self.minus_number = 0
        self.menu_right_count = 0
        self.plus_number = 1.02
        self.filter_number = 0
        self.next_previos = 0
        self.id_child = None
        self.avto_silenium = 0
        self.site=0
        self.oputyvana_number=0
        self.yes_no_maybe_probel=0
        self.post=0
        self.oputyvana_result=[]
        self.list_anketa = []
        self.post_update=0
        self.count=0
        self.id_post=None
        self.id_children=None
        self.word=0
        self.data_krz = []
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "Create an Estimate"))
        # item = self.tableWidget.verticalHeaderItem(0)
        # item.setText(_translate("MainWindow", "1"))
        item = self.tableWidget.horizontalHeaderItem(0)
        item.setText(_translate("MainWindow", "Питання"))
        item = self.tableWidget.horizontalHeaderItem(1)
        item.setText(_translate("MainWindow", "Відповіді"))
        self.label.setText(_translate("MainWindow", "Дата народження:"))
        self.Povtorna_KO.setText(_translate("MainWindow", "Повторна КО"))
        self.label_2.setText(_translate("MainWindow", "Методика обстеження:"))
        self.School.setText(_translate("MainWindow", "Шкільний вік"))
        self.Doschilnul.setText(_translate("MainWindow", "Дошкільний вік"))
        self.label_3.setText(_translate("MainWindow", "Контакти батьків:"))
        self.GoKO.setText(_translate("MainWindow", "Провести КО"))
        self.add_row.setText(_translate("MainWindow", "Додати рядок"))
        self.Show_Metoduk.setText(_translate("MainWindow", "..."))
        self.PIP.setPlaceholderText('ПІП')
        self.Perent_Kontakt.setPlaceholderText('+380XXXXXXXXX')
        self.poshyk_ko.setText(_translate("MainWindow", "Пошук КО"))
        self.all_ko.setText(_translate("MainWindow", "Всі Ко"))
        self.show_ko.setText(_translate("MainWindow", "Показати КО"))
        self.delete_ko.setText(_translate("MainWindow", "Видалити КО"))
        self.add_row_table.setText(_translate("MainWindow", "Додати рядок"))
        self.update_dani.setText(_translate("MainWindow", "Змінити дані"))
        self.all_dani.setText(_translate("MainWindow", "Всі дані"))
        self.Povtorna_KO_2.setText(_translate("MainWindow", "Повторна КО"))
        self.Pervuna_KO.setText(_translate("MainWindow", "Первина КО"))
        self.label_filter_4.setText(_translate("MainWindow", "Дата народження:"))
        self.label_filter_2.setText(_translate("MainWindow", "Дата проведення:"))
        self.first_name.setPlaceholderText("Призвіще:")
        self.last_name.setPlaceholderText("Ім'я:")
        self.pobatkovi.setPlaceholderText("Побатькові:")
        self.Perent_Kontakt_2.setPlaceholderText('+380XXXXXXXXX')
        self.label_irc.setText(_translate("MainWindow", " Автоматизована система ІРЦ,\n використати попередній пароль?"))
        # self.label_URL.setText(_translate("MainWindow", "URL-адреса"))
        self.btn_kompetencia.setText(_translate("MainWindow", "Заповнити компетенцію"))
        self.btn_potrebu.setText(_translate("MainWindow", "Заповнити потреби"))
        self.btn_recomendacii.setText(_translate("MainWindow", "Заповнити рекомендації"))
        self.login.setPlaceholderText('Логін:')
        self.password.setPlaceholderText('Пароль:')
        self.edit_URL.setPlaceholderText('URL:')
        self.add_child_name.setText(_translate("MainWindow", "Додати"))
        self.btn_delete_child.setText(_translate("MainWindow", "Видалити"))
        self.btn_update_child.setText(_translate("MainWindow", "Редагувати"))
        self.btn_vibir_year_krz.setText(_translate("MainWindow", "Показати за роком"))
        self.label_id.setText(_translate("MainWindow", "ID:"))
        self.label_esistement.setText(_translate("MainWindow", "Кількість занять:"))
        self.textEdit_form_1.setPlaceholderText('Текст Заняття:')
        self.textEdit_form_2.setPlaceholderText('Цілі:')
        self.lineEdit_tema.setPlaceholderText('Тема:')
        self.dateEdit_form.setToolTip(
            _translate("MainWindow", "<html><head/><body><p align=\"center\">Дата проведення</p></body></html>"))
        self.textEdit_form_1.setHtml(_translate("MainWindow",
                                                "<!DOCTYPE HTML PUBLIC \"-//W3C//DTD HTML 4.0//EN\" \"http://www.w3.org/TR/REC-html40/strict.dtd\">\n"
                                                "<html><head><meta name=\"qrichtext\" content=\"1\" /><style type=\"text/css\">\n"
                                                "p, li { white-space: pre-wrap; }\n"
                                                "</style></head><body style=\" font-family:\'MS Shell Dlg 2\'; font-size:14pt; font-weight:400; font-style:normal;\">\n"
                                                "<p align=\"justify\" style=\"-qt-paragraph-type:empty; margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><br /></p></body></html>"))
        self.textEdit_form_2.setHtml(_translate("MainWindow",
                                                "<!DOCTYPE HTML PUBLIC \"-//W3C//DTD HTML 4.0//EN\" \"http://www.w3.org/TR/REC-html40/strict.dtd\">\n"
                                                "<html><head><meta name=\"qrichtext\" content=\"1\" /><style type=\"text/css\">\n"
                                                "p, li { white-space: pre-wrap; }\n"
                                                "</style></head><body style=\" font-family:\'MS Shell Dlg 2\'; font-size:14pt; font-weight:400; font-style:normal;\">\n"
                                                "<p align=\"justify\" style=\"-qt-paragraph-type:empty; margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><br /></p></body></html>"))

        self.label_id_2.setText(_translate("MainWindow", "ID:"))
        self.label_esistamete_2.setText(_translate("MainWindow", "Кількість занять:"))
        self.dateEdit_form_2.setToolTip(
            _translate("MainWindow", "<html><head/><body><p align=\"center\"><br/></p></body></html>"))
        self.textEdit_form2.setHtml(_translate("MainWindow",
                                               "<!DOCTYPE HTML PUBLIC \"-//W3C//DTD HTML 4.0//EN\" \"http://www.w3.org/TR/REC-html40/strict.dtd\">\n"
                                               "<html><head><meta name=\"qrichtext\" content=\"1\" /><style type=\"text/css\">\n"
                                               "p, li { white-space: pre-wrap; }\n"
                                               "</style></head><body style=\" font-family:\'MS Shell Dlg 2\'; font-size:14pt; font-weight:400; font-style:normal;\">\n"
                                               "<p align=\"justify\" style=\"-qt-paragraph-type:empty; margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><br /></p></body></html>"))
        self.textEdit_form22.setHtml(_translate("MainWindow",
                                                "<!DOCTYPE HTML PUBLIC \"-//W3C//DTD HTML 4.0//EN\" \"http://www.w3.org/TR/REC-html40/strict.dtd\">\n"
                                                "<html><head><meta name=\"qrichtext\" content=\"1\" /><style type=\"text/css\">\n"
                                                "p, li { white-space: pre-wrap; }\n"
                                                "</style></head><body style=\" font-family:\'MS Shell Dlg 2\'; font-size:14pt; font-weight:400; font-style:normal;\">\n"
                                                "<p align=\"justify\" style=\"-qt-paragraph-type:empty; margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><br /></p></body></html>"))
        # self.btn_show.setText(_translate("MainWindow", "Показати"))
        self.to_word.setText(_translate("MainWindow", "Зберегти у WORD"))
        # self.btn_previos_2.setText(_translate("MainWindow", "Наступне заняття "))

        self.add_menu.clicked.connect(self.menu_right)
        # self.add_metoduk.clicked.connect(self.show_new_window)
        self.add_metoduk.clicked.connect(self.new_window)

        self.minus.clicked.connect(lambda: self.minus_rozmir('-'))
        self.plus.clicked.connect(lambda: self.minus_rozmir('+'))
        self.GoKO.clicked.connect(self.kompetencia_metoduk)
        self.add_row.clicked.connect(self.add_chenge_to_row)
        self.save.clicked.connect(self.save_to_db)
        self.message.clicked.connect(self.filter_bar)
        self.poshyk_ko.clicked.connect(self.search_KO)
        self.all_ko.clicked.connect(self.all_filter_KO)
        self.show_ko.clicked.connect(self.show_result_KO)
        self.next.clicked.connect(lambda: self.table_show_ko('-'))
        self.previous.clicked.connect(lambda: self.table_show_ko('+'))
        self.all_dani.clicked.connect(self.all_dani_children)
        self.update_dani.clicked.connect(lambda: self.mesage('upd'))
        self.delete_ko.clicked.connect(lambda: self.mesage('del'))
        self.add_row_table.clicked.connect(self.add_row_k_p_r)
        self.btn_kompetencia.clicked.connect(lambda: self.selenium_irc('k'))
        self.btn_potrebu.clicked.connect(lambda: self.selenium_irc('p'))
        self.btn_recomendacii.clicked.connect(lambda: self.selenium_irc('r'))
        self.avto_row.clicked.connect(lambda:self.site_metoduk())
        self.Go_qestions.clicked.connect(self.oputyvana_forma)
        self.pushButton_yes.clicked.connect(lambda:self.vidpovidi('y'))
        self.pushButton_no.clicked.connect(lambda:self.vidpovidi('n'))
        self.pushButton_mayby.clicked.connect(lambda:self.vidpovidi('m'))
        self.pushButton_probel.clicked.connect(lambda:self.vidpovidi('p'))
        self.update_qestions.clicked.connect(self.save_anketa_result)
        self.anketa_show.clicked.connect(self.show_anketa)
        self.delete_anketa.clicked.connect(lambda: self.mesage('del_anketa'))
        self.show_anleta_ko.clicked.connect(lambda: self.show_anketa('result') )
        self.add_post.clicked.connect(lambda: self.show_post())
        self.btn_update_child.clicked.connect(self.update_post)

        # self.comboBox_childrean_all.view().pressed.connect(self.test)

        self.add_child_name.clicked.connect(self.add_child_for_post)
        self.add_child_name.clicked.connect(self.con_children_post_id)


        self.btn_delete_child.clicked.connect(lambda: self.mesage('del_child_post_id'))
        self.btn_delete_child.clicked.connect(self.con_children_post_id)

        # self.comboBox_childrean_all.currentIndexChanged.connect(self.curent_id)

        date=datetime.date.today()
        year, mounth, day = str(date).split('-')
        self.dateEdit_form.setDate(QDate(int(year), int(mounth), int(day)))
        self.btn_show.clicked.connect(self.show_post_update)
        self.next_2.clicked.connect(lambda:self.next_previos_post('+'))
        self.previous_2.clicked.connect(lambda:self.next_previos_post('-'))
        self.btn_delete_post_one.clicked.connect(lambda: self.mesage('del_post_tema_all'))
        self.to_word.clicked.connect(self.save_word_posts)
        self.error.clicked.connect(self.hide_error)
        self.rezervne.clicked.connect(self.copy_db_rezerv)
        self.comboBox_childrean_all.activated.connect(self.curent_id)
        self.btn_vibir_year_krz.clicked.connect(self.show_krz_data)


    def show_krz_data(self):
        self.word=1
        try:
            year=self.dateEdit_krz.currentText()
            self.count = 0
            connect = connect('IRCdb.db')
            id = self.con_child_id()
            self.tema = []
            self.hid_work = []
            self.data = []
            self.cili = []
            self.tema.clear()
            self.hid_work.clear()
            self.data.clear()
            self.cili.clear()
            self.textEdit_form22.clear()
            with connect as con:
                cursor = con.cursor()
                cursor.execute(f" SELECT tema, hid_work, data_edit FROM children_post WHERE  id_children={int(id)} AND data_edit LIKE '%{year}%'")
                for i in cursor:
                    self.tema.append(i[0])
                    self.hid_work.append(i[1])
                    self.data.append(i[2])
            self.tema.reverse()
            self.hid_work.reverse()
            self.data.reverse()
            self.lcdNumber_3.display(len(self.data))
            with connect as con:
                cursor = con.cursor()
                cursor.execute(f" SELECT cili FROM children_post_cili WHERE  id_children={int(id)} ")
                for i in cursor:
                    self.cili.append(i[0])
            self.textEdit_form2.setText(self.hid_work[0])
            self.lineEdit_tema_2.setText(self.tema[0])
            year, mounth, day = str(self.data[0]).split('.')
            self.dateEdit_form_2.setDate(QDate(int(day), int(mounth), int(year)))

            for i in self.cili:
                if i == "":
                    pass
                else:
                    self.textEdit_form22.append(i)

            with connect as con:
                cursor = con.cursor()
                cursor.execute(
                    f"""SELECT id, id_children FROM children_post WHERE tema='{self.tema[self.count]}' AND hid_work='{self.hid_work[self.count]}' 
                    AND data_edit='{self.data[self.count]}'""")
                id = cursor.fetchone()
                self.id_post = id[0]
                self.id_children = id[1]



        except Exception as ex:
            self.error.show()
            self.mesage_output('show_post_update')
            print('show_post_update', ex)

    def copy_db_rezerv(self):
        filename = QFileDialog.getSaveFileName(directory = 'IRCdb')
        self.path2 = filename[0]
        shutil.copyfile('IRCdb.db', f'{self.path2+".db"}')

    def hide_error(self):
        self.error.hide()

    def save_word_posts(self):
        try:
            year=self.dateEdit_krz.currentText()
            connect = connect('IRCdb.db')
            self.tema2 = [ ]
            self.hid_work2 = [ ]
            self.data2 = [ ]
            self.cili2 = [ ]
            if self.word==0:
                with connect as con:
                    cursor = con.cursor()
                    cursor.execute(
                        f" SELECT tema, hid_work, data_edit FROM children_post WHERE  id_children={int(self.id_children)} ")
                    for i in cursor:
                        self.tema2.append(i [ 0 ])
                        self.hid_work2.append(i [ 1 ])
                        self.data2.append(i [ 2 ])
                with connect as con:
                    cursor = con.cursor()
                    cursor.execute(f" SELECT cili FROM children_post_cili WHERE  id_children={int(self.id_children)} ")
                    for i in cursor:
                        self.cili2.append(i [ 0 ])
            else:
                with connect as con:
                    cursor = con.cursor()
                    cursor.execute(
                        f" SELECT tema, hid_work, data_edit FROM children_post WHERE  id_children={int(self.id_children)} AND data_edit LIKE '%{year}%' ")
                    for i in cursor:
                        self.tema2.append(i [ 0 ])
                        self.hid_work2.append(i [ 1 ])
                        self.data2.append(i [ 2 ])
                with connect as con:
                    cursor = con.cursor()
                    cursor.execute(f" SELECT cili FROM children_post_cili WHERE  id_children={int(self.id_children)}")
                    for i in cursor:
                        self.cili2.append(i [ 0 ])

            self.doc = docx.Document()
            section = self.doc.sections [ 0 ]
            section.page_height = Cm(29.7)
            section.page_width = Cm(21.0)
            section.left_margin = Mm(15.0)
            section.right_margin = Mm(15.0)
            section.top_margin = Mm(10)
            section.bottom_margin = Mm(10)
            section.header_distance = Mm(10)
            section.footer_distance = Mm(10)
            text_name_combo = self.comboBox_childrean_all.currentText().replace(" (", " ").replace(")", "")
            name_list = text_name_combo.split()
            self.name_add = ''
            for i in name_list:
                self.name_add += str(i) + ' '

            self.doc.styles [ 'Normal' ].font.name = "Arial"
            t = self.doc.add_paragraph()
            t.alignment = 1

            runner = t.add_run(self.name_add)
            runner.font.size = Pt(24)
            table2 = self.doc.add_table(rows=1, cols=2, style='Table Grid')
            row = table2.rows [ 0 ].cells
            row [ 0 ].text = '№'
            row [ 1 ].text = 'Цілі'

            for i in range(len(self.cili2)):
                row = table2.add_row().cells
                table2.cell(i, 0).width = Inches(1.5)
                table2.cell(i, 1).width = Cm(25.0)
                row [ 0 ].text = str(i + 1) + '.'
                row [ 1 ].text = self.cili2 [ i ]

            for row in table2.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(14)

            for i in range(2):
                hdr_cells = table2.rows [ 0 ].cells
                paragraph = hdr_cells [ i ].paragraphs [ 0 ]
                run = paragraph.runs
                font = run [ 0 ].font
                font.size = Pt(14)
                font.bold = True

            self.doc.styles [ 'Normal' ].font.name = "Arial"
            t = self.doc.add_paragraph()
            t.alignment = 1
            runner = t.add_run(' ')
            runner.font.size = Pt(24)

            table = self.doc.add_table(rows=1, cols=4, style='Table Grid')

            row = table.rows [ 0 ].cells
            row [ 0 ].text = '№'
            row [ 1 ].text = 'Дата'
            row [ 2 ].text = 'Тема'
            row [ 3 ].text = 'Зміст'

            for i in range(len(self.tema2)):
                row = table.add_row().cells
                table.cell(i, 0).width = Inches(1.5)
                table.cell(i, 1).width = Cm(4.0)
                table.cell(i, 2).width = Cm(25.0)
                table.cell(i, 3).width = Cm(25.0)
                row [ 0 ].text = str(i + 1) + '.'
                row [ 1 ].text = self.data2 [ i ]
                row [ 2 ].text = self.tema2 [ i ]
                row [ 3 ].text = self.hid_work2 [ i ]

            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(14)

            for i in range(4):
                hdr_cells = table.rows [ 0 ].cells
                paragraph = hdr_cells [ i ].paragraphs [ 0 ]
                run = paragraph.runs
                font = run [ 0 ].font
                font.size = Pt(14)
                font.bold = True

            filename = QFileDialog.getSaveFileName( directory = f'{self.name_add}')
            self.path = filename [ 0 ]
            self.doc.save(self.path +  '.docx')
        except Exception as ex:
            self.error.show()
            self.mesage_output('save_word_posts')
            print('save_word_posts', ex)

    def delete_post(self):
        try:
            connect = connect('IRCdb.db')
            with connect as con:
                cursor = con.cursor()
                cursor.execute(f""" DELETE FROM children_post WHERE id = {self.id_post} """)
            self.show_post_update()
        except Exception as ex:
            self.error.show()
            self.mesage_output('delete_post')
            print('delete_post', ex)

    def update_post_cili(self):
        try:
            connect = connect('IRCdb.db')
            tema = self.lineEdit_tema_2.text().replace("'","`")
            text = self.textEdit_form2.toPlainText().replace("'","`")
            data = self.dateEdit_form_2.text()
            cili = self.textEdit_form22.toPlainText().split('\n')
            with connect as con:
                cursor=con.cursor()
                cursor.execute(f""" UPDATE children_post SET tema='{tema}', hid_work='{text}', data_edit='{data}' WHERE id = {self.id_post} """)

            with connect as con:
                cursor=con.cursor()
                cursor.execute(f""" DELETE FROM children_post_cili WHERE id_children = {self.id_children} """)
            if "" in cili:
                pass
            elif " " in cili:
                pass
            else:
                with connect as con:
                    cursor=con.cursor()
                    for i in cili:
                        cursor.execute(f""" INSERT INTO children_post_cili  (cili, id_children) VALUES ('{i.replace("'","`")}', {self.id_children}); """)
            try:
                if self.lineEdit_name_chenge.text()=='':
                    pass
                else:
                    name = self.lineEdit_name_chenge.text().split()
                    with connect as con:
                        cursor = con.cursor()
                        cursor.execute(
                            f""" UPDATE children_post_id SET first_name='{name[0].replace("'","`")}', 
                            last_name='{name[1].replace("'","`")}', 
                            pobatkovi='{name[2].replace("'","`")}' WHERE id = {self.id_children} """)
                        self.lineEdit_name_chenge.clear()
                    self.con_children_post_id()
                    self.show_post_update()
                self.mesage_output('update_post')
                self.lineEdit_name_chenge.setStyleSheet("QLineEdit{\n"
                                                   "    border: 1px solid;\n"
                                                   "    border-color: rgb(41, 213, 202);\n"
                                                   f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"

                                                   "}\n"
                                                   "QLineEdit:hover {\n"
                                                   "    border: 1px solid;\n"
                                                   "    border-color: rgb(41, 213, 202);\n"
                                                   f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                                   "\n"
                                                   "    \n"
                                                   "}")

            except Exception as ex:
                self.lineEdit_name_chenge.setStyleSheet("QLineEdit{\n"
                                                   "    border: 1px solid;\n"
                                                   "    border-color: red;\n"
                                                   f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"

                                                   "}\n"
                                                   "QLineEdit:hover {\n"
                                                   "    border: 1px solid;\n"
                                                   "    border-color: rgb(41, 213, 202);\n"
                                                   f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                                   "\n"
                                                   "    \n"
                                                   "}")
                print('self.lineEdit_name_chenge EROR..', ex)

        except Exception as ex:
            self.error.show()
            self.mesage_output('update_post_')
            print('update_post', ex)

    def next_previos_post(self,x):
        try:
            len_n = len(self.tema)-1

            if x =='+':
                self.count+=1
            else:
                self.count-=1

            if self.count > len_n:
                self.count = len_n
            elif self.count < 0:
                self.count = 0
            number=[]
            for i in range(len_n+2):
                if i == 0:
                    pass
                else:
                    number.append(i)
            number.reverse()
            self.lcdNumber_3.display(number[self.count])

            connect = connect('IRCdb.db')
            self.textEdit_form2.setText(self.hid_work[self.count])
            self.lineEdit_tema_2.setText(self.tema[self.count])
            year, mounth, day = str(self.data[self.count]).split('.')
            self.dateEdit_form_2.setDate(QDate(int(day), int(mounth), int(year)))

            with connect as con:
                cursor = con.cursor()
                cursor.execute(f"""SELECT id, id_children FROM children_post WHERE tema='{self.tema[self.count]}' AND hid_work='{self.hid_work[self.count]}'
                AND data_edit='{self.data[self.count]}'""")
                id = cursor.fetchone()
                self.id_post=id[0]
                self.id_children = id[1]
        except Exception as ex:
            self.error.show()
            self.mesage_output('next_previos_post')
            print('next_previos_post', ex)

    def show_post_update(self):
        try:
            self.count=0
            connect = connect('IRCdb.db')
            id = self.con_child_id()
            self.tema=[]
            self.hid_work=[]
            self.data = []
            self.cili=[]
            self.tema.clear()
            self.hid_work.clear()
            self.data.clear()
            self.cili.clear()
            self.textEdit_form22.clear()
            self.dateEdit_krz.clear()
            with connect as con:
                cursor = con.cursor()
                cursor.execute(f"SELECT data_edit FROM children_post WHERE id_children = {int(id)}")

                unique_years = set()  # створення set для збереження унікальних років

                for i in cursor:
                    year = i[0][6:10]  # отримання року з дати у форматі 'YYYY-MM-DD'
                    unique_years.add(year)  # додавання року до set

                for year in unique_years:
                    self.dateEdit_krz.addItem(str(year))

            with connect as con:
                cursor = con.cursor()
                cursor.execute(f" SELECT tema, hid_work, data_edit FROM children_post WHERE  id_children={int(id)} ")
                for i in cursor:
                    self.tema.append(i[0])
                    self.hid_work.append(i[1])
                    self.data.append(i[2])
            self.tema.reverse()
            self.hid_work.reverse()
            self.data.reverse()
            with connect as con:
                cursor = con.cursor()
                cursor.execute(f" SELECT cili FROM children_post_cili WHERE  id_children={int(id)} ")
                for i in cursor:
                    self.cili.append(i[0])
            self.textEdit_form2.setText(self.hid_work[0])
            self.lineEdit_tema_2.setText(self.tema[0])
            year, mounth, day = str(self.data[0]).split('.')
            self.dateEdit_form_2.setDate(QDate(int(day), int(mounth), int(year)))

            for i in self.cili:
                if i=="":
                    pass
                else:
                    self.textEdit_form22.append(i)

            with connect as con:
                cursor = con.cursor()
                cursor.execute(
                    f"""SELECT id, id_children FROM children_post WHERE tema='{self.tema[self.count]}' AND hid_work='{self.hid_work[self.count]}' 
                    AND data_edit='{self.data[self.count]}'""")
                id = cursor.fetchone()
                self.id_post = id[0]
                self.id_children=id[1]

        except Exception as ex:
            self.error.show()
            self.mesage_output('show_post_update')
            print('show_post_update',ex)

    def add_post_for_table(self):
        try:
            id = int(self.con_child_id())
            tema = self.lineEdit_tema.text().replace("'","`")
            text_forma1= self.textEdit_form_1.toPlainText().replace("'","`")
            text_cili = self.textEdit_form_2.toPlainText().replace("'","`")
            data = self.dateEdit_form.text()
            connect = connect('IRCdb.db')
            with connect as con:
                cursor = con.cursor()
                cursor.execute(f"INSERT INTO  children_post (tema, hid_work, data_edit, id_children) VALUES "
                               f" ('{tema}', '{text_forma1}', '{data}', {id}) ;")
            if text_cili == '':
                pass
            else:
                with connect as con:
                    cursor = con.cursor()
                    cursor.execute(f"INSERT INTO  children_post_cili (cili, id_children) VALUES "
                                   f" ('{text_cili}', {id}) ;")
            self.lineEdit_tema.clear()
            self.textEdit_form_1.clear()
            self.textEdit_form_2.clear()
            self.curent_id()
        except Exception as ex:
            self.error.show()
            self.mesage_output('add_post_for_table')
            print('add_post_for_table',ex)

    def con_child_id(self):
        try:
            connect = connect('IRCdb.db')
            text_name_combo = self.comboBox_childrean_all.currentText().replace(" (", " ").replace(")", "")
            dani_name = text_name_combo.split()
            if len(dani_name) == 3:
                first_name, last_name, pobatkovi = text_name_combo.split(maxsplit=4)
                unicalinist = 'None'
            else:
                first_name, last_name, pobatkovi, unicalinist = text_name_combo.split(maxsplit=3)

            with connect as con:
                cursor = con.cursor()
                cursor.execute(
                    f"SELECT id FROM  children_post_id WHERE "                                                                                  
                    f"first_name = '{first_name}' AND last_name = '{last_name}'AND pobatkovi = '{pobatkovi}' AND unicalinist='{unicalinist}';")
                id_child = cursor.fetchone()
                id_child_post = id_child[0]
                return id_child_post
        except Exception as ex:
            self.error.show()
            self.mesage_output('con_child_id')
            print('con_child_id', ex)

    def delete_post_child_id(self):
        try:
            connect = connect('IRCdb.db')
            text_name_combo = self.comboBox_childrean_all.currentText().replace(" (", " ").replace(")","")
            dani_name = text_name_combo.split()
            if len(dani_name)==3:
                first_name, last_name, pobatkovi = text_name_combo.split(maxsplit=4)
                unicalinist='None'
            else:
                first_name, last_name, pobatkovi, unicalinist = text_name_combo.split(maxsplit=3)
            with connect as con:
                cursor = con.cursor()
                cursor.execute('pragma foreign_keys=on')
                cursor.execute(
                    f"DELETE FROM  children_post_id WHERE "
                    f"first_name = '{first_name}' AND last_name = '{last_name}'AND pobatkovi = '{pobatkovi}' AND unicalinist='{unicalinist}';")
        except Exception as ex:
            self.error.show()
            self.mesage_output('delete_post_child_id')
            print('delete_post_child_id', ex)

    def add_child_for_post(self):
        try:
            child_name=self.edit_child_name.text().replace("'","`")
            first_name, last_name, pobatkovi = child_name.split()
            unicalinist = self.edit_unical.text().replace("'","`")
            connect = connect('IRCdb.db')

            if unicalinist == '':
                unicalinist='None'
            with connect as con:
                cursor = con.cursor()
                cursor.execute(f" SELECT id FROM children_post_id WHERE first_name = '{first_name}' AND last_name = '{last_name}'AND pobatkovi = '{pobatkovi}' AND unicalinist='{unicalinist}' ")
                text = cursor.fetchone()

                if text != None:
                    self.edit_unical.setPlaceholderText("Ім'я існує, додайте унікальність")
                    self.edit_unical.setStyleSheet("QLineEdit{\n"
                                                       "    border: 1px solid;\n"
                                                       "    border-color: red;\n"
                                                       f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
    
                                                       "}\n"
                                                       "QLineEdit:hover {\n"
                                                       "    border: 1px solid;\n"
                                                       "    border-color: rgb(41, 213, 202);\n"
                                                       f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                                       "\n"
                                                       "    \n"
                                                       "}")
                    pass
                else:
                    cursor.execute(f"INSERT INTO  children_post_id (first_name, last_name, pobatkovi, unicalinist) VALUES "
                            f" ('{first_name}', '{last_name}', '{pobatkovi}', '{unicalinist}') ;")

                    self.edit_unical.setStyleSheet("QLineEdit{\n"
                                                       "    background-color: rgb(255, 255, 255);\n"
                                                       f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                                       "    border: 1px solid;\n"
                                                       "}\n"
                                                       "QLineEdit:hover {\n"
                                                       "    border: 1px solid;\n"
                                                       "    border-color: rgb(41, 213, 202);\n"
                                                       f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                                       "\n"
                                                       "    \n"
                                                       "}")
                    self.edit_child_name.clear()
                    self.edit_unical.clear()
                self.edit_child_name.setStyleSheet("QLineEdit{\n"
                                               "    background-color: rgb(255, 255, 255);\n"
                                               f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                               "    border: 1px solid;\n"
                                               "}\n"
                                               "QLineEdit:hover {\n"
                                               "    border: 1px solid;\n"
                                               "    border-color: rgb(41, 213, 202);\n"
                                               f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                               "\n"
                                               "    \n"
                                               "}")

        except Exception as ex:
            self.edit_child_name.setStyleSheet("QLineEdit{\n"
                                           "    border: 1px solid;\n"
                                           "    border-color: red;\n"
                                           f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"

                                           "}\n"
                                           "QLineEdit:hover {\n"
                                           "    border: 1px solid;\n"
                                           "    border-color: rgb(41, 213, 202);\n"
                                           f"    border-radius: {4 * suma3 / Ui_MainWindow.number}%;\n"
                                           "\n"
                                           "    \n"
                                           "}")
            self.error.show()
            self.mesage_output('add_child_for_post')
            print('add_child_for_post',ex)

    def con_children_post_id(self):
        try:
            connect = connect('IRCdb.db')
            self.comboBox_childrean_all.clear()
            self.slovnuk_child={}
            with connect as con:
                cursor = con.cursor()
                cursor.execute(f"SELECT first_name, last_name,  pobatkovi, unicalinist, id FROM  children_post_id ;")
                for i in cursor:
                    if str(i[3])!='None':
                        name = i [ 0 ] + " " + i [ 1 ] + " " + i [ 2 ]+" (" + str(i [ 3 ]) + ")"
                    else:
                        name = i[0]+" "+i[1]+" "+i[2]
                    id = i[4]
                    self.slovnuk_child[name]=id
                for i in self.slovnuk_child:
                    self.comboBox_childrean_all.addItem(i)
        except Exception as ex:
            self.error.show()
            self.mesage_output('con_children_post_id')
            print('con_children_post_id',ex)

    def curent_id(self):
        self.word=0
        self.textEdit_form_show_cili.clear()
        try:
            connect = connect('IRCdb.db')
            id = self.comboBox_childrean_all.currentText()
            if self.post_update==1:
                self.lcdNumber_11.display(self.slovnuk_child.get(id))
            else:
                self.lcdNumber_1.display(self.slovnuk_child.get(id))
            id_c = int(self.con_child_id())
            with connect as con:
                cursor = con.cursor()
                cursor.execute(f'SELECT COUNT(1) FROM  children_post WHERE id_children ={int(id_c)};')
                c = cursor.fetchone()
                count=c[0]
                if self.post_update == 1:
                    self.lcdNumber_3.display(count)
                    self.show_post_update()
                else:
                    self.lcdNumber_2.display(count)
                    connect = connect('IRCdb.db')
                    with connect as con:
                        cursor = con.cursor()
                        cursor.execute(f" SELECT cili FROM children_post_cili WHERE  id_children={int(id_c)} ")
                        for i in cursor:
                            text=i[0]
                            if text == "":
                                pass
                            else:
                                self.textEdit_form_show_cili.append(text)
        except Exception as ex:
            self.error.show()
            self.mesage_output('curent_id')
            print('curent_id', ex)

    def update_post(self):
        if self.post_update==0:
            self.post_update = 1
            self.post=0
            self.tableWidget.hide()
            self.frame_gurnal_update.show()
            self.curent_id()
            self.minus.hide()
            self.plus.hide()
            self.next_2.show()
            self.previous_2.show()
        else:
            self.post_update = 0
            self.post=1
            self.frame_gurnal_update.hide()
            self.frame_create_gurnal.show()
            self.minus.show()
            self.plus.show()
            self.next_2.hide()
            self.previous_2.hide()
            self.curent_id()

    def show_post(self):
        if self.post==0:
            self.hide_widget()
            self.post=1
            self.frame_megassine.show()
            self.tableWidget.hide()
            self.frame_create_gurnal.show()
            self.con_children_post_id()
            self.curent_id()
            self.add_post.setIcon(QIcon('жrgb.png'))
            self.add_post.setIconSize(
                QSize(int(40 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
        else:
            self.post=0
            self.tableWidget.show()
            self.frame_create_gurnal.hide()
            self.frame_megassine.hide()
            self.add_post.setIcon(QIcon('ж.png'))
            self.add_post.setIconSize(
                QSize(int(40 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))

    def delete_answer_anketa(self):
        try:
            connect = connect('IRCdb.db')
            # data_ko_search = self.dateEdit_2.text().split('.')
            # date = data_ko_search [ 2 ] + '-' + data_ko_search [ 1 ] + '-' + data_ko_search [ 0 ]
            # text = self.comboBox_KO.currentText()
            # first_name, second_name, pobatkovi = text.split()
            name_all = self.comboBox_KO.currentText().replace("/", " ")
            first_name, second_name, pobatkovi, birth, date_edit = name_all.split()
            with connect as con:
                cursor = con.cursor()
                # if self.label_filter_2.isChecked():
                #     cursor.execute(f'DELETE FROM  oputevana_resultat WHERE first_name="{first_name}" AND last_name="{second_name}" AND pobatkovi="{pobatkovi}" AND birth_day="{birth}";')
                # else:
                cursor.execute(f'DELETE FROM oputevana_resultat WHERE first_name="{first_name}" AND last_name="{second_name}" AND pobatkovi="{pobatkovi}" AND birth_day="{birth}"'
                                   f'AND date_edit="{date_edit}";')
        except Exception as ex:
            self.error.show()
            self.mesage_output('delete_answer_anketa')
            print('delete_answer_anketa', ex)

    def show_anketa(self,x):
        try:
            self.table_restart()
            row_result = self.tableWidget.rowCount()
            for i in range(row_result):
                self.tableWidget.removeRow(i)
            anketa_qestions=[]
            anketa_vidpovido=[]
            connect = connect('IRCdb.db')
            text = None
            if x == 'result':
                text = self.PIP.text()
                first_name, second_name, pobatkovi = text.split()
                birth = self.dateEdit.text()
                date_edit = datetime.date.today()
            else:
                # text = self.comboBox_KO.currentText()
                name_all = self.comboBox_KO.currentText().replace("/", " ")
                first_name, second_name, pobatkovi, birth, date_edit = name_all.split()

            # if x == 'result':
            #     data_ko_search = self.dateEdit.text().split('.')
            # else:
            #     data_ko_search = self.dateEdit_2.text().split('.')
            # date = data_ko_search [ 2 ] + '-' + data_ko_search [ 1 ] + '-' + data_ko_search [ 0 ]

            with connect as con:
                cursor = con.cursor()
                item = self.tableWidget.horizontalHeaderItem(0)
                item.setText("Опитування")
                item = self.tableWidget.horizontalHeaderItem(1)
                item.setText("Відповіді")
                # if self.label_filter_2.isChecked():
                #     cursor.execute(f'SELECT COUNT(1) FROM  oputevana_resultat WHERE first_name="{first_name}" AND last_name="{second_name}" AND pobatkovi="{pobatkovi}" AND birth_day="{birth}";')
                # else:
                cursor.execute(f'SELECT COUNT(1) FROM  oputevana_resultat WHERE first_name="{first_name}" AND last_name="{second_name}" AND pobatkovi="{pobatkovi}"'
                                   f'AND birth_day="{birth}" AND date_edit="{date_edit}";')
                count = cursor.fetchone()
                count_kompetencia = count[0]
                self.tableWidget.setRowCount(count_kompetencia)

            with connect as con:
                cursor = con.cursor()
                if self.label_filter_2.isChecked():
                    cursor.execute(f'SELECT putana FROM  oputevana_resultat WHERE first_name="{first_name}" AND last_name="{second_name}" AND pobatkovi="{pobatkovi}" AND birth_day="{birth}";')
                else:
                    cursor.execute(f'SELECT putana FROM  oputevana_resultat WHERE first_name="{first_name}" AND last_name="{second_name}" AND pobatkovi="{pobatkovi}"'
                                   f'AND birth_day="{birth}" AND date_edit="{date_edit}";')
                for i in cursor:
                    anketa_qestions.append(i[0])

            with connect as con:
                cursor = con.cursor()
                if self.label_filter_2.isChecked():
                    cursor.execute(f'SELECT answer FROM  oputevana_resultat WHERE first_name="{first_name}" AND last_name="{second_name}" AND pobatkovi="{pobatkovi}" AND birth_day="{birth}";')
                else:
                    cursor.execute(f'SELECT answer FROM  oputevana_resultat WHERE first_name="{first_name}" AND last_name="{second_name}" AND pobatkovi="{pobatkovi}"'
                                   f'AND birth_day="{birth}" AND date_edit="{date_edit}";')
                for i in cursor:
                    anketa_vidpovido.append(i[0])

            for i in range(count_kompetencia):
                self.tableWidget.setItem(i, 0, QTableWidgetItem(anketa_qestions[i]))
                self.tableWidget.setItem(i, 1, QTableWidgetItem(anketa_vidpovido [ i ]))
        except Exception as ex:
            self.error.show()
            self.mesage_output('show_anketa')
            print('show_anketa', ex)

    def save_anketa_result(self):
        try:
            connect = connect('IRCdb.db')
            first_name, second_name, pobatkovi = self.PIP.text().split()
            data_birth_day = self.dateEdit.text()
            date_edit = datetime.date.today()
            with connect as con:
                cursor = con.cursor()
                for i in range(len(self.list_anketa)):
                    cursor.execute(f"INSERT INTO  oputevana_resultat (first_name, last_name, pobatkovi, birth_day, putana, answer, date_edit) VALUES "
                                   f" ('{first_name}', '{second_name}', '{pobatkovi}', '{data_birth_day}', '{self.list_anketa[i]}', '{self.oputyvana_result[i]}','{date_edit}') ;")
            self.textBrowser.setText('Збережено!')
        except Exception as ex:
            self.error.show()
            self.mesage_output('save_anketa_result')
            print('save_anketa_result', ex)

    def vidpovidi(self,x):
        try:
            self.yes_no_maybe_probel += 1
            if x == 'y':
                self.oputyvana_result.append(self.pushButton_yes.text())
            elif x == 'n':
                self.oputyvana_result.append(self.pushButton_no.text())
            elif x == 'm':
                self.oputyvana_result.append(self.pushButton_mayby.text())
            elif x == 'p':
                self.oputyvana_result.append(self.pushButton_probel.text())
            if self.yes_no_maybe_probel >= len(self.list_anketa):
                self.textBrowser.setText('Дякую за відповіді!')
            else:
                self.textBrowser.setText(self.list_anketa [ self.yes_no_maybe_probel ])
        except Exception as ex:
            self.error.show()
            self.mesage_output('vidpovidi')
            print('vidpovidi', ex)

    def oputyvana_GO(self):
        try:
            self.text = self.comboBox_qestions.currentText().replace('шкл.', 'Шкільний вік').replace('дош.','Дошкільний вік').replace(' ', '_') + 'Anketa'
            connect = connect('IRCdb.db')
            with connect as con:
                cursor = con.cursor()
                cursor.execute(f"SELECT putana FROM  {self.text} ;")
                for i in cursor:
                    self.list_anketa.append(i[0])
            self.textBrowser.setText(self.list_anketa[0])
        except Exception as ex:
            self.error.show()
            self.mesage_output('oputyvana_GO')
            print('oputyvana_GO', ex)

    def oputyvana_forma(self):
        try:
            if  self.oputyvana_number==0:
                self.yes_no_maybe_probel=0
                self.list_anketa.clear()
                self.oputyvana_result.clear()
                self.tableWidget.hide()
                self.frame_qestions.show()
                self.oputyvana_number = 1
                self.oputyvana_GO()
            else:
                self.frame_qestions.hide()
                self.tableWidget.show()
                self.oputyvana_number = 0

        except Exception as ex:
            print('oputyvana_forma',ex)

    def selenium_metoduk(self,x):
        try:
            kompetencia_list = [ ]
            potrebu_list = [ ]
            recomendacii_list = [ ]
            login_user = self.login.text()
            password_user = self.password.text()
            url_user = self.edit_URL.text()
            metoduk = self.comboBox_site.currentText()

            connect = connect('IRCdb.db')

            with connect as con:
                    cursor = con.cursor()
                    cursor.execute(f'DELETE FROM irc_selenium_kompetencia;')

            with connect as con:
                    cursor = con.cursor()
                    cursor.execute(f'DELETE FROM irc_selenium_potrebu;')

            with connect as con:
                    cursor = con.cursor()
                    cursor.execute(f'DELETE FROM irc_selenium_recomendacii;')

            with connect as con:
                    cursor = con.cursor()
                    cursor.execute(f"SELECT qestion FROM  {metoduk.replace('шкл.', 'Шкільний вік').replace('дош.','Дошкільний вік').replace(' ', '_') + 'qestions'} ;")
                    for i in cursor:
                            KO_str = i
                            kompetencia = KO_str [ 0 ]
                            kompetencia_list.append(kompetencia)
            with connect as con:
                    cursor = con.cursor()
                    cursor.execute(f"SELECT potrebu FROM  {metoduk.replace('шкл.', 'Шкільний вік').replace('дош.','Дошкільний вік').replace(' ', '_') + 'potrebu'} ;")
                    for i in cursor:
                            KO_str = i
                            potrebu = KO_str [ 0 ]
                            potrebu_list.append(potrebu)
            with connect as con:
                    cursor = con.cursor()
                    cursor.execute(f"SELECT recomendacii FROM  {metoduk.replace('шкл.', 'Шкільний вік').replace('дош.','Дошкільний вік').replace(' ', '_') + 'recomendacii'} ;")
                    for i in cursor:
                            KO_str = i
                            recomendacii = KO_str [ 0 ]
                            recomendacii_list.append(recomendacii)

            if self.label_irc.isChecked():
                    with connect as con:
                            cursor = con.cursor()
                            cursor.execute(f' UPDATE irc_portal SET url_user="{url_user}" WHERE id=1;')
            else:
                    with connect as con:
                            cursor = con.cursor()
                            cursor.execute(f'DELETE FROM irc_portal;')

                    with connect as con:
                            cursor = con.cursor()
                            cursor.execute(
                                    f"""INSERT INTO irc_portal  (login_user, password_user, url_user) VALUES ('{login_user}', '{password_user}', '{url_user}') """)

            with connect as con:
                    cursor = con.cursor()
                    for i in range(len(kompetencia_list)):
                            cursor.execute(
                                    f"""INSERT INTO irc_selenium_kompetencia  (kompetencia_user) VALUES ('{kompetencia_list [ i ]}') """)

            with connect as con:
                    cursor = con.cursor()
                    for i in range(len(potrebu_list)):
                            cursor.execute(
                                    f"""INSERT INTO irc_selenium_potrebu  (potrebu_user) VALUES ('{potrebu_list [ i ]}') """)

            with connect as con:
                    cursor = con.cursor()
                    for i in range(len(recomendacii_list)):
                            cursor.execute(
                                    f"""INSERT INTO irc_selenium_recomendacii  (recomendacii_user) VALUES ('{recomendacii_list [ i ]}') """)

            if x == 'k':
                    call("selenium_web_add.exe")
            elif x == 'p':
                    call("selenium_web _potrebu_add.exe")
            elif x == 'r':
                    call( "selenium_web _recomendacii_add.exe")
        except Exception as ex:
            self.error.show()
            self.mesage_output('selenium_metoduk')
            print('selenium_metoduk', ex)

    def site_metoduk(self):
            if self.site == 0:
                    self.site = 1
                    self.frame_avtozapovnena.show()
                    self.gridLayout_2.addWidget(self.widget, int(0), int(3), int(1), int(1))
            else:
                    self.site = 0
                    self.frame_avtozapovnena.hide()
                    self.gridLayout_2.addWidget(self.widget, int(0), int(2), int(1), int(1))

    def new_window(self):
        MainWindow.hide()
        call("program_two.exe")
        MainWindow.show()

    def selenium_irc(self, x):
        try:
            kompetencia_list = [ ]
            potrebu_list = [ ]
            recomendacii_list = [ ]
            login_user = self.login.text()
            password_user = self.password.text()
            url_user = self.edit_URL.text()
            id = self.id_children_for_db()
            connect = connect('IRCdb.db')

            with connect as con:
                    cursor = con.cursor()
                    cursor.execute(f'DELETE FROM irc_selenium_kompetencia;')

            with connect as con:
                    cursor = con.cursor()
                    cursor.execute(f'DELETE FROM irc_selenium_potrebu;')

            with connect as con:
                    cursor = con.cursor()
                    cursor.execute(f'DELETE FROM irc_selenium_recomendacii;')

            with connect as con:
                    cursor = con.cursor()
                    cursor.execute(f'SELECT kompetencia FROM  answers_kompetencia WHERE id_children={id};')
                    for i in cursor:
                            KO_str = i
                            kompetencia = KO_str [ 0 ]
                            kompetencia_list.append(kompetencia)
            with connect as con:
                    cursor = con.cursor()
                    cursor.execute(f'SELECT potrebu FROM  answers_potrebu WHERE id_children={id};')
                    for i in cursor:
                            KO_str = i
                            potrebu = KO_str [ 0 ]
                            potrebu_list.append(potrebu)
            with connect as con:
                    cursor = con.cursor()
                    cursor.execute(f'SELECT recomendacii FROM  answers_recomendacii WHERE id_children={id};')
                    for i in cursor:
                            KO_str = i
                            recomendacii = KO_str [ 0 ]
                            recomendacii_list.append(recomendacii)

            if self.label_irc.isChecked():
                    with connect as con:
                            cursor = con.cursor()
                            cursor.execute(f' UPDATE irc_portal SET url_user="{url_user}" WHERE id=1;')
            else:
                    with connect as con:
                            cursor = con.cursor()
                            cursor.execute(f'DELETE FROM irc_portal;')

                    with connect as con:
                            cursor = con.cursor()
                            cursor.execute(
                                    f"""INSERT INTO irc_portal  (login_user, password_user, url_user) VALUES ('{login_user}', '{password_user}', '{url_user}') """)

            with connect as con:
                    cursor = con.cursor()
                    for i in range(len(kompetencia_list)):
                            cursor.execute(
                                    f"""INSERT INTO irc_selenium_kompetencia  (kompetencia_user) VALUES ('{kompetencia_list [ i ]}') """)

            with connect as con:
                    cursor = con.cursor()
                    for i in range(len(potrebu_list)):
                            cursor.execute(
                                    f"""INSERT INTO irc_selenium_potrebu  (potrebu_user) VALUES ('{potrebu_list [ i ]}') """)

            with connect as con:
                    cursor = con.cursor()
                    for i in range(len(recomendacii_list)):
                            cursor.execute(
                                    f"""INSERT INTO irc_selenium_recomendacii  (recomendacii_user) VALUES ('{recomendacii_list [ i ]}') """)

            if x == 'k':
                self.kompetencia_aoto_ko()
            elif x == 'p':
                self.potrebu_auto_ko()
            elif x == 'r':
                self.recomendacii_auto_ko()
        except Exception as ex:
            self.error.show()
            self.mesage_output('selenium_irc')
            print('selenium_irc', ex)

    def kompetencia_aoto_ko(self):
        browser = webdriver.Firefox()
        url_2 = None
        login = None
        password = None
        option = webdriver.FirefoxOptions()
        connect = connect('IRCdb.db')
        kompetencia_list = [ ]

        try:
            with connect as con:
                cursor = con.cursor()
                cursor.execute(f"""SELECT login_user, password_user, url_user FROM  irc_portal;""")
                user = str(cursor.fetchone())
                user_dani = user.replace(user [ 0 ], '').replace(user [ 1 ], '').replace(user [ -1 ], '').replace(",",
                                                                                                                  '')
            login, password, url_2 = user_dani.split()

            with connect as con:
                cursor = con.cursor()
                cursor.execute(f'SELECT kompetencia_user FROM  irc_selenium_kompetencia ;')
                for i in cursor:
                    KO_str = i
                    kompetencia = KO_str [ 0 ]
                    kompetencia_list.append(
                        str(kompetencia).replace("`", "'").replace(" ", "").replace("    ", "").replace("  ",
                                                                                                        "").replace(".",
                                                                                                                    "").replace(
                            "'", "").replace(",", ""))

            browser.get(url_2)
            time.sleep(2)

            login_input = browser.find_element(By.ID, "username")
            login_input.clear()
            login_input.send_keys(login)

            password_input = browser.find_element(By.ID, "password")
            password_input.clear()
            password_input.send_keys(password)

            btn_vhid = browser.find_element(By.ID, "login")
            btn_vhid.click()


            time.sleep(15)
            # btn_show_ko = browser.find_element(By.NAME, 'select-ko-button')
            # btn_show_ko.click()

            time.sleep(2)

            text_input = browser.find_elements(By.TAG_NAME, 'input')


            text_id = [ ]

            for i in range(len(text_input)):
                text_id.append(text_input [ i ].get_attribute('id'))

            text_chetbox = browser.find_elements(By.TAG_NAME, "tr")
            number = [ ]
            for i in range(len(text_chetbox)):
                text = text_chetbox [ i ].text.replace(" ", "").replace("    ", "").replace("  ", "").replace(".",
                                                                                                              "").replace(
                    "`", "").replace("'", "").replace(",", "")
                if text in kompetencia_list:
                    number.append(i)
                    kompetencia_list.remove(text)
            for i in number:
                id_click = browser.find_element(By.ID, text_id [ i  ])
                id_click.click()



        except Exception as ex:
            print(ex)

    def potrebu_auto_ko(self):
        browser = webdriver.Firefox()
        url_2 = None
        login = None
        password = None
        option = webdriver.FirefoxOptions()
        connect = connect('IRCdb.db')
        potrebu_list = [ ]

        try:
            with connect as con:
                cursor = con.cursor()
                cursor.execute(f"""SELECT login_user, password_user, url_user FROM  irc_portal;""")
                user = str(cursor.fetchone())
                user_dani = user.replace(user [ 0 ], '').replace(user [ 1 ], '').replace(user [ -1 ], '').replace(",",
                                                                                                                  '')
            login, password, url_2 = user_dani.split()

            with connect as con:
                cursor = con.cursor()
                cursor.execute(f'SELECT potrebu_user FROM  irc_selenium_potrebu ;')
                for i in cursor:
                    KO_str = i
                    potrebu = KO_str [ 0 ]
                    potrebu_list.append(
                        str(potrebu).replace("`", "'").replace(" ", "").replace("    ", "").replace("  ", "").replace(
                            ".", "").replace("'", "").replace(",", ""))
            browser.get(url_2)
            time.sleep(2)

            login_input = browser.find_element(By.ID, "username")
            login_input.clear()
            login_input.send_keys(login)

            password_input = browser.find_element(By.ID, "password")
            password_input.clear()
            password_input.send_keys(password)

            btn_vhid = browser.find_element(By.ID, "login")
            btn_vhid.click()


            time.sleep(15)
            # btn_show_ko = browser.find_element(By.NAME, 'select-ko-button')
            # btn_show_ko.click()

            text_input = browser.find_elements(By.TAG_NAME, 'input')
            text_id = [ ]
            for i in range(len(text_input)):
                text_id.append(text_input [ i ].get_attribute('id'))

            text_chetbox = browser.find_elements(By.TAG_NAME, "tr")

            number = [ ]

            for i in range(len(text_chetbox)):
                text = text_chetbox [ i ].text.replace(" ", "").replace("    ", "").replace("  ", "").replace(".",
                                                                                                              "").replace(
                    "`", "").replace("'", "").replace(",", "")
                if text in potrebu_list:
                    number.append(i)
                    potrebu_list.remove(text)
            for i in number:
                id_click = browser.find_element(By.ID, text_id [ i ])
                id_click.click()


        except Exception as ex:
            print(ex)

    def recomendacii_auto_ko(self):
        browser = webdriver.Firefox()
        url_2 = None
        login = None
        password = None
        option = webdriver.FirefoxOptions()
        connect = connect('IRCdb.db')
        recomendacii_list = [ ]

        try:
            with connect as con:
                cursor = con.cursor()
                cursor.execute(f"""SELECT login_user, password_user, url_user FROM  irc_portal;""")
                user = str(cursor.fetchone())
                user_dani = user.replace(user [ 0 ], '').replace(user [ 1 ], '').replace(user [ -1 ], '').replace(",",
                                                                                                                  '')
            login, password, url_2 = user_dani.split()

            with connect as con:
                cursor = con.cursor()
                cursor.execute(f'SELECT recomendacii_user FROM  irc_selenium_recomendacii ;')
                for i in cursor:
                    KO_str = i
                    recomendacii = KO_str [ 0 ]
                    recomendacii_list.append(
                        str(recomendacii).replace("`", "'").replace(" ", "").replace("    ", "").replace("  ",
                                                                                                         "").replace(
                            ".", "").replace("'", "").replace(",", ""))

            browser.get(url_2)
            time.sleep(2)

            login_input = browser.find_element(By.ID, "username")
            login_input.clear()
            login_input.send_keys(login)

            password_input = browser.find_element(By.ID, "password")
            password_input.clear()
            password_input.send_keys(password)

            btn_vhid = browser.find_element(By.ID, "login")
            btn_vhid.click()


            time.sleep(15)

            # btn_show_ko = browser.find_element(By.NAME, 'select-ko-button')
            # btn_show_ko.click()

            text_input = browser.find_elements(By.TAG_NAME, 'input')
            text_id = [ ]
            for i in range(len(text_input)):
                text_id.append(text_input [ i ].get_attribute('id'))

            text_chetbox = browser.find_elements(By.TAG_NAME, "tr")

            number = [ ]

            for i in range(len(text_chetbox)):
                text = text_chetbox [ i ].text.replace(" ", "").replace("    ", "").replace("  ", "").replace(".",
                                                                                                              "").replace(
                    "`", "").replace("'", "").replace(",", "")
                if text in recomendacii_list:
                    number.append(i)
                    recomendacii_list.remove(text)
            for i in number:
                id_click = browser.find_element(By.ID, text_id [ i  ])
                id_click.click()


        except Exception as ex:
            print(ex)

    def avtozapovnena(self):
        if self.avto_silenium == 0:
                self.avto_silenium = 1
                self.tableWidget.hide()
                self.frame_avto.show()
                self.login.show()
                self.password.show()
                self.label_irc.show()
                self.label_URL.show()
                self.edit_URL.show()
                self.btn_kompetencia.show()
                self.btn_potrebu.show()
                self.btn_recomendacii.show()

        else:
                self.avto_silenium = 0
                self.tableWidget.show()
                self.frame_avto.hide()
                self.login.hide()
                self.password.hide()
                self.label_irc.hide()
                self.label_URL.hide()
                self.edit_URL.hide()
                self.btn_kompetencia.hide()
                self.btn_potrebu.hide()
                self.btn_recomendacii.hide()

    def add_row_k_p_r(self):
        try:
            self.btn_delete = QtWidgets.QPushButton()
            font = QtGui.QFont()
            font.setPointSize(12)
            self.btn_delete.setStyleSheet(
                    "QPushButton{\n"
                    "    background-color: rgb(255, 255, 255);\n"
                    f"    border-radius: {3 * suma3}%;\n"
                    f"    border: 2px solid;\n"
                    "}\n"
                    "QPushButton:hover{"
                    "border-color: red;"
                    "}"
                    "\n"
                    "}\n"
    
                    "")
            self.btn_delete.setFont(font)
            self.btn_delete.setText('X')
            if self.next_previos == 0:
                    count_row = self.tableWidget.rowCount()
                    self.tableWidget.setRowCount(count_row + 1)
                    self.tableWidget.setCellWidget(count_row, 1, self.btn_delete)

            elif self.next_previos == 1:
                    count_row = self.tableWidget.rowCount()
                    self.tableWidget.setRowCount(count_row + 1)
                    self.tableWidget.setCellWidget(count_row, 1, self.btn_delete)

            elif self.next_previos == 2:
                    count_row = self.tableWidget.rowCount()
                    self.tableWidget.setRowCount(count_row + 1)
                    self.tableWidget.setCellWidget(count_row, 1, self.btn_delete)
            self.btn_delete.clicked.connect(lambda: self.mesage('del_add'))
        except Exception as ex:
            self.error.show()
            self.mesage_output('add_row_k_p_r')
            print('add_row_k_p_r', ex)

    def delete_row_k_p_r(self):
        try:
            for i in self.tableWidget.selectedIndexes():
                    index = i.row()
                    self.tableWidget.removeRow(index)
        except Exception as ex:
            self.error.show()
            self.mesage_output('delete_row_k_p_r')
            print('delete_row_k_p_r', ex)

    def mesage_output(self, x):
        try:
            msgBox = QMessageBox()
            msgBox.setWindowFlag(Qt.FramelessWindowHint)
            msgBox.setIcon(QMessageBox.Information)

            if x == 'update_post':
                msgBox.setText("Збережено!")
                msgBox.setWindowTitle("Збережено")
            elif x == 'save_word_posts':
                msgBox.setText("ERROR save_word_posts"+'\n'+'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'delete_post':
                msgBox.setText("ERROR delete_post" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x=='update_post_':
                msgBox.setText("ERROR update_post_" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x=='next_previos_post':
                msgBox.setText("ERROR next_previos_post" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x=='show_post_update':
                msgBox.setText("ERROR show_post_update" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x=='add_post_for_table':
                msgBox.setText("ERROR add_post_for_table" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x=='con_child_id':
                msgBox.setText("ERROR con_child_id" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x=='delete_post_child_id':
                msgBox.setText("ERROR delete_post_child_id" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x=='add_child_for_post':
                msgBox.setText("ERROR add_child_for_post" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'con_children_post_id':
                msgBox.setText("ERROR con_children_post_id" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'curent_id':
                msgBox.setText("ERROR curent_id" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'delete_answer_anketa':
                msgBox.setText("ERROR delete_answer_anketa" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'show_anketa':
                msgBox.setText("ERROR show_anketa" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'save_anketa_result':
                msgBox.setText("ERROR save_anketa_result" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'vidpovidi':
                msgBox.setText("ERROR vidpovidi" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'oputyvana_GO':
                msgBox.setText("ERROR oputyvana_GO" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'selenium_metoduk':
                msgBox.setText("ERROR selenium_metoduk" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'selenium_irc':
                msgBox.setText("ERROR selenium_irc" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'add_row_k_p_r':
                msgBox.setText("ERROR add_row_k_p_r" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'delete_row_k_p_r':
                msgBox.setText("ERROR delete_row_k_p_r" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'delete_child':
                msgBox.setText("ERROR delete_child" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'update_dani_children':
                msgBox.setText("ERROR update_dani_children" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'id_children_for_db':
                msgBox.setText("ERROR show_anketa" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'all_dani_children':
                msgBox.setText("ERROR all_dani_children" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'show_result_KO':
                msgBox.setText("ERROR show_result_KO" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'table_show_ko':
                msgBox.setText("ERROR table_show_ko" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'all_filter_KO':
                msgBox.setText("ERROR all_filter_KO" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'search_KO':
                msgBox.setText("ERROR search_KO" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'show_new_window':
                msgBox.setText("ERROR show_new_window" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'save_to_db':
                msgBox.setText("ERROR save_to_db" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'add_chenge_to_row':
                msgBox.setText("ERROR add_chenge_to_row" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'kompetencia_metoduk':
                msgBox.setText("ERROR kompetencia_metoduk" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'con_nazva_metoduk':
                msgBox.setText("ERROR con_nazva_metoduk" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")
            elif x == 'menu_right':
                msgBox.setText("ERROR menu_right" + '\n' + 'Перезаватажте програму')
                msgBox.setWindowTitle("ERROR")

            font = QtGui.QFont()
            font.setPointSize(14)
            msgBox.setFont(font)
            msgBox.setStyleSheet('QMessageBox{'
                                 'background-color: white;'
                                 f'border-radius: {7 * suma3}%;'
                                 'border: 2px solid;'
                                 '}'
                                 "QPushButton{min-width: 100px;}"
                                 )
            msgBox.setStandardButtons(QMessageBox.Ok | QMessageBox.Cancel)
            returnValue = msgBox.exec()
            if returnValue == QMessageBox.Ok and x == 'update_post':
                self.show_post_update()
            elif returnValue == QMessageBox.Ok and x == 'save_word_posts':
                pass
            elif   returnValue == QMessageBox.Ok and x == 'delete_post':
                pass
            elif returnValue == QMessageBox.Ok and x=='update_post_':
                pass
            elif returnValue == QMessageBox.Ok and x == 'next_previos_post':
                pass
            elif returnValue == QMessageBox.Ok and x == 'show_post_update':
                pass
            elif returnValue == QMessageBox.Ok and x == 'add_post_for_table':
                pass
            elif returnValue == QMessageBox.Ok and x == 'con_child_id':
                pass
            elif returnValue == QMessageBox.Ok and x == 'delete_post_child_id':
                pass
            elif returnValue == QMessageBox.Ok and x == 'add_child_for_post':
                pass
            elif returnValue == QMessageBox.Ok and x == 'con_children_post_id':
                pass
            elif returnValue == QMessageBox.Ok and x == 'curent_id':
                pass
            elif returnValue == QMessageBox.Ok and x == 'delete_answer_anketa':
                pass
            elif returnValue == QMessageBox.Ok and x == 'show_anketa':
                pass
            elif returnValue == QMessageBox.Ok and x == 'save_anketa_result':
                pass
            elif returnValue == QMessageBox.Ok and x == 'vidpovidi':
                pass
            elif returnValue == QMessageBox.Ok and x == 'oputyvana_GO':
                pass
            elif returnValue == QMessageBox.Ok and x == 'selenium_metoduk':
                pass
            elif returnValue == QMessageBox.Ok and x == 'selenium_irc':
                pass
            elif returnValue == QMessageBox.Ok and x == 'add_row_k_p_r':
                pass
            elif returnValue == QMessageBox.Ok and x == 'delete_row_k_p_r':
                pass
            elif returnValue == QMessageBox.Ok and x == 'delete_child':
                pass
            elif returnValue == QMessageBox.Ok and x == 'update_dani_children':
                pass
            elif returnValue == QMessageBox.Ok and x == 'id_children_for_db':
                pass
            elif returnValue == QMessageBox.Ok and x == 'all_dani_children':
                pass
            elif returnValue == QMessageBox.Ok and x == 'show_result_KO':
                pass
            elif returnValue == QMessageBox.Ok and x == 'table_show_ko':
                pass
            elif returnValue == QMessageBox.Ok and x == 'all_filter_KO':
                pass
            elif returnValue == QMessageBox.Ok and x == 'search_KO':
                pass
            elif returnValue == QMessageBox.Ok and x == 'show_new_window':
                pass
            elif returnValue == QMessageBox.Ok and x == 'save_to_db':
                pass
            elif returnValue == QMessageBox.Ok and x == 'add_chenge_to_row':
                pass
            elif returnValue == QMessageBox.Ok and x == 'kompetencia_metoduk':
                pass
            elif returnValue == QMessageBox.Ok and x == 'con_nazva_metoduk':
                pass
            elif returnValue == QMessageBox.Ok and x == 'menu_right':
                pass
        except Exception as ex:
            print('mesage_output', ex)

    def mesage(self, x):
        try:
            msgBox = QMessageBox()
            msgBox.setWindowFlag(Qt.FramelessWindowHint)
            msgBox.setIcon(QMessageBox.Information)

            if x == 'del':
                if self.comboBox_KO.currentText() == '':
                    x = 'error'
                    pass
                else:
                    msgBox.setText("Ви впевнені що бажаєте видаити КО?!" + "\n" + self.comboBox_KO.currentText())
                    msgBox.setWindowTitle("Видалення")
            if x == 'error':
                    msgBox.setText("Помилка не вибрана дитина!")
                    msgBox.setWindowTitle("Помилка")
            elif x == 'upd':
                    msgBox.setText("Ви впевнені що бажаєте змінити дані?")
                    msgBox.setWindowTitle("Змінити дані")
            elif x == 'del_add':
                    msgBox.setText("Ви впевнені що бажаєте видалити?")
                    msgBox.setWindowTitle("Змінити КО")
            elif x == 'del_k_p_r':
                    msgBox.setText("Ви впевнені що бажаєте видалити?")
                    msgBox.setWindowTitle("Змінити КО")
            elif x == 'del_anketa':
                msgBox.setText("Ви впевнені що бажаєте видалити опитування?")
            elif x =='del_child_post_id':
                if self.comboBox_childrean_all.currentText()=='':
                    msgBox.setText("Помилка, не вибрана дитина!")
                    pass
                else:
                    msgBox.setText("Ви впевнені що бажаєте видалити?")
            elif x == 'del_post_tema_all':
                if self.lineEdit_tema_2 == '':
                    msgBox.setText("Помилка, не вибраний допис!")
                    pass
                else:
                    msgBox.setText("Ви впевнені що бажаєте видалити?")

            font = QtGui.QFont()
            font.setPointSize(14)
            msgBox.setFont(font)
            msgBox.setStyleSheet('QMessageBox{'
                                 'background-color: white;'
                                 f'border-radius: {7 * suma3}%;'
                                 'border: 2px solid;'
                                 '}'
                                 "QPushButton{min-width: 100px;}"
                                 )
            msgBox.setStandardButtons(QMessageBox.Ok | QMessageBox.Cancel)
            returnValue = msgBox.exec()

            if returnValue == QMessageBox.Ok and x == 'error':
                    pass
            elif returnValue == QMessageBox.Ok and x == 'del':
                    self.delete_child()
            elif returnValue == QMessageBox.Ok and x == 'upd':
                    self.update_dani_children()
            elif returnValue == QMessageBox.Ok and x == 'del_add':
                    self.delete_row_k_p_r()
            elif returnValue == QMessageBox.Ok and x == 'del_k_p_r':
                    self.delete_row_k_p_r()
            elif returnValue == QMessageBox.Ok and x == 'del_anketa':
                self.delete_answer_anketa()
            elif returnValue == QMessageBox.Ok and x == 'del_child_post_id':
                if self.comboBox_childrean_all.currentText() == '':
                    pass
                else:
                    self.delete_post_child_id()
            elif returnValue == QMessageBox.Ok and x == 'del_post_tema_all':
                if self.lineEdit_tema_2 == '':
                    pass
                else:
                    self.delete_post()
        except Exception as ex:
            print('mesage', ex)
    # стара фунція
    # def delete_child(self):
    #     try:
    #         id = self.id_children_for_db()
    #         connect = sl.connect('IRCdb.db')
    #         with connect as con:
    #                 cursor = con.cursor()
    #                 cursor.execute(f'SELECT first_name, last_name, pobatkovi, data_edit FROM children WHERE id = {int(id)};')
    #                 text = cursor.fetchone()
    #                 first_name = text[0]
    #                 last_name = text[1]
    #                 pobatkovi= text[2]
    #                 birth_day=text[3]
    #         with connect as con:
    #                 cursor = con.cursor()
    #                 cursor.execute('pragma foreign_keys=on')
    #                 cursor.execute(f'DELETE FROM children WHERE id = {int(id)};')
    #         with connect as con:
    #             cursor = con.cursor()
    #             cursor.execute(f'DELETE FROM oputevana_resultat WHERE first_name="{first_name}" AND last_name="{last_name}" AND pobatkovi="{pobatkovi}" AND birth_day="{birth_day}";')
    #         self.all_filter_KO()
    #     except Exception as ex:
    #         self.error.show()
    #         self.mesage_output('delete_child')
    #         print('delete_child', ex)
    def delete_child(self):
        try:
            id = self.id_children_for_db()
            connect = connect('IRCdb.db')
            with connect as con:
                    cursor = con.cursor()
                    cursor.execute(f'SELECT first_name, last_name, pobatkovi, birth_day, data_edit FROM children WHERE id = {int(id)};')
                    text = cursor.fetchone()
                    first_name = text[0]
                    last_name = text[1]
                    pobatkovi= text[2]
                    birth_day=text[3]
                    data_edit=text[4]
            with connect as con:
                    cursor = con.cursor()
                    cursor.execute('pragma foreign_keys=on')
                    cursor.execute(f'DELETE FROM children WHERE id = {int(id)};')
            with connect as con:
                cursor = con.cursor()
                cursor.execute(f'DELETE FROM oputevana_resultat WHERE first_name="{first_name}" AND last_name="{last_name}" AND pobatkovi="{pobatkovi}" AND birth_day="{birth_day}"'
                               f'AND date_edit="{data_edit}";')
            self.all_filter_KO()
        except Exception as ex:
            self.error.show()
            self.mesage_output('delete_child')
            print('delete_child', ex)

    def update_dani_children(self):
        try:
                id = self.id_children_for_db()
                connect = connect('IRCdb.db')
                first_name = self.first_name.text().replace("'", "`").replace(" ", "")
                second_name = self.last_name.text().replace("'", "`").replace(" ", "")
                pobatkovi = self.pobatkovi.text().replace("'", "`").replace(" ", "")
                with connect as con:
                        cursor = con.cursor()
                        if self.first_name.text() != '' and self.last_name.text() != '' and self.pobatkovi.text() != '' and self.Perent_Kontakt_2.text() != '' and self.Pervuna_KO.isChecked() and self.label_filter_4.isChecked():
                                cursor.execute(
                                        f' UPDATE children SET first_name="{first_name}", last_name="{second_name}", pobatkovi="{pobatkovi}", povtorna_ko="Первина КО", birth_day="{self.dateEdit_2.text()}",  perent="{self.Perent_Kontakt_2.text()}" WHERE id={int(id)};')


                        elif self.first_name.text() != '' and self.last_name.text() != '' and self.pobatkovi.text() != '' and self.Perent_Kontakt_2.text() != '' and self.Povtorna_KO_2.isChecked() and self.label_filter_4.isChecked():
                                cursor.execute(
                                        f' UPDATE children SET first_name="{first_name}", last_name="{second_name}", pobatkovi="{pobatkovi}", povtorna_ko="Повторна КО", birth_day="{self.dateEdit_2.text()}",  perent="{self.Perent_Kontakt_2.text()}" WHERE id={int(id)};')


                        elif self.first_name.text() != '' and self.last_name.text() != '' and self.pobatkovi.text() != '' and self.Povtorna_KO_2.isChecked() and self.label_filter_4.isChecked():
                                cursor.execute(
                                        f' UPDATE children SET first_name="{first_name}", last_name="{second_name}", pobatkovi="{pobatkovi}", povtorna_ko="Повторна КО", birth_day="{self.dateEdit_2.text()}" WHERE id={int(id)};')


                        elif self.first_name.text() != '' and self.last_name.text() != '' and self.pobatkovi.text() != '' and self.Pervuna_KO.isChecked() and self.label_filter_4.isChecked():
                                cursor.execute(
                                        f' UPDATE children SET first_name="{first_name}", last_name="{second_name}", pobatkovi="{pobatkovi}", povtorna_ko="Первина КО", birth_day="{self.dateEdit_2.text()}" WHERE id={int(id)};')


                        elif self.first_name.text() != '' and self.last_name.text() != '' and self.pobatkovi.text() != '' and self.label_filter_4.isChecked():
                                cursor.execute(
                                        f' UPDATE children SET first_name="{first_name}", last_name="{second_name}", pobatkovi="{pobatkovi}", birth_day="{self.dateEdit_2.text()}" WHERE id={int(id)};')


                        elif self.first_name.text() != '' and self.last_name.text() != '' and self.pobatkovi.text() != '' and self.Povtorna_KO_2.isChecked():
                                cursor.execute(
                                        f' UPDATE children SET first_name="{first_name}", last_name="{second_name}", pobatkovi="{pobatkovi}", povtorna_ko="Повторна КО" WHERE id={int(id)};')


                        elif self.first_name.text() != '' and self.last_name.text() != '' and self.pobatkovi.text() != '' and self.Pervuna_KO.isChecked():
                                cursor.execute(
                                        f' UPDATE children SET first_name="{first_name}", last_name="{second_name}", pobatkovi="{pobatkovi}", povtorna_ko="Первина КО" WHERE id={int(id)};')


                        elif self.first_name.text() != '' and self.last_name.text() != '' and self.pobatkovi.text() != '' and self.Perent_Kontakt_2.text() != '':
                                cursor.execute(
                                        f' UPDATE children SET first_name="{first_name}", last_name="{second_name}", pobatkovi="{pobatkovi}", perent="{self.Perent_Kontakt_2.text()}" WHERE id={int(id)};')


                        elif self.first_name.text() != '' and self.last_name.text() != '' and self.pobatkovi.text() != '':
                                cursor.execute(
                                        f' UPDATE children SET first_name="{first_name}", last_name="{second_name}", pobatkovi="{pobatkovi}" WHERE id={int(id)};')


                        elif self.first_name.text() == '' and self.last_name.text() == '' and self.pobatkovi.text() == '' and self.Perent_Kontakt_2.text() == '' and self.Pervuna_KO.isChecked() and self.label_filter_4.isChecked():
                                cursor.execute(
                                        f' UPDATE children SET povtorna_ko="Первина КО", birth_day="{self.dateEdit_2.text()}" WHERE id={int(id)};')


                        elif self.first_name.text() == '' and self.last_name.text() == '' and self.pobatkovi.text() == '' and self.Perent_Kontakt_2.text() == '' and self.Povtorna_KO_2.isChecked() and self.label_filter_4.isChecked():
                                cursor.execute(
                                        f' UPDATE children SET povtorna_ko="Повторна КО", birth_day="{self.dateEdit_2.text()}" WHERE id={int(id)};')


                        elif self.first_name.text() == '' and self.last_name.text() == '' and self.pobatkovi.text() == '' and self.Perent_Kontakt_2.text() != '' and self.Povtorna_KO_2.isChecked() and self.label_filter_4.isChecked():
                                cursor.execute(
                                        f' UPDATE children SET povtorna_ko="Повторна КО", birth_day="{self.dateEdit_2.text()}", perent="{self.Perent_Kontakt_2.text()}" WHERE id={int(id)};')


                        elif self.first_name.text() == '' and self.last_name.text() == '' and self.pobatkovi.text() == '' and self.Perent_Kontakt_2.text() != '' and self.Pervuna_KO.isChecked() and self.label_filter_4.isChecked():
                                cursor.execute(
                                        f' UPDATE children SET povtorna_ko="Первина КО", birth_day="{self.dateEdit_2.text()}", perent="{self.Perent_Kontakt_2.text()}" WHERE id={int(id)};')


                        elif self.first_name.text() == '' and self.last_name.text() == '' and self.pobatkovi.text() == '' and self.Perent_Kontakt_2.text() != '':
                                cursor.execute(
                                        f' UPDATE children SET perent="{self.Perent_Kontakt_2.text()}" WHERE id={int(id)};')


                        elif self.first_name.text() == '' and self.last_name.text() == '' and self.pobatkovi.text() == '' and self.Perent_Kontakt_2.text() == '' and self.Pervuna_KO.isChecked():
                                cursor.execute(f' UPDATE children SET povtorna_ko="Первина КО" WHERE id={int(id)};')


                        elif self.first_name.text() == '' and self.last_name.text() == '' and self.pobatkovi.text() == '' and self.Perent_Kontakt_2.text() == '' and self.Povtorna_KO_2.isChecked():
                                cursor.execute(f' UPDATE children SET povtorna_ko="Повторна КО" WHERE id={int(id)};')


                        elif self.first_name.text() == '' and self.last_name.text() == '' and self.pobatkovi.text() == '' and self.Perent_Kontakt_2.text() == '' and self.label_filter_4.isChecked():
                                cursor.execute(
                                        f' UPDATE children SET birth_day="{self.dateEdit_2.text()}" WHERE id={int(id)};')


                self.Perent_Kontakt_2.clear()
                self.first_name.clear()
                self.last_name.clear()
                self.pobatkovi.clear()
                self.Povtorna_KO_2.setChecked(False)
                self.Pervuna_KO.setChecked(False)
                self.label_filter_4.setChecked(False)
                self.label_filter_2.setChecked(False)
                self.all_filter_KO()

        except Exception as ex:
            self.error.show()
            self.mesage_output('update_dani_children')
            print('update_dani_children: ', ex)

    def id_children_for_db(self):
        try:
            # first_name, last_name, pobatkovi = self.comboBox_KO.currentText().split()
            name_all = self.comboBox_KO.currentText().replace("/", " ")
            first_name, last_name, pobatkovi, birth, date_edit = name_all.split()
            connect = connect('IRCdb.db')
            # data_ko_search = self.dateEdit_2.text().split('.')
            # data_birth = data_ko_search[0] + '.' + data_ko_search[1] + '.' + data_ko_search[2]
            # data_edit = data_ko_search [ 2 ] + '-' + data_ko_search [ 1 ] + '-' + data_ko_search [ 0 ]

            with connect as con:
                    cursor = con.cursor()
                    # if self.label_filter_4.isChecked():
                    #     cursor.execute(f'SELECT id from children WHERE first_name="{first_name}" AND last_name="{last_name}" AND pobatkovi="{pobatkovi}" AND birth_day="{data_birth}";')
                    # elif self.label_filter_2.isChecked():
                    #     cursor.execute(f'SELECT id from children WHERE first_name="{first_name}" AND last_name="{last_name}" AND pobatkovi="{pobatkovi}" AND data_edit="{data_edit}";')
                    # else:
                    cursor.execute(f'SELECT id from children WHERE first_name="{first_name}" AND last_name="{last_name}" AND pobatkovi="{pobatkovi}"'
                                       f'AND birth_day="{birth}" AND data_edit="{date_edit}";')
                    for i in cursor:
                            string = i[0]
                            id_ = string
                            id_child = id_
                    return id_child
        except Exception as ex:
            self.error.show()
            self.mesage_output('id_children_for_db')
            print('id_children_for_db', ex)

    def all_dani_children(self):
        try:
                id_child = self.id_children_for_db()
                connect = connect('IRCdb.db')
                with connect as con:
                        cursor = con.cursor()
                        cursor.execute(f'SELECT first_name, last_name, pobatkovi, birth_day, povtorna_ko, metoduka, zaklad, perent, data_edit from children WHERE id={id_child};')
                        dani = str(cursor.fetchone()).replace("(", "").replace(")", "").replace(",", "").replace("'","").replace("КО", '')
                        first_name, last_name, pobatkovi, birth_day, povtorna_ko, metoduka, age, perent, data_edit = dani.split()

                msgBox = QMessageBox()
                msgBox.setWindowFlag(Qt.FramelessWindowHint)
                msgBox.setIcon(QMessageBox.Information)
                msgBox.setText("Призвіще: " + f"{first_name}" + "\n"
                                                                "Ім'я: " + f"{last_name}" + "\n"
                                                                                            "Побатькові: " + f"{pobatkovi}" + "\n"
                                                                                                                              "Дата народження: " + f"{birth_day}" + "\n"
                                                                                                                                                                     "КО: " + f"{povtorna_ko}" + "\n"
                                                                                                                                                                                                 "Методика: " + f"{str(metoduka).replace('_', ' ').replace('Шкільний вік', 'шкл.').replace('Дошкільний вік', 'дош.').replace('qestions', '')}" + "\n"
                                                                                                                                                                                                                                                                                                                                                 "Контакти: " + f"{perent}" + "\n"
                                                                                                                                                                                                                                                                                                                                                                              "Дата проведення: " + f"{data_edit}" + "\n"
                               )
                msgBox.setWindowTitle("Дані")
                font = QtGui.QFont()
                font.setPointSize(14)
                msgBox.setFont(font)
                msgBox.setStyleSheet('QMessageBox{'
                                     'background-color: white;'
                                     f'border-radius: {7 * suma3}%;'
                                     'border: 2px solid;'
                                     '}'
                                     "QPushButton{min-width: 800px;}"
                                     )
                msgBox.setStandardButtons(QMessageBox.Ok)
                returnValue = msgBox.exec()
        except Exception as ex:
                self.error.show()
                self.mesage_output('all_dani_children')
                print('all_dani_children: ', ex)
                self.comboBox_KO.addItem('Поле пусте!')

    def show_result_KO(self):
        try:
                if self.minus_number == 0:
                        self.minus_number = 1.00
                self.tableWidget.setColumnCount(2)
                self.tableWidget.horizontalHeader().resizeSection(0, int(520 * suma3 / self.minus_number))
                self.tableWidget.horizontalHeader().resizeSection(1, int(70 * suma3 / self.minus_number))
                item = self.tableWidget.horizontalHeaderItem(0)
                item.setText("Компетенція")
                item = self.tableWidget.horizontalHeaderItem(1)
                item.setText("del")
                # first_name, last_name, pobatkovi = self.comboBox_KO.currentText().split()
                name_all = self.comboBox_KO.currentText().replace("/"," ")
                first_name, last_name, pobatkovi, birth, date_edit = name_all.split()
                data_ko_search = self.dateEdit_2.text().split('.')
                data = data_ko_search [ 2 ] + '-' + data_ko_search [ 1 ] + '-' + data_ko_search [ 0 ]
                data_birth = data_ko_search [ 0 ] + '.' + data_ko_search [ 1 ] + '.' + data_ko_search [ 2 ]

                connect = connect('IRCdb.db')
                with connect as con:
                        cursor = con.cursor()
                        # if self.Povtorna_KO_2.isChecked() and self.label_filter_2.isChecked():
                        #         cursor.execute(
                        #                 f'SELECT id from children WHERE first_name="{first_name}" AND last_name="{last_name}"  AND pobatkovi="{pobatkovi}" AND povtorna_ko="Повторна КО" AND data_edit="{data}";')
                        # elif self.Pervuna_KO.isChecked() and self.label_filter_2.isChecked():
                        #         cursor.execute(
                        #                 f'SELECT id from children WHERE first_name="{first_name}" AND last_name="{last_name}"  AND pobatkovi="{pobatkovi}" AND povtorna_ko="Первина КО" AND data_edit="{data}";')
                        # elif self.label_filter_2.isChecked():
                        #         cursor.execute(
                        #                 f'SELECT id from children WHERE first_name="{first_name}" AND last_name="{last_name}" AND pobatkovi="{pobatkovi}" AND  data_edit="{data}";')
                        # elif self.Povtorna_KO_2.isChecked():
                        #         cursor.execute(
                        #                 f'SELECT id from children WHERE first_name="{first_name}" AND last_name="{last_name}" AND pobatkovi="{pobatkovi}" AND povtorna_ko="Повторна КО";')
                        # elif self.Pervuna_KO.isChecked():
                        #         cursor.execute(
                        #                 f'SELECT id from children WHERE first_name="{first_name}" AND last_name="{last_name}" AND pobatkovi="{pobatkovi}" AND povtorna_ko="Первина КО";')
                        # elif self.label_filter_4.isChecked():
                        #     cursor.execute(
                        #         f'SELECT id from children WHERE first_name="{first_name}" AND last_name="{last_name}" AND pobatkovi="{pobatkovi}" AND birth_day="{data_birth}";')
                        # else:
                        cursor.execute(
                                        f'SELECT id from children WHERE first_name="{first_name}" AND last_name="{last_name}" AND pobatkovi="{pobatkovi}" '
                                        f'AND birth_day="{birth}" AND data_edit="{date_edit}" ;')
                        for i in cursor:
                                string = str(i)
                                id_ = int(string.replace(string [ 0 ], '').replace(string [ -1 ], '').replace(
                                        string [ -2 ], ''))
                                self.id_child = id_
                                self.table_show_ko('0')
        except Exception as ex:
            self.error.show()
            self.mesage_output('show_result_KO')
            print('show_result_KO:', ex)

    def table_show_ko(self, x):
        try:
                connect = connect('IRCdb.db')
                if x == '+':
                        self.next_previos += 1
                elif x == '-':
                        self.next_previos -= 1
                if self.next_previos > 2:
                        self.next_previos = 2
                elif self.next_previos < 0:
                        self.next_previos = 0
                with connect as con:
                        cursor = con.cursor()
                        if self.next_previos == 0:
                                item = self.tableWidget.horizontalHeaderItem(0)
                                item.setText("Компетенція")
                                cursor.execute(
                                        f'SELECT COUNT(1) FROM  answers_kompetencia WHERE id_children={self.id_child};')
                        elif self.next_previos == 1:
                                item = self.tableWidget.horizontalHeaderItem(0)
                                item.setText("Потреби")
                                cursor.execute(
                                        f'SELECT COUNT(1) FROM  answers_potrebu WHERE id_children={self.id_child};')
                        elif self.next_previos == 2:
                                item = self.tableWidget.horizontalHeaderItem(0)
                                item.setText("Рекомендації")
                                cursor.execute(
                                        f'SELECT COUNT(1) FROM  answers_recomendacii WHERE id_children={self.id_child};')
                        count = str(cursor.fetchone())
                        count_kompetencia = int(
                                count.replace(count [ 0 ], '').replace(count [ -1 ], '').replace(count [ -2 ], ''))
                        self.tableWidget.setRowCount(count_kompetencia)
                row_k = 0
                with connect as con:
                        cursor = con.cursor()
                        if self.next_previos == 0:
                                cursor.execute(
                                        f'SELECT kompetencia FROM  answers_kompetencia WHERE id_children={self.id_child};')
                        elif self.next_previos == 1:
                                cursor.execute(
                                        f'SELECT potrebu FROM  answers_potrebu WHERE id_children={self.id_child};')
                        elif self.next_previos == 2:
                                cursor.execute(
                                        f'SELECT recomendacii FROM  answers_recomendacii WHERE id_children={self.id_child};')
                        for i in cursor:
                                KO_str = i
                                KO = KO_str [ 0 ]
                                self.tableWidget.setItem(row_k, 0, QTableWidgetItem(KO))

                                self.btn_delete = QtWidgets.QPushButton()
                                font = QtGui.QFont()
                                font.setPointSize(12)
                                self.btn_delete.setStyleSheet(
                                        "QPushButton{\n"
                                        "    background-color: rgb(255, 255, 255);\n"
                                        f"    border-radius: {3 * suma3}%;\n"
                                        f"    border: 2px solid;\n"
                                        "}\n"
                                        "QPushButton:hover{"
                                        "border-color: red;"
                                        "}"
                                        "\n"
                                        "}\n"

                                        "")
                                self.btn_delete.setFont(font)
                                self.btn_delete.setText('X')
                                self.tableWidget.setCellWidget(row_k, 1, self.btn_delete)
                                row_k += 1
                                self.btn_delete.clicked.connect(lambda: self.mesage('del_k_p_r'))
        except Exception as ex:
                self.error.show()
                self.mesage_output('table_show_ko')
                self.mesage('error')
                print('table_show_ko: ', ex)

    def table_restart(self):
        if self.minus_number == 0:
                self.minus_number = 1.00
        self.tableWidget.setColumnCount(2)
        self.tableWidget.setRowCount(1)
        item = QtWidgets.QTableWidgetItem()
        font = QtGui.QFont()
        font.setPointSize(14)
        font.setBold(True)
        font.setWeight(75)
        item.setFont(font)
        self.tableWidget.setVerticalHeaderItem(0, item)
        item = QtWidgets.QTableWidgetItem()
        font = QtGui.QFont()
        font.setPointSize(14)
        font.setBold(True)
        font.setWeight(75)
        item.setFont(font)
        self.tableWidget.setHorizontalHeaderItem(0, item)
        item = QtWidgets.QTableWidgetItem()
        font = QtGui.QFont()
        font.setPointSize(14)
        font.setBold(True)
        font.setWeight(75)
        item.setFont(font)
        self.tableWidget.setHorizontalHeaderItem(1, item)
        item = QtWidgets.QTableWidgetItem()
        font = QtGui.QFont()
        font.setPointSize(14)
        font.setBold(True)
        font.setWeight(75)
        item.setFont(font)
        self.tableWidget.setHorizontalHeaderItem(2, item)
        self.tableWidget.horizontalHeader().resizeSection(0, int(450 * suma3 / self.minus_number))
        self.tableWidget.horizontalHeader().resizeSection(1, int(140 * suma3 / self.minus_number))

        item = self.tableWidget.horizontalHeaderItem(0)
        item.setText("Питання")
        item = self.tableWidget.horizontalHeaderItem(1)
        item.setText('Відповіді')
        item = self.tableWidget.horizontalHeaderItem(2)

    def all_filter_KO(self):
        try:
            self.comboBox_KO.clear()
            data_ko_search = self.dateEdit_2.text().split('.')
            data = data_ko_search [ 2 ] + '-' + data_ko_search [ 1 ] + '-' + data_ko_search [ 0 ]
            data_birth = data_ko_search [ 0 ] + '.' + data_ko_search [ 1 ] + '.' + data_ko_search [ 2 ]

            connect = connect('IRCdb.db')
            with connect as con:
                    cursor = con.cursor()
                    if self.Pervuna_KO.isChecked() and self.label_filter_2.isChecked():
                            cursor.execute(
                                    f'SELECT first_name, last_name, pobatkovi, birth_day, data_edit from children WHERE povtorna_ko="Первина КО" AND data_edit="{data}";')
                    elif self.Povtorna_KO_2.isChecked() and self.label_filter_2.isChecked():
                            cursor.execute(
                                    f'SELECT first_name, last_name, pobatkovi, birth_day, data_edit from children WHERE povtorna_ko="Повторна КО" AND data_edit="{data}";')
                    elif self.Povtorna_KO_2.isChecked():
                            cursor.execute(
                                    f'SELECT first_name, last_name, pobatkovi, birth_day, data_edit from children WHERE povtorna_ko="Повторна КО";')
                    elif self.Pervuna_KO.isChecked():
                            cursor.execute(
                                    f'SELECT first_name, last_name, pobatkovi, birth_day, data_edit from children WHERE povtorna_ko="Первина КО";')
                    elif self.label_filter_2.isChecked():
                            cursor.execute(
                                    f'SELECT first_name, last_name, pobatkovi, birth_day, data_edit from children WHERE data_edit="{data}";')
                    elif self.label_filter_4.isChecked():
                        cursor.execute(
                            f'SELECT first_name, last_name, pobatkovi, birth_day, data_edit from children WHERE birth_day="{data_birth}";')
                    else:
                            cursor.execute(f'SELECT first_name, last_name, pobatkovi, birth_day, data_edit from children;')

                    for i in cursor:
                            child = str(i).replace('(', '').replace(')', '').replace("'", '').replace(',', '')
                            first_name, last_name, pobatkovi, birth, date_edit = child.split()
                            # self.comboBox_KO.addItem(child)
                            self.comboBox_KO.addItem(first_name+' '+last_name+' '+pobatkovi+'\n'+birth+'/'+date_edit+'\n')

            data_ko_search.clear()
        except Exception as ex:
            self.error.show()
            self.mesage_output('all_filter_KO')
            print('all_filter_KO', ex)

    def search_KO(self):
        try:
            self.comboBox_KO.clear()
            connect = connect('IRCdb.db')
            first_name = self.first_name.text().replace("'", "`")
            second_name = self.last_name.text().replace("'", "`")
            pobatkovi = self.pobatkovi.text().replace("'", "`")
            data_ko_search = self.dateEdit_2.text().split('.')
            data = data_ko_search [ 2 ] + '-' + data_ko_search [ 1 ] + '-' + data_ko_search [ 0 ]
            with connect as con:
                    cursor = con.cursor()
                    if self.Povtorna_KO_2.isChecked() and self.label_filter_2.isChecked() and second_name == '' and pobatkovi == '':
                            cursor.execute(
                                    f'SELECT first_name, last_name, pobatkovi, birth_day, data_edit from children WHERE first_name="{first_name}" AND povtorna_ko="Повторна КО" AND data_edit="{data}";')
                    elif self.Povtorna_KO_2.isChecked() and self.label_filter_2.isChecked() and first_name == '' and pobatkovi == '':
                            cursor.execute(
                                    f'SELECT first_name, last_name, pobatkovi, birth_day, data_edit from children WHERE last_name="{second_name}" AND povtorna_ko="Повторна КО" AND data_edit="{data}";')
                    elif self.Povtorna_KO_2.isChecked() and self.label_filter_2.isChecked() and pobatkovi == '':
                            cursor.execute(
                                    f'SELECT first_name, last_name, pobatkovi, birth_day, data_edit from children WHERE first_name="{first_name}" AND last_name="{second_name}" AND povtorna_ko="Повторна КО" AND data_edit="{data}";')
                    elif second_name == '' and pobatkovi == '':
                            if self.Povtorna_KO_2.isChecked():
                                    cursor.execute(
                                            f'SELECT first_name, last_name, pobatkovi, birth_day, data_edit from children WHERE first_name="{first_name}" AND povtorna_ko="Повторна КО" ;')
                            elif self.label_filter_2.isChecked():
                                    cursor.execute(
                                            f'SELECT first_name, last_name, pobatkovi, birth_day, data_edit from children WHERE first_name="{first_name}" AND data_edit="{data}" ;')
                            else:
                                    cursor.execute(
                                            f'SELECT first_name, last_name, pobatkovi, birth_day, data_edit from children WHERE first_name="{first_name}";')
                    elif first_name == '' and pobatkovi == '':
                            if self.Povtorna_KO_2.isChecked():
                                    cursor.execute(
                                            f'SELECT first_name, last_name, pobatkovi, birth_day, data_edit from children WHERE last_name="{second_name}" AND povtorna_ko="Повторна КО" ;')
                            elif self.label_filter_2.isChecked():
                                    cursor.execute(
                                            f'SELECT first_name, last_name, pobatkovi, birth_day, data_edit from children WHERE last_name="{second_name}" AND data_edit="{data}" ;')
                            else:
                                    cursor.execute(
                                            f'SELECT first_name, last_name, pobatkovi, birth_day, data_edit from children WHERE last_name="{second_name}";')
                    elif pobatkovi == '':
                            if self.Povtorna_KO_2.isChecked():
                                    cursor.execute(
                                            f'SELECT first_name, last_name, pobatkovi, birth_day, data_edit from children WHERE first_name="{first_name}" AND last_name="{second_name}" AND povtorna_ko="Повторна КО";')
                            elif self.label_filter_2.isChecked():
                                    cursor.execute(
                                            f'SELECT first_name, last_name, pobatkovi, birth_day, data_edit from children WHERE first_name="{first_name}" AND last_name="{second_name}" AND data_edit="{data}";')
                            else:
                                    cursor.execute(
                                            f'SELECT first_name, last_name, pobatkovi, birth_day, data_edit from children WHERE first_name="{first_name}" AND last_name="{second_name}" ;')
                    else:
                            if self.Povtorna_KO_2.isChecked() and self.label_filter_2.isChecked():
                                    cursor.execute(
                                            f'SELECT first_name, last_name, pobatkovi, birth_day, data_edit from children WHERE first_name="{first_name}" AND last_name="{second_name}"  AND pobatkovi="{pobatkovi}" AND povtorna_ko="Повторна КО" AND data_edit="{data}";')
                            elif self.Povtorna_KO_2.isChecked():
                                    cursor.execute(
                                            f'SELECT first_name, last_name, pobatkovi, birth_day, data_edit from children WHERE first_name="{first_name}" AND last_name="{second_name}"  AND pobatkovi="{pobatkovi}" AND povtorna_ko="Повторна КО";')
                            elif self.label_filter_2.isChecked():
                                    cursor.execute(
                                            f'SELECT first_name, last_name, pobatkovi, birth_day, data_edit from children WHERE first_name="{first_name}" AND last_name="{second_name}"  AND pobatkovi="{pobatkovi}" AND  data_edit="{data}";')
                            else:
                                    cursor.execute(
                                            f'SELECT first_name, last_name, pobatkovi, birth_day, data_edit from children WHERE first_name="{first_name}" AND last_name="{second_name}"  AND pobatkovi="{pobatkovi}";')
                    data_ko_search.clear()
                    for i in cursor:
                            child = str(i).replace('(', '').replace(')', '').replace("'", '').replace(',', '')
                            first_name, last_name, pobatkovi, birth, date_edit = child.split()
                            self.comboBox_KO.addItem(first_name+' '+last_name+' '+pobatkovi+'\n'+birth+'/'+date_edit+'\n')
                            # self.comboBox_KO.addItem(child)
        except Exception as ex:
            self.error.show()
            self.mesage_output('search_KO')
            print('search_KO', ex)

    def filter_bar(self):
        if self.minus_number == 0:
                self.minus_number = 1
        if self.filter_number == 0:
                self.hide_widget()
                self.message.setIcon(QIcon('chenge_rgb.png'))
                self.message.setIconSize(
                        QSize(int(40 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
                self.filter_number+=1
                self.frame_two.show()
                self.next.show()
                self.previous.show()
                self.plus.hide()
                self.minus.hide()


        else:
                self.message.setIcon(QIcon('chenge.png'))
                self.message.setIconSize(
                        QSize(int(40 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
                self.filter_number = 0
                self.frame_two.hide()
                self.next.hide()
                self.previous.hide()
                self.plus.show()
                self.minus.show()
                self.frame_avtozapovnena.hide()
                self.site=0

    def show_new_window(self):
        try:
            if self.minus_number == 0:
                self.minus_number = 1
                self.w = MyWin()
                self.w.show()
                self.add_metoduk.setIcon(QIcon('2rgb.png'))
                self.add_metoduk.setIconSize(QSize(int(45 * suma3 / self.minus_number), int(45 * suma3 / self.minus_number)))
                MainWindow.hide()

                def back():
                        if self.minus_number == 0:
                                self.minus_number = 1
                        MainWindow.show()
                        self.add_metoduk.setGeometry(
                                QtCore.QRect(int(4 * suma3 / self.minus_number), int(62 * suma3 / self.minus_number),
                                             int(55 * suma3 / self.minus_number), int(55 * suma3 / self.minus_number)))
                        self.add_metoduk.setIcon(QIcon('2.png'))

                self.w.pushButton.clicked.connect(back)
        except Exception as ex:
            self.error.show()
            self.mesage_output('show_new_window')
            print('show_new_window', ex)

    def save_to_db(self):
        try:
                if self.menu_right_count == 1:
                        self.list_checket_variant = [ ]
                        self.potrebu = [ ]
                        self.recomendacii = [ ]
                        self.id_grup_qestion = [ ]
                        first_name, second_name, pobatkovi = self.PIP.text().split()
                        data_birth_day = self.dateEdit.text()
                        povtorna = None
                        if self.Povtorna_KO.isChecked():
                                povtorna = self.Povtorna_KO.text()
                        else:
                                povtorna = 'Первина КО'
                        age = None
                        if self.School.isChecked():
                                age = self.School.text().replace(" ", "_")
                        elif self.Doschilnul.isChecked():
                                age = self.Doschilnul.text().replace(" ", "_")
                        text_kompetencia = self.comboBox.currentText().replace('шкл.', 'Шкільний вік').replace('дош.','Дошкільний вік').replace(' ', '_') + 'qestions'
                        kontacktu = self.Perent_Kontakt.text()
                        connect = connect('IRCdb.db')
                        with connect as con:
                                cursor = con.cursor()
                                cursor.execute(f"SELECT date('now');")
                                data = str(cursor.fetchone())
                                data_ko = data.replace(data [ 0 ], '').replace(data [ -1 ], '').replace(data [ -2 ],'').replace(data [ 1 ], '')

                        c = self.tableWidget.rowCount()
                        for i in range(c):
                                if self.tableWidget.cellWidget(i, 1).checkState():
                                        text_table_is_checket = self.tableWidget.item(i, 0).text().replace("'", "`")
                                        self.list_checket_variant.append(text_table_is_checket)
                                        check = QtWidgets.QCheckBox()
                                        check.setChecked(False)
                                        check.setStyleSheet("""
                                                                                                                                              QCheckBox:hover{
                                                                                                                                                  background-color: rgb(41, 213, 202);
                                                                                                                                              }
                                                                                                                                              QCheckBox::indicator {
                                                                                                                                                  width: 80px;
                                                                                                                                                  height: 30px;
                                                                                                                                              }
                                                                                                                                          """)
                                        self.tableWidget.setCellWidget(i, 1, check)

                        for i in range(len(self.list_checket_variant)):
                                with connect as con:
                                        cursor = con.cursor()
                                        cursor.execute(f'SELECT grup_id from {text_kompetencia} WHERE qestion="{self.list_checket_variant [ i ]}";')
                                        for i in cursor:
                                                string = str(i)
                                                grup_id = int(string.replace(string [ 0 ], '').replace(string [ -1 ],'').replace(
                                                        string [ -2 ], ''))
                                                self.id_grup_qestion.append(grup_id)

                        for i in list(set(self.id_grup_qestion)):
                                with connect as con:
                                        cursor = con.cursor()
                                        cursor.execute(
                                                f'SELECT potrebu from {text_kompetencia.replace("qestions", "potrebu")} WHERE grup_id={i};')
                                        for i in cursor:
                                                potrebu = i [ 0 ]
                                                self.potrebu.append(potrebu)

                        for i in list(set(self.id_grup_qestion)):
                                with connect as con:
                                        cursor = con.cursor()
                                        cursor.execute(
                                                f'SELECT recomendacii from {text_kompetencia.replace("qestions", "recomendacii")} WHERE grup_id={i};')
                                        for i in cursor:
                                                recomendacii = i [ 0 ]
                                                self.recomendacii.append(recomendacii)

                        with connect as con:
                                cursor = con.cursor()
                                cursor.execute(
                                        f"""INSERT INTO  children (first_name, last_name, pobatkovi, birth_day, povtorna_ko, metoduka, zaklad, perent, data_edit) VALUES ('{first_name.replace("'", "`")}','{second_name.replace("'", "`")}', '{pobatkovi.replace("'", "`")}', '{data_birth_day}',
                     '{povtorna}', '{text_kompetencia}','{age}','{kontacktu}','{data_ko}') """)

                        id_pip = None
                        with connect as con:
                                cursor = con.cursor()
                                cursor.execute(
                                        f"""SELECT id from children WHERE first_name='{first_name.replace("'","`")}' AND last_name='{second_name.replace("'","`")}' AND pobatkovi='{pobatkovi.replace("'","`")}';""" )
                                for i in cursor:
                                        id_ = str(i)
                                        id_pip = int(id_.replace(id_ [ 0 ], '').replace(id_ [ -1 ], '').replace(id_ [ -2 ],''))

                        with connect as con:
                                cursor = con.cursor()
                                for i in range(len(self.list_checket_variant)):
                                        cursor.execute(
                                                f"""INSERT INTO  answers_kompetencia (kompetencia, id_children) VALUES ('{self.list_checket_variant [ i ]}', {id_pip}) """)
                        with connect as con:
                                cursor = con.cursor()
                                for i in range(len(self.potrebu)):
                                        cursor.execute(
                                                f"""INSERT INTO  answers_potrebu (potrebu, id_children) VALUES ('{self.potrebu [ i ]}', {id_pip}) """)
                        with connect as con:
                                cursor = con.cursor()
                                for i in range(len(self.recomendacii)):
                                        cursor.execute(
                                                f"""INSERT INTO  answers_recomendacii (recomendacii, id_children) VALUES ('{self.recomendacii [ i ]}', {id_pip}) """)

                        self.tableWidget.setRowCount(1)
                        self.tableWidget.setItem(0, 0, QTableWidgetItem('Дані збережено!'))
                elif self.filter_number == 1:
                        id = self.id_children_for_db()
                        connect = connect('IRCdb.db')

                        if self.next_previos == 0:

                                with connect as con:
                                        cursor = con.cursor()
                                        cursor.execute(
                                                f'DELETE FROM answers_kompetencia WHERE id_children = {int(id)};')
                                count_row = self.tableWidget.rowCount()
                                with connect as con:
                                        cursor = con.cursor()
                                        for i in range(int(count_row)):
                                                text_table_kompetencia = self.tableWidget.item(i, 0).text().replace("'","`")
                                                cursor.execute(
                                                        f"""INSERT INTO answers_kompetencia  (kompetencia, id_children) VALUES ('{text_table_kompetencia}', {int(id)}) """)
                        elif self.next_previos == 1:

                                with connect as con:
                                        cursor = con.cursor()
                                        cursor.execute(f'DELETE FROM answers_potrebu WHERE id_children = {int(id)};')
                                count_row = self.tableWidget.rowCount()
                                with connect as con:
                                        cursor = con.cursor()
                                        for i in range(int(count_row)):
                                                text_table_potrebu = self.tableWidget.item(i, 0).text().replace("'","`")
                                                cursor.execute(
                                                        f"""INSERT INTO answers_potrebu  (potrebu, id_children) VALUES ('{text_table_potrebu}', {int(id)}) """)
                        elif self.next_previos == 2:
                                with connect as con:
                                        cursor = con.cursor()
                                        cursor.execute(
                                                f'DELETE FROM answers_recomendacii WHERE id_children = {int(id)};')
                                count_row = self.tableWidget.rowCount()
                                with connect as con:
                                        cursor = con.cursor()
                                        for i in range(int(count_row)):
                                                text_table_recomendacii = self.tableWidget.item(i, 0).text().replace("'","`")
                                                cursor.execute(
                                                        f"""INSERT INTO answers_recomendacii  (recomendacii, id_children) VALUES ('{text_table_recomendacii}', {int(id)}) """)
                        self.tableWidget.setRowCount(1)
                        self.tableWidget.setItem(0, 0, QTableWidgetItem('Успішно збережено!'))
                elif self.post==1:
                    self.add_post_for_table()
                elif self.post_update==1:
                    self.update_post_cili()
        except Exception as ex:
            self.error.show()
            self.mesage_output('save_to_db')
            print('save_to_db:', ex)

    def add_chenge_to_row(self):
        try:
            c = self.tableWidget.rowCount()
            self.tableWidget.setRowCount(c + 1)
            self.tableWidget.setItem(c, 0, QTableWidgetItem(' '))
            check = QtWidgets.QCheckBox()
            check.setStyleSheet("""
                                    QCheckBox:hover{
                                        background-color: rgb(41, 213, 202);
                                    }
                                    QCheckBox::indicator {
                                        width: 80px;
                                        height: 30px;
                                    }
                                """)
            self.tableWidget.setCellWidget(c, 1, check)
        except Exception as ex:
            self.error.show()
            self.mesage_output('add_chenge_to_row')
            print('add_chenge_to_row', ex)

    def kompetencia_metoduk(self):
        try:
                text_kompetencia = self.comboBox.currentText().replace('шкл.', 'Шкільний вік').replace('дош.','Дошкільний вік').replace(' ', '_') + 'qestions'
                connect = connect('IRCdb.db')
                with connect as con:
                        cursor = con.cursor()
                        cursor.execute(f'SELECT COUNT(1) FROM  {text_kompetencia};')
                        count = str(cursor.fetchone())
                        count_rows = count.replace(count [ 0 ], '').replace(count [ -1 ], '').replace(count [ -2 ], '')
                        for i in range(int(count_rows) + 1):
                                self.tableWidget.setRowCount(i)
                        cursor.execute(f'SELECT qestion FROM  {text_kompetencia};')
                        c = 0
                        for i in cursor:
                                string = i
                                grup = string [ 0 ]
                                self.tableWidget.setItem(c, 0, QTableWidgetItem(grup))
                                self.CheckBox = QtWidgets.QCheckBox()
                                self.CheckBox.setStyleSheet("""
                                                                     QCheckBox:hover{
                                                                         background-color: rgb(41, 213, 202);
                                                                     }
                                                                     QCheckBox::indicator {
                                                                         width: 80px;
                                                                         height: 30px;
                                                                     }
                                                                 """)

                                self.tableWidget.setCellWidget(c, 1, self.CheckBox)
                                c += 1
        except Exception as ex:
            self.error.show()
            self.mesage_output('kompetencia_metoduk')
            print('kompetencia_metoduk:', ex)

    def con_nazva_metoduk(self):
        try:
            self.comboBox.clear()
            connect = connect('IRCdb.db')
            with connect as con:
                    cursor = con.cursor()
                    cursor.execute('SELECT name from sqlite_master where type= "table"')
                    for i in cursor:
                            table = str(i).replace('(', '').replace(')', '').replace("'", '').replace(',', '')
                            if 'qestions' in table:
                                    if self.site==1:
                                            self.comboBox_site.addItem(table.replace('_', ' ').replace('qestions', '').replace('Шкільний вік','шкл.').replace('Дошкільний вік', 'дош.'))
                                    else:
                                            self.comboBox.addItem(table.replace('_', ' ').replace('qestions', '').replace('Шкільний вік','шкл.').replace('Дошкільний вік', 'дош.'))
        except Exception as ex:
            self.error.show()
            self.mesage_output('con_nazva_metoduk')
            print('con_nazva_metoduk', ex)

    def menu_right(self):
        try:
            self.table_restart()
            row_result = self.tableWidget.rowCount()
            for i in range(row_result):
                    self.tableWidget.removeRow(i)
            if self.minus_number == 0:
                    self.minus_number = 1
            if self.menu_right_count == 0:
                self.comboBox_qestions.clear()
                connect = connect('IRCdb.db')
                with connect as con:
                    cursor = con.cursor()
                    cursor.execute('SELECT name from sqlite_master where type= "table"')
                    for i in cursor:
                        table = i [ 0 ]
                        if 'Anketa' in table:
                            self.comboBox_qestions.addItem(table.replace('_', ' ').replace('Anketa', '').replace('Шкільний вік', 'шкл.').replace('Дошкільний вік', 'дош.'))

                    self.hide_widget()
                    self.add_menu.setIcon(QIcon('1argb.png'))
                    self.con_nazva_metoduk()
                    self.frame.show()
                    self.menu_right_count += 1

            else:
                    self.add_menu.setIcon(QIcon('1.png'))
                    self.add_menu.setIconSize(QSize(int(45 * suma3 / Ui_MainWindow.number), int(45 * suma3 / Ui_MainWindow.number)))
                    self.menu_right_count = 0
                    self.frame.hide()
        except Exception as ex:
            self.error.show()
            self.mesage_output('menu_right')
            print('menu_right', ex)

    def minus_rozmir(self, n):
        if n == '-':
                if Ui_MainWindow.number == 0:
                        Ui_MainWindow.number = 1.02
                elif Ui_MainWindow.number >= 0:
                        Ui_MainWindow.number += 0.02
        if n == '+':
                if Ui_MainWindow.number == 0:
                        Ui_MainWindow.number = 1.00

                elif Ui_MainWindow.number >= 0:
                        Ui_MainWindow.number -= 0.02

        self.setupUi(MainWindow)

    def hide_widget(self):
            if self.menu_right_count == 0:
                    self.frame_two.hide()
                    self.filter_number=0
                    self.post = 0
                    self.message.setIcon(QIcon('chenge.png'))
                    self.message.setIconSize(QSize(int(40 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
                    self.next.hide()
                    self.previous.hide()
                    self.plus.show()
                    self.minus.show()
                    self.frame_create_gurnal.hide()
                    self.frame_megassine.hide()
                    self.add_post.setIcon(QIcon('ж.png'))
                    self.add_post.setIconSize(QSize(int(40 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
                    self.tableWidget.show()
                    self.post_update = 0
                    self.frame_gurnal_update.hide()
                    self.next_2.hide()
                    self.previous_2.hide()

            elif self.filter_number==0:
                    self.frame.hide()
                    self.menu_right_count=0
                    self.post = 0
                    self.add_menu.setIcon(QIcon('1.png'))
                    self.add_menu.setIconSize(QSize(int(45 * suma3 / Ui_MainWindow.number), int(45 * suma3 / Ui_MainWindow.number)))
                    self.frame_create_gurnal.hide()
                    self.frame_megassine.hide()
                    self.add_post.setIcon(QIcon('ж.png'))
                    self.add_post.setIconSize(
                        QSize(int(40 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
                    self.tableWidget.show()
                    self.post_update = 0
                    self.frame_gurnal_update.hide()
                    self.next_2.hide()
                    self.previous_2.hide()
            elif self.post==0:
                self.frame_two.hide()
                self.filter_number = 0
                self.message.setIcon(QIcon('chenge.png'))
                self.message.setIconSize(QSize(int(40 * suma3 / Ui_MainWindow.number), int(40 * suma3 / Ui_MainWindow.number)))
                self.next.hide()
                self.previous.hide()
                self.plus.show()
                self.minus.show()
                self.frame_create_gurnal.hide()
                self.frame_megassine.hide()
                self.frame.hide()
                self.menu_right_count = 0
                self.add_menu.setIcon(QIcon('1.png'))
                self.add_menu.setIconSize(QSize(int(45 * suma3 / Ui_MainWindow.number), int(45 * suma3 / Ui_MainWindow.number)))
                self.frame_create_gurnal.hide()
                self.frame_megassine.hide()
                self.post_update = 0
                self.frame_gurnal_update.hide()
                self.next_2.hide()
                self.previous_2.hide()



if __name__ == "__main__":
        import sys
        app = QtWidgets.QApplication(sys.argv)
        app.setWindowIcon(QIcon('LOGO.png'))

        MainWindow = QtWidgets.QMainWindow()
        ui = Ui_MainWindow()
        ui.setupUi(MainWindow)
        MainWindow.show()
        sys.exit(app.exec_())